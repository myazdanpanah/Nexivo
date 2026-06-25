# -*- coding: utf-8 -*-

import pandas as pd
import re
import os
import sys
import jdatetime
import json
import hashlib
import pyodbc
import shutil
import time
import subprocess
import tempfile
import abc
import datetime
import logging
import traceback
import webbrowser
import requests
from tkinter import filedialog, messagebox, ttk
import tkinter as tk
import customtkinter as ctk
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from docxtpl import DocxTemplate
from pathlib import Path
from jinja2 import Environment

logging.basicConfig(filename='app_errors.log', level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("green")

COLOR_PRIMARY = "#529e98"
COLOR_ACCENT = "#f16724"
COLOR_COMPANY = "#4CAF50"
COLOR_HELP = "#2196F3"
COLOR_ABOUT = "#9C27B0"
COLOR_SETTINGS = "#FF9800"

MONTH_LIST = ["فروردین", "اردیبهشت", "خرداد", "تیر", "مرداد", "شهریور",
              "مهر", "آبان", "آذر", "دی", "بهمن", "اسفند"]
SERVICE_TYPES = ["ویزا", "اتوبوس", "قطار", "گشت", "تور", "CIP"]

# ==================== توابع کمکی ====================
def to_persian_digits(text):
    persian_digits = "۰۱۲۳۴۵۶۷۸۹"
    english_digits = "0123456789"
    trans_table = str.maketrans(english_digits, persian_digits)
    return text.translate(trans_table)

def to_english_digits(text):
    persian_digits = "۰۱۲۳۴۵۶۷۸۹"
    english_digits = "0123456789"
    trans_table = str.maketrans(persian_digits, english_digits)
    return text.translate(trans_table)

def persian_number(value):
    return to_persian_digits(str(value))

def english_number(value):
    return to_english_digits(str(value))

def persian_to_english_number(value):
    if pd.isna(value):
        return 0
    value = str(value)
    persian_digits = "۰۱۲۳۴۵۶۷۸۹"
    english_digits = "0123456789"
    for p, e in zip(persian_digits, english_digits):
        value = value.replace(p, e)
    value = value.replace(",", "")
    value = re.sub(r"[^\d]", "", value)
    return int(value) if value else 0

def persian_date_to_gregorian(date_str):
    """
    تبدیل تاریخ شمسی (با ارقام فارسی یا انگلیسی) به میلادی با فرمت YYYY-MM-DD
    مثال: "1404/01/15" یا "۱۴۰۴/۰۱/۱۵" -> "2025-04-04"
    """
    if not date_str or date_str.strip() == "":
        return None
    # حذف فاصله‌ها و تبدیل ارقام فارسی به انگلیسی
    date_str = date_str.strip()
    # تبدیل ارقام فارسی به انگلیسی
    persian_digits = "۰۱۲۳۴۵۶۷۸۹"
    english_digits = "0123456789"
    trans_table = str.maketrans(persian_digits, english_digits)
    date_str = date_str.translate(trans_table)
    # حذف جداکننده‌های غیرمجاز (فقط اسلش مجاز است)
    if '/' not in date_str:
        return None
    parts = date_str.split('/')
    if len(parts) != 3:
        return None
    try:
        year = int(parts[0])
        month = int(parts[1])
        day = int(parts[2])
        # بررسی محدوده سال شمسی (۱۳۰۰ تا ۱۵۰۰)
        if not (1300 <= year <= 1500):
            return None
        greg_date = jdatetime.date(year, month, day).togregorian()
        return greg_date.strftime("%Y-%m-%d")
    except Exception:
        return None

def gregorian_to_persian(date_value):
    """تبدیل تاریخ میلادی (date object یا string) به رشته شمسی با فرمت YYYY/MM/DD"""
    if date_value is None:
        return ""
    try:
        # اگر از نوع datetime.date یا datetime.datetime بود
        if hasattr(date_value, 'year'):
            greg_date = date_value
        else:
            # اگر string بود مانند "2026-05-05"
            date_str = str(date_value)
            if '-' in date_str:
                parts = date_str.split('-')
                greg_date = datetime.date(int(parts[0]), int(parts[1]), int(parts[2]))
            else:
                return date_str
        # تبدیل به شمسی
        persian_date = jdatetime.date.fromgregorian(date=greg_date)
        return f"{persian_date.year}/{persian_date.month:02d}/{persian_date.day:02d}"
    except Exception:
        return str(date_value)
    
def extract_flight_date(text):
    if pd.isna(text):
        return ""
    match = re.search(r"تاریخ حرکت:\s*([0-9۰-۹\/]+)", str(text))
    return match.group(1) if match else ""

def extract_hotel_dates(row):
    checkin = str(row.get("تاریخ ورود", "")) if pd.notna(row.get("تاریخ ورود")) else ""
    checkout = str(row.get("تاریخ خروج", "")) if pd.notna(row.get("تاریخ خروج")) else ""
    if checkin and checkout:
        return f"ورود: {checkin}  خروج: {checkout}"
    elif checkin:
        return f"ورود: {checkin}"
    elif checkout:
        return f"خروج: {checkout}"
    return ""

def detect_file_type(df):
    if "هتل" in df.columns:
        return "هتل"
    elif "مسیر" in df.columns:
        return "پرواز"
    else:
        return "خدمات"

def process_flight_data(df):
    df = df.copy()
    df["نام مسافر"] = df["نام مسافر"].apply(
        lambda x: str(x).rstrip(', ').strip() if pd.notna(x) else x
    )
    output_df = pd.DataFrame()
    output_df["ردیف"] = range(1, len(df) + 1)
    output_df["قرارداد"] = df["شماره قرارداد"].astype(str)
    output_df["نام مسافر"] = df["نام مسافر"]
    output_df["مسیر"] = df["مسیر"]
    output_df["شماره بلیط"] = df["شماره بلیط"].astype(str)
    output_df["تاریخ پرواز"] = df["تاریخ پرواز"].apply(extract_flight_date)
    output_df["بدهکار"] = df["بدهکار"].apply(persian_to_english_number)
    output_df["بستانکار"] = df["بستانکار"].apply(persian_to_english_number)
    output_df["مانده"] = output_df["بدهکار"] - output_df["بستانکار"]
    return output_df

def process_hotel_data(df):
    df = df.copy()
    df["نام مسافر"] = df["نام مسافر"].apply(
        lambda x: str(x).strip() if pd.notna(x) else x
    )
    
    passenger_count_col = None
    exact_matches = ["تعداد مسافر", "تعداد نفرات", "تعداد", "مسافرین", "تعداد افراد"]
    for col in df.columns:
        if col.strip() in exact_matches:
            passenger_count_col = col
            break
    if not passenger_count_col:
        for col in df.columns:
            col_lower = col.strip()
            if ("تعداد" in col_lower or "نفرات" in col_lower or "مسافر" in col_lower) and "نام" not in col_lower:
                passenger_count_col = col
                break
    if passenger_count_col:
        df["تعداد نفرات"] = df[passenger_count_col].apply(
            lambda x: max(1, persian_to_english_number(x)) if pd.notna(x) else 1
        )
    else:
        df["تعداد نفرات"] = 1

    output_df = pd.DataFrame()
    output_df["ردیف"] = range(1, len(df) + 1)
    output_df["قرارداد"] = df["شماره قرارداد"].astype(str)
    output_df["نام مسافر"] = df["نام مسافر"]
    output_df["هتل"] = df["هتل"]
    output_df["نوع اتاق"] = df["اتاق"]
    output_df["تعداد نفرات"] = df["تعداد نفرات"]
    output_df["تاریخ ورود و خروج"] = df.apply(extract_hotel_dates, axis=1)
    output_df["بدهکار"] = df["بدهکار"].apply(persian_to_english_number)
    output_df["بستانکار"] = df["بستانکار"].apply(persian_to_english_number)
    output_df["مانده"] = output_df["بدهکار"] - output_df["بستانکار"]
    return output_df

def process_service_data(df):
    df = df.copy()
    df["نام مسافر"] = df["نام مسافر"].apply(
        lambda x: str(x).rstrip(', ').strip() if pd.notna(x) else x
    )
    output_df = pd.DataFrame()
    output_df["ردیف"] = range(1, len(df) + 1)
    output_df["قرارداد"] = df["شماره قرارداد"].astype(str)
    output_df["نام مسافر"] = df["نام مسافر"]
    output_df["شرح خدمات"] = df.get("شرح خدمات", "")
    output_df["تاریخ"] = df.get("تاریخ", "")
    output_df["توضیحات"] = df.get("توضیحات", "")
    output_df["بدهکار"] = df["بدهکار"].apply(persian_to_english_number)
    output_df["بستانکار"] = df["بستانکار"].apply(persian_to_english_number)
    output_df["مانده"] = output_df["بدهکار"] - output_df["بستانکار"]
    return output_df

def generate_word(template_path, output_word_path, context):
    doc = DocxTemplate(template_path)
    doc.render(context)
    doc.save(output_word_path)

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

def get_base_path():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

def safe_filename(text):
    text = str(text).strip()
    text = re.sub(r'[\\/*?:"<>|]', '-', text)
    text = text.replace('/', '-').replace('\\', '-')
    text = text.rstrip('. ')
    if not text:
        text = "بدون_نام"
    return text

def ensure_simple_template_exists(template_path):
    if not os.path.exists(template_path):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'B Nazanin'
        style.font.size = Pt(16)
        p = doc.add_paragraph()
        p.add_run("تاریخ: {{ letter_date }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("شماره: {{ letter_number }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        run = p.add_run("شرکت محترم {{ payer }}")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("احتراما بدین وسیله صورت حساب مربوط به {{ period_range }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("بابت {{ service_type }} به مبلغ {{ total_debt }} ریال")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("جهت اطلاع و انجام اقدامات لازم به پیوست ایفاد می‌گردد.")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("خواهشمند است دستور فرمائید مبلغ مربوطه به حساب این شرکت واریز شود.")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("مشخصات حساب بانکی شرکت به شرح زیر است:")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("نام بانک: {{ bank_name1 }} - شماره حساب: {{ account_number1 }} - شماره شبا: {{ shaba_number1 }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("نام بانک: {{ bank_name2 }} - شماره حساب: {{ account_number2 }} - شماره شبا: {{ shaba_number2 }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("نام صاحب حساب: {{ company_name }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("{{ manager_name }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("{{ manager_position }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        run = p.add_run("برابر قانون و با توجه به ماهیت کار آژانس (ارائه خدمات واسطه‌ای) برای هریک از خدمات ارائه شده، صورتحساب به نام و کد ملی مسافر صادر و به سامانه مودیان ارسال می‌گردد.")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.save(template_path)

def ensure_creditor_template_exists(template_path):
    if not os.path.exists(template_path):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'B Nazanin'
        style.font.size = Pt(16)
        p = doc.add_paragraph()
        p.add_run("تاریخ: {{ letter_date }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("شماره: {{ letter_number }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        run = p.add_run("شرکت محترم {{ payer }}")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("احتراماً، بدین وسیله اعلام می‌دارد که بابت خدمات {{ service_type }} ارائه شده در بازه {{ period_range }}، مبلغ **{{ total_credit }}** ریال بستانکار این شرکت می‌باشد.")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("خواهشمند است نسبت به صدور فاکتور رسمی به نام این شرکت اقدام فرمایید تا مبلغ مربوطه در اسرع وقت پرداخت گردد.")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("اطلاعات حساب شرکت برای واریز بستانکاری:")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("بانک {{ bank_name1 }} - شماره حساب {{ account_number1 }} - شبا {{ shaba_number1 }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("بانک {{ bank_name2 }} - شماره حساب {{ account_number2 }} - شبا {{ shaba_number2 }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("با احترام،")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("{{ manager_name }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("{{ manager_position }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.save(template_path)

def ensure_subsidiary_template_exists(template_path):
    if not os.path.exists(template_path):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'B Nazanin'
        style.font.size = Pt(16)
        p = doc.add_paragraph()
        p.add_run("تاریخ: {{ letter_date }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("شماره: {{ letter_number }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        run = p.add_run("شرکت محترم {{ payer }}")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("احتراما بدین وسیله صورت حساب مربوط به {{ subsidiary }} در بازه {{ period_range }}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("بابت {{ service_type }} به مبلغ {{ total_debt }} ریال")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("جهت اطلاع و انجام اقدامات لازم به پیوست ایفاد می‌گردد.")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("خواهشمند است دستور فرمائید مبلغ مربوطه به حساب این شرکت واریز شود.")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("مشخصات حساب بانکی شرکت به شرح زیر است:")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("بانک ملت شماره حساب: 78466591\t شماره شبا: IR650120000000000078466591")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("بانک سامان شماره حساب: 1-4559598-810-820\t شماره شبا: IR015980455981000820056019")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("شرکت خدمات مسافرت هوائی گردشگری و زیارتی سفرهای کرمان خودرو")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("حمیدرضا فدائی")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        p.add_run("مدیر امور اداری و مالی")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph()
        run = p.add_run("** برابر قانون و با توجه به ماهیت کار آژانس (ارائه خدمات واسطه‌ای) برای هریک از خدمات ارائه شده ، صورتحساب به نام و کد ملی مسافر صادر و به سامانه مودیان ارسال می‌گردد. **")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.save(template_path)

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class FileItem(ctk.CTkFrame):
    def __init__(self, parent, file_path, file_type, on_remove, on_type_change):
        super().__init__(parent, fg_color="transparent")
        self.file_path = file_path
        self.on_remove = on_remove
        self.on_type_change = on_type_change
        self.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(self, text=os.path.basename(file_path), font=("B Nazanin", 16)).grid(row=0, column=0, padx=5, pady=2, sticky="w")
        self.type_combo = ctk.CTkComboBox(self, values=["پرواز", "هتل", "خدمات"], command=self.type_changed, font=("B Nazanin", 16), dropdown_font=("B Nazanin", 16), width=120)
        self.type_combo.set(file_type)
        self.type_combo.grid(row=0, column=1, padx=5, pady=2)
        ctk.CTkButton(self, text="❌", command=self.remove, width=40, fg_color=COLOR_ACCENT, hover_color="#d45a1c").grid(row=0, column=2, padx=5, pady=2)
    def type_changed(self, choice):
        self.on_type_change(self.file_path, choice)
    def remove(self):
        self.on_remove(self)

class DatabaseInterface(abc.ABC):
    @abc.abstractmethod
    def connect(self): pass
    @abc.abstractmethod
    def get_user(self, username, password_hash): pass
    @abc.abstractmethod
    def get_all_users(self): pass
    @abc.abstractmethod
    def add_user(self, username, password_hash, can_view, can_pay, can_issue, can_manage_users, can_manage_settings, can_delete_history, can_delete_payment, can_export, can_manage_companies, can_submit, can_approve): pass
    @abc.abstractmethod
    def update_user(self, user_id, username, password_hash, can_view, can_pay, can_issue, can_manage_users, can_manage_settings, can_delete_history, can_delete_payment, can_export, can_manage_companies, can_submit, can_approve): pass
    @abc.abstractmethod
    def delete_user(self, user_id): pass
    @abc.abstractmethod
    def get_all_payers(self): pass
    @abc.abstractmethod
    def add_payer(self, code, name, parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email): pass
    @abc.abstractmethod
    def update_payer(self, old_code, new_name, new_code, new_parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email): pass
    @abc.abstractmethod
    def delete_payer(self, code): pass
    @abc.abstractmethod
    def add_invoice(self, letter_number, payer, payer_code, invoice_type, amount, issue_date, file_path): pass
    @abc.abstractmethod
    def get_invoices(self, filters=None): pass
    @abc.abstractmethod
    def delete_invoice(self, invoice_id): pass
    @abc.abstractmethod
    def delete_all_invoices(self): pass
    @abc.abstractmethod
    def add_payment(self, payment_code, payer_code, payer_name, payment_date, amount, tracking_number, description): pass
    @abc.abstractmethod
    def get_payments(self, payer_code=None): pass
    @abc.abstractmethod
    def delete_payment(self, payment_id): pass
    @abc.abstractmethod
    def add_manual_debt(self, payer_code, description, amount, date, notes): pass
    @abc.abstractmethod
    def get_manual_debts(self, payer_code=None): pass
    @abc.abstractmethod
    def delete_manual_debt(self, debt_id): pass
    @abc.abstractmethod
    def add_credit(self, payer_code, amount, credit_date, description, invoice_number): pass
    @abc.abstractmethod
    def get_credits(self, payer_code=None): pass
    @abc.abstractmethod
    def delete_credit(self, credit_id): pass
    @abc.abstractmethod
    def get_next_letter_number(self, payer_code, year): pass
    @abc.abstractmethod
    def update_letter_sequence(self, payer_code, year, last_number): pass
    @abc.abstractmethod
    def get_company_settings(self): pass
    @abc.abstractmethod
    def save_company_settings(self, settings): pass
    @abc.abstractmethod
    def reserve_letter_numbers(self, payer_code, year, count): pass
    @abc.abstractmethod
    def release_letter_numbers(self, payer_code, year, numbers): pass
    @abc.abstractmethod
    def backup(self, backup_path): pass
    @abc.abstractmethod
    def restore(self, backup_path): pass
    @abc.abstractmethod
    def create_approval_request(self, requester_id, requester_name, payer_code, payer_name, letter_number, letter_date, start_day, end_day, month, year, period_range, service_data, flight_data, hotel_data): pass
    @abc.abstractmethod
    def get_pending_approval_requests(self): pass
    @abc.abstractmethod
    def get_approval_requests(self, user_id=None): pass
    @abc.abstractmethod
    def get_approval_request_by_id(self, req_id): pass
    @abc.abstractmethod
    def approve_request(self, req_id, approver_id): pass
    @abc.abstractmethod
    def reject_request(self, req_id, approver_id, reason): pass
    @abc.abstractmethod
    def delete_approval_request(self, req_id): pass
    def peek_next_letter_number(self, payer_code, year): pass

class SQLServerDatabase(DatabaseInterface):
    def __init__(self, config):
        self.config = config
        self.ensure_tables()

    def connect(self):
        conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={self.config['server']};DATABASE={self.config['database']};UID={self.config['username']};PWD={self.config['password']}"
        return pyodbc.connect(conn_str, autocommit=False)

    def ensure_tables(self):
        try:
            conn = self.connect()
            cursor = conn.cursor()
            # جدول users با ستون can_edit_template
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='users' AND xtype='U')
                CREATE TABLE users (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    username NVARCHAR(50) UNIQUE,
                    password_hash NVARCHAR(64),
                    can_view BIT DEFAULT 1,
                    can_pay BIT DEFAULT 0,
                    can_issue BIT DEFAULT 0,
                    can_manage_users BIT DEFAULT 0,
                    can_manage_settings BIT DEFAULT 0,
                    can_delete_history BIT DEFAULT 0,
                    can_delete_payment BIT DEFAULT 0,
                    can_export BIT DEFAULT 0,
                    can_manage_companies BIT DEFAULT 0,
                    can_submit BIT DEFAULT 0,
                    can_approve BIT DEFAULT 0
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'can_edit_template')
                ALTER TABLE users ADD can_edit_template BIT DEFAULT 0
            """)
            # بقیه جداول (مطابق کد اصلی شما)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='payers' AND xtype='U')
                CREATE TABLE payers (
                    code NVARCHAR(5) PRIMARY KEY,
                    name NVARCHAR(100),
                    parent_code NVARCHAR(5) NULL,
                    national_id NVARCHAR(20),
                    economic_code NVARCHAR(20),
                    registration_number NVARCHAR(20),
                    address NVARCHAR(MAX),
                    postal_code NVARCHAR(20),
                    phone NVARCHAR(20),
                    customer_type NVARCHAR(10) DEFAULT 'legal',
                    mobile NVARCHAR(20),
                    email NVARCHAR(100)
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='invoices' AND xtype='U')
                CREATE TABLE invoices (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    letter_number NVARCHAR(50),
                    payer NVARCHAR(100),
                    payer_code NVARCHAR(5),
                    invoice_type NVARCHAR(50),
                    amount BIGINT,
                    issue_date DATE,
                    file_path NVARCHAR(500)
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='payments' AND xtype='U')
                CREATE TABLE payments (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    payment_code NVARCHAR(20),
                    payer_code NVARCHAR(5),
                    payer_name NVARCHAR(100),
                    payment_date DATE,
                    amount BIGINT,
                    tracking_number NVARCHAR(50),
                    description NVARCHAR(500)
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='manual_debts' AND xtype='U')
                CREATE TABLE manual_debts (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    payer_code NVARCHAR(5),
                    description NVARCHAR(200),
                    amount BIGINT,
                    date DATE,
                    notes NVARCHAR(500)
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='letter_sequences' AND xtype='U')
                CREATE TABLE letter_sequences (
                    payer_code NVARCHAR(5),
                    year NVARCHAR(2),
                    last_number INT,
                    PRIMARY KEY (payer_code, year)
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='available_letter_numbers' AND xtype='U')
                CREATE TABLE available_letter_numbers (
                    payer_code NVARCHAR(5),
                    year NVARCHAR(2),
                    number INT,
                    PRIMARY KEY (payer_code, year, number)
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='company_settings' AND xtype='U')
                CREATE TABLE company_settings (
                    id INT PRIMARY KEY DEFAULT 1,
                    company_name NVARCHAR(200),
                    national_id NVARCHAR(20),
                    economic_code NVARCHAR(20),
                    registration_number NVARCHAR(20),
                    address NVARCHAR(MAX),
                    postal_code NVARCHAR(20),
                    phone NVARCHAR(20),
                    bank_name1 NVARCHAR(100),
                    account_number1 NVARCHAR(50),
                    shaba_number1 NVARCHAR(50),
                    bank_name2 NVARCHAR(100),
                    account_number2 NVARCHAR(50),
                    shaba_number2 NVARCHAR(50),
                    logo_path NVARCHAR(500),
                    signature_path NVARCHAR(500),
                    stamp_path NVARCHAR(500),
                    manager_name NVARCHAR(100),
                    manager_position NVARCHAR(100),
                    approved_output_path NVARCHAR(500) NULL
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='credits' AND xtype='U')
                CREATE TABLE credits (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    payer_code NVARCHAR(5),
                    amount BIGINT,
                    credit_date DATE,
                    description NVARCHAR(200),
                    invoice_number NVARCHAR(50),
                    created_at DATETIME DEFAULT GETDATE()
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='approval_requests' AND xtype='U')
                CREATE TABLE approval_requests (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    requester_id INT,
                    requester_name NVARCHAR(50),
                    payer_code NVARCHAR(5),
                    payer_name NVARCHAR(100),
                    letter_number NVARCHAR(50),
                    letter_date NVARCHAR(20),
                    start_day INT,
                    end_day INT,
                    month NVARCHAR(20),
                    year INT,
                    period_range NVARCHAR(100),
                    service_data_json NVARCHAR(MAX),
                    flight_data_json NVARCHAR(MAX),
                    hotel_data_json NVARCHAR(MAX),
                    status NVARCHAR(20) DEFAULT 'pending',
                    created_at DATETIME DEFAULT GETDATE(),
                    approved_by INT NULL,
                    approved_at DATETIME NULL,
                    rejection_reason NVARCHAR(500) NULL,
                    reserved_numbers NVARCHAR(MAX) NULL,
                    final_letter_number NVARCHAR(50) NULL
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='final_files' AND xtype='U')
                CREATE TABLE final_files (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    request_id INT NOT NULL,
                    letter_number NVARCHAR(50),
                    service_type NVARCHAR(50),
                    file_name NVARCHAR(200),
                    file_path NVARCHAR(500) NULL,
                    file_data VARBINARY(MAX) NULL,
                    created_at DATETIME DEFAULT GETDATE(),
                    last_accessed DATETIME NULL,
                    is_deleted_from_disk BIT DEFAULT 0,
                    FOREIGN KEY (request_id) REFERENCES approval_requests(id) ON DELETE CASCADE
                )
            """)
            cursor.execute("""
                IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='service_categories' AND xtype='U')
                CREATE TABLE service_categories (
                    id INT IDENTITY(1,1) PRIMARY KEY,
                    parent_id INT NULL,
                    name NVARCHAR(100) NOT NULL,
                    code NVARCHAR(20) NOT NULL UNIQUE,
                    is_active BIT DEFAULT 1,
                    display_order INT DEFAULT 0,
                    created_at DATETIME DEFAULT GETDATE(),
                    FOREIGN KEY (parent_id) REFERENCES service_categories(id) ON DELETE CASCADE
                )
            """)
            cursor.execute("IF NOT EXISTS (SELECT 1 FROM service_categories) INSERT INTO service_categories (name, code, is_active, display_order) VALUES ('ویزا','VISA',1,1), ('اتوبوس','BUS',1,2), ('قطار','TRAIN',1,3), ('گشت','TOUR',1,4), ('تور','PACKAGE',1,5), ('CIP','CIP',1,6)")
            conn.commit()
            conn.close()
        except Exception as e:
            logging.error(f"خطا در ایجاد یا به‌روزرسانی جداول: {e}")

    # ==================== متدهای مربوط به users ====================
    def get_user(self, username, password_hash):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, username, can_view, can_pay, can_issue, can_manage_users,
                       can_manage_settings, can_delete_history, can_delete_payment,
                       can_export, can_manage_companies, can_submit, can_approve, can_edit_template
                FROM users WHERE username=? AND password_hash=?
            """, (username, password_hash))
            row = cursor.fetchone()
            if row:
                return (row[0], row[1], bool(row[2]), bool(row[3]), bool(row[4]), bool(row[5]),
                        bool(row[6]), bool(row[7]), bool(row[8]), bool(row[9]), bool(row[10]),
                        bool(row[11]), bool(row[12]), bool(row[13]))
            return None

    def get_all_users(self):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, username, can_view, can_pay, can_issue, can_manage_users,
                       can_manage_settings, can_delete_history, can_delete_payment,
                       can_export, can_manage_companies, can_submit, can_approve, can_edit_template
                FROM users ORDER BY id
            """)
            rows = cursor.fetchall()
            return [(row[0], row[1], bool(row[2]), bool(row[3]), bool(row[4]), bool(row[5]),
                     bool(row[6]), bool(row[7]), bool(row[8]), bool(row[9]), bool(row[10]),
                     bool(row[11]), bool(row[12]), bool(row[13])) for row in rows]

    def add_user(self, username, password_hash, can_view, can_pay, can_issue, can_manage_users,
                 can_manage_settings, can_delete_history, can_delete_payment, can_export,
                 can_manage_companies, can_submit, can_approve, can_edit_template):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO users (username, password_hash, can_view, can_pay, can_issue,
                can_manage_users, can_manage_settings, can_delete_history, can_delete_payment,
                can_export, can_manage_companies, can_submit, can_approve, can_edit_template)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """, (username, password_hash, can_view, can_pay, can_issue, can_manage_users,
                  can_manage_settings, can_delete_history, can_delete_payment, can_export,
                  can_manage_companies, can_submit, can_approve, can_edit_template))
            conn.commit()

    def update_user(self, user_id, username, password_hash, can_view, can_pay, can_issue,
                    can_manage_users, can_manage_settings, can_delete_history, can_delete_payment,
                    can_export, can_manage_companies, can_submit, can_approve, can_edit_template):
        with self.connect() as conn:
            cursor = conn.cursor()
            if password_hash:
                cursor.execute("""
                    UPDATE users SET username=?, password_hash=?, can_view=?, can_pay=?,
                    can_issue=?, can_manage_users=?, can_manage_settings=?, can_delete_history=?,
                    can_delete_payment=?, can_export=?, can_manage_companies=?, can_submit=?,
                    can_approve=?, can_edit_template=? WHERE id=?
                """, (username, password_hash, can_view, can_pay, can_issue, can_manage_users,
                      can_manage_settings, can_delete_history, can_delete_payment, can_export,
                      can_manage_companies, can_submit, can_approve, can_edit_template, user_id))
            else:
                cursor.execute("""
                    UPDATE users SET username=?, can_view=?, can_pay=?, can_issue=?,
                    can_manage_users=?, can_manage_settings=?, can_delete_history=?,
                    can_delete_payment=?, can_export=?, can_manage_companies=?, can_submit=?,
                    can_approve=?, can_edit_template=? WHERE id=?
                """, (username, can_view, can_pay, can_issue, can_manage_users,
                      can_manage_settings, can_delete_history, can_delete_payment, can_export,
                      can_manage_companies, can_submit, can_approve, can_edit_template, user_id))
            conn.commit()

    def delete_user(self, user_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM users WHERE id=?", (user_id,))
            conn.commit()

    # ==================== متدهای مربوط به payers ====================
    def get_all_payers(self):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT code, name, parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email FROM payers ORDER BY name")
            rows = cursor.fetchall()
            payers = []
            for row in rows:
                payers.append({"code": row[0], "name": row[1], "parent_code": row[2], "national_id": row[3], "economic_code": row[4], "registration_number": row[5], "address": row[6], "postal_code": row[7], "phone": row[8], "customer_type": row[9], "mobile": row[10] if len(row)>10 else "", "email": row[11] if len(row)>11 else ""})
            return payers

    def add_payer(self, code, name, parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email):
        parent_code_val = parent_code if parent_code else None
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO payers (code, name, parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)", (code, name, parent_code_val, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email))
            conn.commit()

    def update_payer(self, old_code, new_name, new_code, new_parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email):
        with self.connect() as conn:
            cursor = conn.cursor()
            if new_parent_code == "" or new_parent_code is None:
                new_parent_code = None
            if old_code == new_code:
                cursor.execute("UPDATE payers SET name=?, parent_code=?, national_id=?, economic_code=?, registration_number=?, address=?, postal_code=?, phone=?, customer_type=?, mobile=?, email=? WHERE code=?", (new_name, new_parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email, old_code))
            else:
                cursor.execute("SELECT COUNT(*) FROM payers WHERE parent_code=?", (old_code,))
                child_count = cursor.fetchone()[0]
                if child_count > 0:
                    raise Exception("این شرکت دارای زیرمجموعه است. برای تغییر کد ابتدا زیرمجموعه‌ها را حذف یا جابه‌جا کنید.")
                cursor.execute("UPDATE payers SET code=?, name=?, parent_code=?, national_id=?, economic_code=?, registration_number=?, address=?, postal_code=?, phone=?, customer_type=?, mobile=?, email=? WHERE code=?", (new_code, new_name, new_parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email, old_code))
            if new_code == new_parent_code:
                cursor.execute("UPDATE payers SET parent_code=NULL WHERE code=?", (new_code,))
            conn.commit()

    def delete_payer(self, code):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE payers SET parent_code=NULL WHERE parent_code=?", (code,))
            cursor.execute("DELETE FROM payers WHERE code=?", (code,))
            conn.commit()

    # ==================== متدهای مربوط به invoices ====================
    def add_invoice(self, letter_number, payer, payer_code, invoice_type, amount, issue_date, file_path):
        gregorian_date = persian_date_to_gregorian(issue_date)
        if gregorian_date is None:
            gregorian_date = issue_date
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO invoices (letter_number, payer, payer_code, invoice_type, amount, issue_date, file_path) VALUES (?,?,?,?,?,?,?)", (letter_number, payer, payer_code, invoice_type, amount, gregorian_date, file_path))
            conn.commit()

    def get_invoices(self, filters=None):
        with self.connect() as conn:
            cursor = conn.cursor()
            query = "SELECT id, letter_number, payer, payer_code, invoice_type, amount, issue_date, file_path FROM invoices WHERE 1=1"
            params = []
            if filters:
                if filters.get('payer'):
                    query += " AND payer LIKE ?"
                    params.append(f"%{filters['payer']}%")
                if filters.get('payer_code'):
                    query += " AND payer_code = ?"
                    params.append(filters['payer_code'])
                if filters.get('start_date'):
                    query += " AND issue_date >= ?"
                    params.append(filters['start_date'])
                if filters.get('end_date'):
                    query += " AND issue_date <= ?"
                    params.append(filters['end_date'])
            query += " ORDER BY id DESC"
            cursor.execute(query, params)
            rows = cursor.fetchall()
            return rows

    def delete_invoice(self, invoice_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT payer_code, letter_number, issue_date FROM invoices WHERE id=?", (invoice_id,))
            row = cursor.fetchone()
            if row:
                payer_code = row[0]
                letter_number = row[1]
                issue_date = row[2]
                if issue_date:
                    try:
                        jalali_date = jdatetime.date.fromgregorian(date=issue_date)
                        year = str(jalali_date.year % 100).zfill(2)
                    except:
                        match = re.search(r'\d{4}', letter_number)
                        if match:
                            year = match.group()[-2:]
                        else:
                            today = jdatetime.date.today()
                            year = str(today.year % 100).zfill(2)
                else:
                    match = re.search(r'\d{4}', letter_number)
                    if match:
                        year = match.group()[-2:]
                    else:
                        today = jdatetime.date.today()
                        year = str(today.year % 100).zfill(2)
                num_match = re.search(r'(\d+)(?!.*\d)', letter_number)
                if num_match:
                    num = int(num_match.group(1))
                    cursor.execute("INSERT INTO available_letter_numbers (payer_code, year, number) VALUES (?, ?, ?)", (payer_code, year, num))
            cursor.execute("DELETE FROM invoices WHERE id=?", (invoice_id,))
            conn.commit()

    def delete_all_invoices(self):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT payer_code, letter_number, issue_date FROM invoices")
            rows = cursor.fetchall()
            for row in rows:
                payer_code = row[0]
                letter_number = row[1]
                issue_date = row[2]
                if issue_date:
                    try:
                        jalali_date = jdatetime.date.fromgregorian(date=issue_date)
                        year = str(jalali_date.year % 100).zfill(2)
                    except:
                        match = re.search(r'\d{4}', letter_number)
                        if match:
                            year = match.group()[-2:]
                        else:
                            today = jdatetime.date.today()
                            year = str(today.year % 100).zfill(2)
                else:
                    match = re.search(r'\d{4}', letter_number)
                    if match:
                        year = match.group()[-2:]
                    else:
                        today = jdatetime.date.today()
                        year = str(today.year % 100).zfill(2)
                num_match = re.search(r'(\d+)(?!.*\d)', letter_number)
                if num_match:
                    num = int(num_match.group(1))
                    cursor.execute("INSERT INTO available_letter_numbers (payer_code, year, number) VALUES (?, ?, ?)", (payer_code, year, num))
            cursor.execute("DELETE FROM invoices")
            conn.commit()

    # ==================== متدهای مربوط به payments ====================
    def add_payment(self, payment_code, payer_code, payer_name, payment_date, amount, tracking_number, description):
        gregorian_date = persian_date_to_gregorian(payment_date)
        if gregorian_date is None:
            raise ValueError(f"تاریخ پرداخت نامعتبر: {payment_date}")
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO payments (payment_code, payer_code, payer_name, payment_date, amount, tracking_number, description) VALUES (?,?,?,?,?,?,?)", (payment_code, payer_code, payer_name, gregorian_date, amount, tracking_number, description))
            conn.commit()

    def get_payments(self, payer_code=None):
        with self.connect() as conn:
            cursor = conn.cursor()
            if payer_code:
                cursor.execute("SELECT id, payment_date, amount, tracking_number, description, payer_code FROM payments WHERE payer_code=? ORDER BY payment_date DESC", (payer_code,))
            else:
                cursor.execute("SELECT id, payment_date, amount, tracking_number, description, payer_code FROM payments ORDER BY payment_date DESC")
            rows = cursor.fetchall()
            return rows

    def delete_payment(self, payment_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM payments WHERE id=?", (payment_id,))
            conn.commit()

    # ==================== متدهای مربوط به manual_debts ====================
    def add_manual_debt(self, payer_code, description, amount, date, notes):
        if isinstance(date, str) and '/' in date:
            try:
                parts = date.split('/')
                if len(parts) == 3:
                    year = int(english_number(parts[0]))
                    month = int(english_number(parts[1]))
                    day = int(english_number(parts[2]))
                    if 1300 <= year <= 1500:
                        greg_date = jdatetime.date(year, month, day).togregorian()
                        date = greg_date.strftime("%Y-%m-%d")
            except:
                pass
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO manual_debts (payer_code, description, amount, date, notes) VALUES (?,?,?,?,?)", (payer_code, description, amount, date, notes))
            conn.commit()

    def get_manual_debts(self, payer_code=None):
        with self.connect() as conn:
            cursor = conn.cursor()
            if payer_code:
                cursor.execute("SELECT id, description, amount, date, notes FROM manual_debts WHERE payer_code=? ORDER BY date DESC", (payer_code,))
            else:
                cursor.execute("SELECT id, description, amount, date, notes, payer_code FROM manual_debts ORDER BY date DESC")
            rows = cursor.fetchall()
            return rows

    def delete_manual_debt(self, debt_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM manual_debts WHERE id=?", (debt_id,))
            conn.commit()

    # ==================== متدهای مربوط به credits ====================
    def add_credit(self, payer_code, amount, credit_date, description, invoice_number=""):
        gregorian_date = persian_date_to_gregorian(credit_date)
        if gregorian_date is None:
            raise ValueError(f"تاریخ بستانکاری نامعتبر: {credit_date}")
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO credits (payer_code, amount, credit_date, description, invoice_number) VALUES (?,?,?,?,?)", (payer_code, amount, gregorian_date, description, invoice_number))
            conn.commit()

    def get_credits(self, payer_code=None):
        with self.connect() as conn:
            cursor = conn.cursor()
            if payer_code:
                cursor.execute("SELECT id, amount, credit_date, description, invoice_number FROM credits WHERE payer_code=? ORDER BY credit_date DESC", (payer_code,))
            else:
                cursor.execute("SELECT id, amount, credit_date, description, invoice_number, payer_code FROM credits ORDER BY credit_date DESC")
            return cursor.fetchall()

    def delete_credit(self, credit_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM credits WHERE id=?", (credit_id,))
            conn.commit()

    # ==================== متدهای مربوط به letter sequences ====================
    def get_next_letter_number(self, payer_code, year):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT MIN(number) FROM available_letter_numbers WHERE payer_code = ? AND year = ?", (payer_code, year))
            row = cursor.fetchone()
            if row and row[0] is not None:
                return row[0]
            cursor.execute("SELECT last_number FROM letter_sequences WHERE payer_code = ? AND year = ?", (payer_code, year))
            row = cursor.fetchone()
            return row[0] + 1 if row else 1

    def update_letter_sequence(self, payer_code, year, last_number):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT 1 FROM available_letter_numbers WHERE payer_code = ? AND year = ? AND number = ?", (payer_code, year, last_number))
            if cursor.fetchone():
                cursor.execute("DELETE FROM available_letter_numbers WHERE payer_code = ? AND year = ? AND number = ?", (payer_code, year, last_number))
            else:
                cursor.execute("MERGE INTO letter_sequences AS target USING (VALUES (?, ?, ?)) AS source (payer_code, year, last_number) ON target.payer_code = source.payer_code AND target.year = source.year WHEN MATCHED THEN UPDATE SET last_number = source.last_number WHEN NOT MATCHED THEN INSERT (payer_code, year, last_number) VALUES (source.payer_code, source.year, source.last_number);", (payer_code, year, last_number))
            conn.commit()

    def reserve_letter_numbers(self, payer_code, year, count):
        numbers = []
        for _ in range(count):
            next_num = self.get_next_letter_number(payer_code, year)
            numbers.append(next_num)
            self.update_letter_sequence(payer_code, year, next_num)
        return numbers

    def release_letter_numbers(self, payer_code, year, numbers):
        with self.connect() as conn:
            cursor = conn.cursor()
            for num in numbers:
                cursor.execute("IF NOT EXISTS (SELECT 1 FROM available_letter_numbers WHERE payer_code=? AND year=? AND number=?) INSERT INTO available_letter_numbers (payer_code, year, number) VALUES (?, ?, ?)", (payer_code, year, num, payer_code, year, num))
            conn.commit()

    def peek_next_letter_number(self, payer_code, year):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT MIN(number) FROM available_letter_numbers WHERE payer_code = ? AND year = ?", (payer_code, year))
            row = cursor.fetchone()
            if row and row[0] is not None:
                return row[0]
            cursor.execute("SELECT last_number FROM letter_sequences WHERE payer_code = ? AND year = ?", (payer_code, year))
            row = cursor.fetchone()
            return row[0] + 1 if row else 1

    # ==================== متدهای مربوط به company_settings ====================
    def get_company_settings(self):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM company_settings WHERE id=1")
            row = cursor.fetchone()
            if row:
                cols = [desc[0] for desc in cursor.description]
                return dict(zip(cols, row))
            return None

    def save_company_settings(self, settings):
        with self.connect() as conn:
            cursor = conn.cursor()
            try:
                cursor.execute("SELECT id FROM company_settings WHERE id=1")
                if cursor.fetchone():
                    set_clause = ", ".join([f"{k}=?" for k in settings.keys()])
                    cursor.execute(f"UPDATE company_settings SET {set_clause} WHERE id=1", list(settings.values()))
                else:
                    placeholders = ", ".join(["?"] * len(settings))
                    cursor.execute(f"INSERT INTO company_settings ({', '.join(settings.keys())}) VALUES ({placeholders})", list(settings.values()))
                conn.commit()
            except Exception as e:
                conn.rollback()
                raise e

    # ==================== متدهای مربوط به approval_requests ====================
    def create_approval_request(self, requester_id, requester_name, payer_code, payer_name, letter_number, letter_date, start_day, end_day, month, year, period_range, service_data, flight_data, hotel_data):
        import json
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO approval_requests
                (requester_id, requester_name, payer_code, payer_name, letter_number, letter_date,
                start_day, end_day, month, year, period_range, service_data_json, flight_data_json,
                hotel_data_json, status)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending')
            """, (requester_id, requester_name, payer_code, payer_name, letter_number, letter_date,
                start_day, end_day, month, year, period_range,
                json.dumps(service_data, ensure_ascii=False),
                json.dumps(flight_data, ensure_ascii=False),
                json.dumps(hotel_data, ensure_ascii=False)))
            cursor.execute("SELECT SCOPE_IDENTITY()")
            new_id = cursor.fetchone()[0]
            conn.commit()
            return new_id

    def get_pending_approval_requests(self):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, requester_name, payer_name, letter_number, letter_date, status, created_at FROM approval_requests WHERE status = 'pending' ORDER BY created_at DESC")
            return cursor.fetchall()

    def get_approval_requests(self, user_id=None):
        with self.connect() as conn:
            cursor = conn.cursor()
            if user_id is not None:
                cursor.execute("SELECT id, requester_id, requester_name, payer_name, letter_number, letter_date, status, created_at, rejection_reason FROM approval_requests WHERE requester_id = ? ORDER BY created_at DESC", (user_id,))
            else:
                cursor.execute("SELECT id, requester_id, requester_name, payer_name, letter_number, letter_date, status, created_at, rejection_reason FROM approval_requests ORDER BY created_at DESC")
            return cursor.fetchall()

    def get_approval_request_by_id(self, req_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM approval_requests WHERE id = ?", (req_id,))
            row = cursor.fetchone()
            if row:
                cols = [desc[0] for desc in cursor.description]
                return dict(zip(cols, row))
            return None

    def approve_request(self, req_id, approver_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE approval_requests SET status = 'approved', approved_by = ?, approved_at = GETDATE() WHERE id = ?", (approver_id, req_id))
            conn.commit()

    def reject_request(self, req_id, approver_id, reason):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE approval_requests SET status = 'rejected', approved_by = ?, approved_at = GETDATE(), rejection_reason = ? WHERE id = ?", (approver_id, reason, req_id))
            conn.commit()

    def delete_approval_request(self, req_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM approval_requests WHERE id = ?", (req_id,))
            conn.commit()

    # ==================== متدهای مربوط به final_files ====================
    def save_final_file(self, request_id, letter_number, service_type, file_name, file_path, file_data):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO final_files (request_id, letter_number, service_type, file_name, file_path, file_data)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (request_id, letter_number, service_type, file_name, file_path, pyodbc.Binary(file_data)))
            conn.commit()
            return cursor.execute("SELECT SCOPE_IDENTITY()").fetchone()[0]

    def get_final_file(self, file_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT file_path, file_data, file_name, request_id, letter_number, service_type FROM final_files WHERE id = ?", (file_id,))
            row = cursor.fetchone()
            if row:
                return {'file_path': row[0], 'file_data': row[1], 'file_name': row[2], 'request_id': row[3], 'letter_number': row[4], 'service_type': row[5]}
            return None

    def get_final_files_by_request(self, request_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, file_name, service_type, letter_number FROM final_files WHERE request_id = ?", (request_id,))
            return cursor.fetchall()

    def get_final_file_by_letter_number(self, letter_number):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT TOP 1 id, file_path, file_data, file_name FROM final_files WHERE letter_number = ?", (letter_number,))
            row = cursor.fetchone()
            if row:
                return {'id': row[0], 'file_path': row[1], 'file_data': row[2], 'file_name': row[3]}
            return None

    def update_final_file_disk_path(self, file_id, disk_path):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE final_files SET file_path = ?, is_deleted_from_disk = 0 WHERE id = ?", (disk_path, file_id))
            conn.commit()

    def mark_disk_deleted(self, file_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE final_files SET is_deleted_from_disk = 1 WHERE id = ?", (file_id,))
            conn.commit()

    # ==================== متدهای مربوط به service_categories ====================
    def get_all_service_categories(self, include_inactive=False):
        with self.connect() as conn:
            cursor = conn.cursor()
            if include_inactive:
                cursor.execute("SELECT id, parent_id, name, code, is_active FROM service_categories ORDER BY display_order, id")
            else:
                cursor.execute("SELECT id, parent_id, name, code, is_active FROM service_categories WHERE is_active=1 ORDER BY display_order, id")
            rows = cursor.fetchall()
            return [{'id': r[0], 'parent_id': r[1], 'name': r[2], 'code': r[3], 'is_active': r[4]} for r in rows]

    def get_service_category_tree(self):
        categories = self.get_all_service_categories()
        tree = []
        mapping = {}
        for cat in categories:
            cat['children'] = []
            mapping[cat['id']] = cat
        for cat in categories:
            if cat['parent_id'] is None:
                tree.append(cat)
            else:
                parent = mapping.get(cat['parent_id'])
                if parent:
                    parent['children'].append(cat)
        return tree

    def add_service_category(self, name, code, parent_id=None, is_active=True, display_order=0):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO service_categories (parent_id, name, code, is_active, display_order) VALUES (?,?,?,?,?)", (parent_id, name, code, is_active, display_order))
            conn.commit()
            cursor.execute("SELECT SCOPE_IDENTITY()")
            return cursor.fetchone()[0]

    def update_service_category(self, cat_id, name=None, code=None, parent_id=None, is_active=None, display_order=None):
        with self.connect() as conn:
            cursor = conn.cursor()
            updates = []
            params = []
            if name is not None:
                updates.append("name=?")
                params.append(name)
            if code is not None:
                updates.append("code=?")
                params.append(code)
            if parent_id is not None:
                updates.append("parent_id=?")
                params.append(parent_id)
            if is_active is not None:
                updates.append("is_active=?")
                params.append(is_active)
            if display_order is not None:
                updates.append("display_order=?")
                params.append(display_order)
            if not updates:
                return
            params.append(cat_id)
            cursor.execute(f"UPDATE service_categories SET {', '.join(updates)} WHERE id=?", params)
            conn.commit()

    def delete_service_category(self, cat_id):
        with self.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE service_categories SET parent_id = NULL WHERE parent_id = ?", (cat_id,))
            cursor.execute("DELETE FROM service_categories WHERE id = ?", (cat_id,))
            conn.commit()

    # ==================== پشتیبان‌گیری و بازیابی ====================
    def backup(self, backup_path):
        final_dir = os.path.dirname(backup_path)
        os.makedirs(final_dir, exist_ok=True)
        temp_dir = os.path.join(os.environ.get('TEMP', 'C:\\Temp'), 'KKTCOIG_Backup_Temp')
        os.makedirs(temp_dir, exist_ok=True)
        temp_backup_path = os.path.join(temp_dir, os.path.basename(backup_path))
        ps_script_template = '''
$connString = "DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password}"
$connection = New-Object System.Data.Odbc.OdbcConnection($connString)
$connection.Open()
$command = $connection.CreateCommand()
$command.CommandText = "BACKUP DATABASE [{database}] TO DISK = '{temp_path}'"
$command.ExecuteNonQuery()
$connection.Close()
'''
        ps_script = ps_script_template.format(
            server=self.config['server'],
            database=self.config['database'],
            username=self.config['username'],
            password=self.config['password'],
            temp_path=temp_backup_path
        )
        ps_script_path = os.path.join(temp_dir, 'backup_script.ps1')
        with open(ps_script_path, 'w', encoding='utf-8') as f:
            f.write(ps_script)
        try:
            subprocess.run(['powershell.exe', '-ExecutionPolicy', 'Bypass', '-File', ps_script_path], check=True, capture_output=True, text=True, timeout=60)
        except subprocess.CalledProcessError as e:
            error_msg = e.stderr if e.stderr else str(e)
            raise Exception(f"خطا در اجرای اسکریپت پشتیبان‌گیری: {error_msg}")
        except subprocess.TimeoutExpired:
            raise Exception("زمان اجرای اسکریپت پشتیبان‌گیری بیش از حد طول کشید.")
        max_wait = 10
        waited = 0
        while not os.path.exists(temp_backup_path) and waited < max_wait:
            time.sleep(1)
            waited += 1
        if not os.path.exists(temp_backup_path):
            raise Exception("فایل پشتیبان در مسیر موقت ایجاد نشد.")
        try:
            shutil.copy2(temp_backup_path, backup_path)
        except Exception as e:
            raise Exception(f"خطا در کپی فایل پشتیبان به مقصد نهایی: {str(e)}")
        finally:
            try:
                os.remove(temp_backup_path)
                os.remove(ps_script_path)
            except:
                pass
        if not os.path.exists(backup_path) or os.path.getsize(backup_path) == 0:
            raise Exception("فایل پشتیبان نهایی ایجاد نشد یا خالی است.")

    def restore(self, backup_path):
        conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={self.config['server']};DATABASE={self.config['database']};UID={self.config['username']};PWD={self.config['password']}"
        conn = pyodbc.connect(conn_str, autocommit=True)
        cursor = conn.cursor()
        try:
            cursor.execute(f"ALTER DATABASE [{self.config['database']}] SET SINGLE_USER WITH ROLLBACK IMMEDIATE")
            cursor.execute(f"RESTORE DATABASE [{self.config['database']}] FROM DISK = ? WITH REPLACE", (backup_path,))
            cursor.execute(f"ALTER DATABASE [{self.config['database']}] SET MULTI_USER")
        except Exception as e:
            try:
                cursor.execute(f"ALTER DATABASE [{self.config['database']}] SET MULTI_USER")
            except:
                pass
            raise Exception(f"خطا در بازیابی SQL Server: {str(e)}")
        finally:
            conn.close()

class UserManager:
    def __init__(self, db: DatabaseInterface):
        self.db = db
        self._init_admin()

    def _init_admin(self):
        if not self.db.get_all_users():
            self.db.add_user("admin", hashlib.sha256("admin".encode()).hexdigest(),
                             1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1, 1)  # can_edit_template=1

    def hash_password(self, pwd):
        return hashlib.sha256(pwd.encode()).hexdigest()

    def authenticate(self, username, password):
        return self.db.get_user(username, self.hash_password(password))

    def add_user(self, username, password, can_view, can_pay, can_issue, can_manage_users,
                 can_manage_settings, can_delete_history, can_delete_payment, can_export,
                 can_manage_companies, can_submit, can_approve, can_edit_template):
        password_hash = self.hash_password(password)
        try:
            self.db.add_user(username, password_hash, can_view, can_pay, can_issue,
                             can_manage_users, can_manage_settings, can_delete_history,
                             can_delete_payment, can_export, can_manage_companies,
                             can_submit, can_approve, can_edit_template)
            return True
        except Exception:
            return False

    def update_user(self, user_id, username, password, can_view, can_pay, can_issue,
                    can_manage_users, can_manage_settings, can_delete_history, can_delete_payment,
                    can_export, can_manage_companies, can_submit, can_approve, can_edit_template):
        password_hash = self.hash_password(password) if password else None
        try:
            self.db.update_user(user_id, username, password_hash, can_view, can_pay, can_issue,
                                can_manage_users, can_manage_settings, can_delete_history,
                                can_delete_payment, can_export, can_manage_companies,
                                can_submit, can_approve, can_edit_template)
            return True
        except Exception:
            return False

    def delete_user(self, user_id):
        self.db.delete_user(user_id)

    def get_all_users(self):
        return self.db.get_all_users()

class DatabaseManager:
    def __init__(self, config):
        self.config = config
    def backup_database(self, backup_path):
        final_dir = os.path.dirname(backup_path)
        os.makedirs(final_dir, exist_ok=True)
        temp_dir = os.path.join(os.environ.get('TEMP', 'C:\\Temp'), 'KKTCOIG_Backup_Temp')
        os.makedirs(temp_dir, exist_ok=True)
        temp_backup_path = os.path.join(temp_dir, os.path.basename(backup_path))
        ps_script_template = '''
$connString = "DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password}"
$connection = New-Object System.Data.Odbc.OdbcConnection($connString)
$connection.Open()
$command = $connection.CreateCommand()
$command.CommandText = "BACKUP DATABASE [{database}] TO DISK = '{temp_path}'"
$command.ExecuteNonQuery()
$connection.Close()
'''
        ps_script = ps_script_template.format(
            server=self.config['server'],
            database=self.config['database'],
            username=self.config['username'],
            password=self.config['password'],
            temp_path=temp_backup_path
        )
        ps_script_path = os.path.join(temp_dir, 'backup_script.ps1')
        with open(ps_script_path, 'w', encoding='utf-8') as f:
            f.write(ps_script)
        try:
            subprocess.run(['powershell.exe', '-ExecutionPolicy', 'Bypass', '-File', ps_script_path], check=True, capture_output=True, text=True, timeout=60)
        except subprocess.CalledProcessError as e:
            error_msg = e.stderr if e.stderr else str(e)
            raise Exception(f"خطا در اجرای اسکریپت پشتیبان‌گیری: {error_msg}")
        except subprocess.TimeoutExpired:
            raise Exception("زمان اجرای اسکریپت پشتیبان‌گیری بیش از حد طول کشید.")
        max_wait = 10
        waited = 0
        while not os.path.exists(temp_backup_path) and waited < max_wait:
            time.sleep(1)
            waited += 1
        if not os.path.exists(temp_backup_path):
            raise Exception("فایل پشتیبان در مسیر موقت ایجاد نشد.")
        try:
            shutil.copy2(temp_backup_path, backup_path)
        except Exception as e:
            raise Exception(f"خطا در کپی فایل پشتیبان به مقصد نهایی: {str(e)}")
        finally:
            try:
                os.remove(temp_backup_path)
                os.remove(ps_script_path)
            except:
                pass
        if not os.path.exists(backup_path) or os.path.getsize(backup_path) == 0:
            raise Exception("فایل پشتیبان نهایی ایجاد نشد یا خالی است.")
    def restore_database(self, backup_path):
        conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={self.config['server']};DATABASE={self.config['database']};UID={self.config['username']};PWD={self.config['password']}"
        conn = pyodbc.connect(conn_str, autocommit=True)
        cursor = conn.cursor()
        try:
            cursor.execute(f"ALTER DATABASE [{self.config['database']}] SET SINGLE_USER WITH ROLLBACK IMMEDIATE")
            cursor.execute(f"RESTORE DATABASE [{self.config['database']}] FROM DISK = ? WITH REPLACE", (backup_path,))
            cursor.execute(f"ALTER DATABASE [{self.config['database']}] SET MULTI_USER")
        except Exception as e:
            try:
                cursor.execute(f"ALTER DATABASE [{self.config['database']}] SET MULTI_USER")
            except:
                pass
            raise Exception(f"خطا در بازیابی SQL Server: {str(e)}")
        finally:
            conn.close()

def generate_backup_filename():
    return f"KKTCOIG-Backup-{jdatetime.datetime.now().strftime('%Y%m%d-%H%M%S')}"

class ManagePayersDialog(ctk.CTkToplevel):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.title("مدیریت طرف‌های حساب")
        self.geometry("900x800")
        self.transient(parent)
        self.grab_set()
        self.font_farsi = ("B Nazanin", 18)
        self.font_farsi_bold = ("B Nazanin", 18, "bold")
        ctk.CTkLabel(self, text="مدیریت طرف‌های حساب", font=("B Nazanin", 22, "bold")).pack(pady=15)
        list_frame = ctk.CTkFrame(self)
        list_frame.pack(pady=15, padx=20, fill="both", expand=True)
        scrollbar = ctk.CTkScrollbar(list_frame)
        scrollbar.pack(side="right", fill="y")
        import tkinter as tk
        self.listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, font=self.font_farsi, bg="#2b2b2b", fg="white", selectbackground=COLOR_PRIMARY, selectforeground="white", borderwidth=0, highlightthickness=0)
        self.listbox.pack(side="left", fill="both", expand=True)
        scrollbar.configure(command=self.listbox.yview)
        self.refresh_listbox()
        input_frame = ctk.CTkFrame(self)
        input_frame.pack(pady=15, padx=20, fill="x")
        col1 = ctk.CTkFrame(input_frame, fg_color="transparent")
        col1.pack(side="left", fill="both", expand=True, padx=5)
        col2 = ctk.CTkFrame(input_frame, fg_color="transparent")
        col2.pack(side="left", fill="both", expand=True, padx=5)
        ctk.CTkLabel(col1, text="کد شرکت", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.entry_code = ctk.CTkEntry(col1, width=150, font=self.font_farsi)
        self.entry_code.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col1, text="نام شرکت", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.entry_name = ctk.CTkEntry(col1, width=250, font=self.font_farsi)
        self.entry_name.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        self.is_subsidiary_var = ctk.BooleanVar(value=False)
        self.subsidiary_checkbox = ctk.CTkCheckBox(col1, text="شرکت زیرمجموعه", variable=self.is_subsidiary_var, font=self.font_farsi, command=self.toggle_parent_combo)
        self.subsidiary_checkbox.grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col1, text="شرکت مادر", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.parent_combo = ctk.CTkComboBox(col1, values=self.get_parent_names(), font=self.font_farsi, dropdown_font=self.font_farsi, width=200, state="disabled")
        self.parent_combo.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col1, text="کد ملی", font=self.font_farsi).grid(row=4, column=0, padx=5, pady=5, sticky="e")
        self.entry_national_id = ctk.CTkEntry(col1, width=150, font=self.font_farsi)
        self.entry_national_id.grid(row=4, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col1, text="کد اقتصادی", font=self.font_farsi).grid(row=5, column=0, padx=5, pady=5, sticky="e")
        self.entry_economic_code = ctk.CTkEntry(col1, width=150, font=self.font_farsi)
        self.entry_economic_code.grid(row=5, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col1, text="شماره ثبت", font=self.font_farsi).grid(row=6, column=0, padx=5, pady=5, sticky="e")
        self.entry_reg_number = ctk.CTkEntry(col1, width=150, font=self.font_farsi)
        self.entry_reg_number.grid(row=6, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col2, text="آدرس", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.entry_address = ctk.CTkEntry(col2, width=300, font=self.font_farsi)
        self.entry_address.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col2, text="کد پستی", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.entry_postal_code = ctk.CTkEntry(col2, width=150, font=self.font_farsi)
        self.entry_postal_code.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col2, text="تلفن", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.entry_phone = ctk.CTkEntry(col2, width=150, font=self.font_farsi)
        self.entry_phone.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col2, text="شماره موبایل", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.entry_mobile = ctk.CTkEntry(col2, width=150, font=self.font_farsi)
        self.entry_mobile.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col2, text="ایمیل", font=self.font_farsi).grid(row=4, column=0, padx=5, pady=5, sticky="e")
        self.entry_email = ctk.CTkEntry(col2, width=200, font=self.font_farsi)
        self.entry_email.grid(row=4, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(col2, text="نوع مشتری", font=self.font_farsi).grid(row=5, column=0, padx=5, pady=5, sticky="e")
        self.customer_type_combo = ctk.CTkComboBox(col2, values=["حقوقی", "حقیقی"], font=self.font_farsi, dropdown_font=self.font_farsi, width=150)
        self.customer_type_combo.grid(row=5, column=1, padx=5, pady=5, sticky="w")
        self.customer_type_combo.set("حقوقی")
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(pady=15)
        self.btn_add = ctk.CTkButton(btn_frame, text="افزودن", command=self.add_payer, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100)
        self.btn_add.pack(side="left", padx=5)
        self.btn_edit = ctk.CTkButton(btn_frame, text="ویرایش", command=self.edit_payer, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100, state="disabled")
        self.btn_edit.pack(side="left", padx=5)
        self.btn_delete = ctk.CTkButton(btn_frame, text="حذف", command=self.delete_payer, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100, state="disabled")
        self.btn_delete.pack(side="left", padx=5)
        self.btn_close = ctk.CTkButton(btn_frame, text="بستن", command=self.close_dialog, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100)
        self.btn_close.pack(side="left", padx=5)
        self.listbox.bind("<<ListboxSelect>>", self.on_select)
    def toggle_parent_combo(self):
        if self.is_subsidiary_var.get():
            self.parent_combo.configure(state="normal")
            self.parent_combo.configure(values=self.get_parent_names())
        else:
            self.parent_combo.configure(state="disabled")
            self.parent_combo.set("")
    def get_parent_names(self):
        items = []
        for p in self.app.payers_list:
            if hasattr(self, 'current_code') and p['code'] == self.current_code:
                continue
            items.append(f"{p['code']} - {p['name']}")
        return items
    def refresh_listbox(self):
        self.listbox.delete(0, "end")
        for payer in self.app.payers_list:
            parent_text = ""
            if payer.get('parent_code'):
                parent = next((p for p in self.app.payers_list if p['code'] == payer['parent_code']), None)
                if parent:
                    parent_text = f" (زیرمجموعه {parent['name']})"
            display_text = f"{payer['code']} - {payer['name']}{parent_text}"
            self.listbox.insert("end", display_text)
    def on_select(self, event):
        selection = self.listbox.curselection()
        if selection:
            self.selected_index = selection[0]
            self.btn_edit.configure(state="normal")
            self.btn_delete.configure(state="normal")
            payer = self.app.payers_list[self.selected_index]
            self.current_code = payer['code']
            self.entry_name.delete(0, "end")
            self.entry_name.insert(0, payer['name'])
            self.entry_code.delete(0, "end")
            self.entry_code.insert(0, payer['code'])
            has_parent = payer.get('parent_code') and payer['parent_code'] != ""
            self.is_subsidiary_var.set(has_parent)
            if has_parent:
                self.parent_combo.configure(state="normal")
                parent = next((p for p in self.app.payers_list if p['code'] == payer['parent_code']), None)
                if parent:
                    self.parent_combo.set(f"{parent['code']} - {parent['name']}")
                else:
                    self.parent_combo.set("")
            else:
                self.parent_combo.configure(state="disabled")
                self.parent_combo.set("")
            self.entry_national_id.delete(0, "end")
            self.entry_national_id.insert(0, payer.get('national_id', ''))
            self.entry_economic_code.delete(0, "end")
            self.entry_economic_code.insert(0, payer.get('economic_code', ''))
            self.entry_reg_number.delete(0, "end")
            self.entry_reg_number.insert(0, payer.get('registration_number', ''))
            self.entry_address.delete(0, "end")
            self.entry_address.insert(0, payer.get('address', ''))
            self.entry_postal_code.delete(0, "end")
            self.entry_postal_code.insert(0, payer.get('postal_code', ''))
            self.entry_phone.delete(0, "end")
            self.entry_phone.insert(0, payer.get('phone', ''))
            self.entry_mobile.delete(0, "end")
            self.entry_mobile.insert(0, payer.get('mobile', ''))
            self.entry_email.delete(0, "end")
            self.entry_email.insert(0, payer.get('email', ''))
            cust_type = payer.get('customer_type', 'legal')
            self.customer_type_combo.set("حقوقی" if cust_type == 'legal' else "حقیقی")
        else:
            self.selected_index = None
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
            self.clear_entries()
    def clear_entries(self):
        self.entry_name.delete(0, "end")
        self.entry_code.delete(0, "end")
        self.is_subsidiary_var.set(False)
        self.parent_combo.configure(state="disabled")
        self.parent_combo.set("")
        self.entry_national_id.delete(0, "end")
        self.entry_economic_code.delete(0, "end")
        self.entry_reg_number.delete(0, "end")
        self.entry_address.delete(0, "end")
        self.entry_postal_code.delete(0, "end")
        self.entry_phone.delete(0, "end")
        self.entry_mobile.delete(0, "end")
        self.entry_email.delete(0, "end")
        self.customer_type_combo.set("حقوقی")
    def add_payer(self):
        name = self.entry_name.get().strip()
        code = self.entry_code.get().strip()
        if not name or not code:
            messagebox.showerror("خطا", "نام و کد شرکت نمی‌تواند خالی باشد.")
            return
        if not code.isdigit():
            messagebox.showerror("خطا", "کد شرکت باید عدد باشد.")
            return
        if len(code) > 5:
            messagebox.showerror("خطا", "کد شرکت حداکثر ۵ رقم می‌تواند باشد.")
            return
        if any(p['code'] == code for p in self.app.payers_list):
            messagebox.showerror("خطا", "این کد قبلاً ثبت شده است.")
            return
        parent_code = None
        if self.is_subsidiary_var.get():
            parent_text = self.parent_combo.get()
            if parent_text:
                parts = parent_text.split(" - ", 1)
                if len(parts) == 2:
                    candidate_code = parts[0].strip()
                    if any(p['code'] == candidate_code for p in self.app.payers_list):
                        parent_code = candidate_code
                    else:
                        messagebox.showerror("خطا", "کد شرکت مادر نامعتبر است.")
                        return
                else:
                    messagebox.showerror("خطا", "فرمت شرکت مادر صحیح نیست.")
                    return
            else:
                messagebox.showerror("خطا", "لطفاً شرکت مادر را انتخاب کنید.")
                return
        national_id = self.entry_national_id.get().strip()
        economic_code = self.entry_economic_code.get().strip()
        reg_number = self.entry_reg_number.get().strip()
        address = self.entry_address.get().strip()
        postal_code = self.entry_postal_code.get().strip()
        phone = self.entry_phone.get().strip()
        mobile = self.entry_mobile.get().strip()
        email = self.entry_email.get().strip()
        customer_type = 'legal' if self.customer_type_combo.get() == "حقوقی" else 'person'
        if self.app.add_payer_db(name, code, parent_code, national_id, economic_code, reg_number, address, postal_code, phone, customer_type, mobile, email):
            self.refresh_listbox()
            self.clear_entries()
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
        else:
            messagebox.showerror("خطا", "خطا در افزودن شرکت.")
    def edit_payer(self):
        if self.selected_index is None:
            return
        name = self.entry_name.get().strip()
        code = self.entry_code.get().strip()
        old_code = self.app.payers_list[self.selected_index]['code']
        if not name or not code:
            messagebox.showerror("خطا", "نام و کد شرکت نمی‌تواند خالی باشد.")
            return
        if not code.isdigit():
            messagebox.showerror("خطا", "کد شرکت باید عدد باشد.")
            return
        if len(code) > 5:
            messagebox.showerror("خطا", "کد شرکت حداکثر ۵ رقم می‌تواند باشد.")
            return
        if any(p['code'] == code and p['code'] != old_code for p in self.app.payers_list):
            messagebox.showerror("خطا", "این کد قبلاً ثبت شده است.")
            return
        parent_code = None
        if self.is_subsidiary_var.get():
            parent_text = self.parent_combo.get()
            if parent_text:
                parts = parent_text.split(" - ", 1)
                if len(parts) == 2:
                    candidate_code = parts[0].strip()
                    if any(p['code'] == candidate_code for p in self.app.payers_list):
                        parent_code = candidate_code
                    else:
                        messagebox.showerror("خطا", "کد شرکت مادر نامعتبر است.")
                        return
                else:
                    messagebox.showerror("خطا", "فرمت شرکت مادر صحیح نیست.")
                    return
            else:
                messagebox.showerror("خطا", "لطفاً شرکت مادر را انتخاب کنید.")
                return
        national_id = self.entry_national_id.get().strip()
        economic_code = self.entry_economic_code.get().strip()
        reg_number = self.entry_reg_number.get().strip()
        address = self.entry_address.get().strip()
        postal_code = self.entry_postal_code.get().strip()
        phone = self.entry_phone.get().strip()
        mobile = self.entry_mobile.get().strip()
        email = self.entry_email.get().strip()
        customer_type = 'legal' if self.customer_type_combo.get() == "حقوقی" else 'person'
        if self.app.update_payer_db(old_code, name, code, parent_code, national_id, economic_code, reg_number, address, postal_code, phone, customer_type, mobile, email):
            self.refresh_listbox()
            self.clear_entries()
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
            self.selected_index = None
        else:
            messagebox.showerror("خطا", "خطا در ویرایش شرکت.")
    def delete_payer(self):
        if self.selected_index is None:
            return
        code = self.app.payers_list[self.selected_index]['code']
        if messagebox.askyesno("تأیید حذف", "آیا از حذف این آیتم اطمینان دارید؟\nتوجه: شرکت‌های زیرمجموعه آن نیز از حالت زیرمجموعه خارج می‌شوند."):
            self.app.delete_payer_db(code)
            self.refresh_listbox()
            self.clear_entries()
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
            self.selected_index = None
    def close_dialog(self):
        self.destroy()

class ManageUsersDialog(ctk.CTkToplevel):
    def __init__(self, parent, user_manager):
        super().__init__(parent)
        self.title("مدیریت کاربران")
        self.geometry("900x750")
        self.transient(parent)
        self.grab_set()
        self.user_manager = user_manager
        self.font_farsi = ("B Nazanin", 16)
        self.font_farsi_bold = ("B Nazanin", 16, "bold")
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(fill="both", expand=True, padx=15, pady=15)
        ctk.CTkLabel(main_frame, text="لیست کاربران", font=("B Nazanin", 20, "bold")).pack(pady=5)
        self.users_frame = ctk.CTkScrollableFrame(main_frame, height=250)
        self.users_frame.pack(fill="both", expand=True, padx=5, pady=5)
        self.refresh_users()
        form_frame = ctk.CTkFrame(main_frame)
        form_frame.pack(pady=10, fill="x")
        ctk.CTkLabel(form_frame, text="نام کاربری", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.entry_username = ctk.CTkEntry(form_frame, width=180, font=self.font_farsi)
        self.entry_username.grid(row=0, column=1, padx=5, pady=5)
        ctk.CTkLabel(form_frame, text="رمز عبور", font=self.font_farsi).grid(row=0, column=2, padx=5, pady=5, sticky="e")
        self.entry_password = ctk.CTkEntry(form_frame, width=180, font=self.font_farsi, show="*")
        self.entry_password.grid(row=0, column=3, padx=5, pady=5)
        chk_frame = ctk.CTkFrame(form_frame, fg_color="transparent")
        chk_frame.grid(row=1, column=0, columnspan=4, pady=5)
        self.var_view = ctk.BooleanVar(value=True)
        self.var_issue = ctk.BooleanVar(value=False)
        self.var_pay = ctk.BooleanVar(value=False)
        self.var_manage_users = ctk.BooleanVar(value=False)
        self.var_manage_settings = ctk.BooleanVar(value=False)
        self.var_delete_history = ctk.BooleanVar(value=False)
        self.var_delete_payment = ctk.BooleanVar(value=False)
        self.var_export = ctk.BooleanVar(value=False)
        self.var_manage_companies = ctk.BooleanVar(value=False)
        self.var_submit = ctk.BooleanVar(value=True)
        self.var_approve = ctk.BooleanVar(value=False)
        self.var_edit_template = ctk.BooleanVar(value=False)  # جدید

        ctk.CTkCheckBox(chk_frame, text="مشاهده", variable=self.var_view, font=self.font_farsi).grid(row=0, column=0, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="ایجاد بدهی", variable=self.var_issue, font=self.font_farsi).grid(row=0, column=1, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="ثبت پرداخت", variable=self.var_pay, font=self.font_farsi).grid(row=0, column=2, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="مدیریت شرکت‌ها", variable=self.var_manage_companies, font=self.font_farsi).grid(row=0, column=3, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="مدیریت کاربران", variable=self.var_manage_users, font=self.font_farsi).grid(row=1, column=0, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="تنظیمات سیستم", variable=self.var_manage_settings, font=self.font_farsi).grid(row=1, column=1, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="حذف تاریخچه", variable=self.var_delete_history, font=self.font_farsi).grid(row=1, column=2, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="حذف پرداخت", variable=self.var_delete_payment, font=self.font_farsi).grid(row=1, column=3, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="خروجی گرفتن", variable=self.var_export, font=self.font_farsi).grid(row=2, column=0, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="ارسال برای تأیید", variable=self.var_submit, font=self.font_farsi).grid(row=2, column=1, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="تأیید درخواست‌ها", variable=self.var_approve, font=self.font_farsi).grid(row=2, column=2, padx=5, sticky="w")
        ctk.CTkCheckBox(chk_frame, text="ویرایش قالب", variable=self.var_edit_template, font=self.font_farsi).grid(row=2, column=3, padx=5, sticky="w")

        role_frame = ctk.CTkFrame(form_frame, fg_color="transparent")
        role_frame.grid(row=2, column=0, columnspan=4, pady=5)
        ctk.CTkLabel(role_frame, text="انتخاب نقش", font=self.font_farsi).pack(side="left", padx=5)
        self.role_combo = ctk.CTkComboBox(role_frame, values=["ادمین", "مدیر", "کاربر", "مدیر شرکت‌ها", "کاربر تأیید"], font=self.font_farsi, dropdown_font=self.font_farsi, width=150, command=self.apply_role_preset)
        self.role_combo.pack(side="left", padx=5)

        btn_frame = ctk.CTkFrame(main_frame)
        btn_frame.pack(pady=10)
        self.btn_add = ctk.CTkButton(btn_frame, text="افزودن", command=self.add_user, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100)
        self.btn_add.pack(side="left", padx=5)
        self.btn_edit = ctk.CTkButton(btn_frame, text="ویرایش", command=self.edit_user, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100, state="disabled")
        self.btn_edit.pack(side="left", padx=5)
        self.btn_delete = ctk.CTkButton(btn_frame, text="حذف", command=self.delete_user, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100, state="disabled")
        self.btn_delete.pack(side="left", padx=5)
        self.btn_close = ctk.CTkButton(btn_frame, text="بستن", command=self.destroy, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100)
        self.btn_close.pack(side="left", padx=5)
        self.selected_user_id = None

    def refresh_users(self):
        for widget in self.users_frame.winfo_children():
            widget.destroy()
        users = self.user_manager.get_all_users()
        if not users:
            ctk.CTkLabel(self.users_frame, text="هیچ کاربری یافت نشد.", font=self.font_farsi).pack()
            return
        headers = ["ردیف", "نام کاربری", "مشاهده", "ایجاد بدهی", "پرداخت", "مدیریت شرکت‌ها", "مدیریت کاربران", "تنظیمات", "حذف تاریخچه", "حذف پرداخت", "خروجی", "ارسال برای تأیید", "تأیید درخواست", "ویرایش قالب", "انتخاب"]
        for col, h in enumerate(headers):
            lbl = ctk.CTkLabel(self.users_frame, text=h, font=self.font_farsi_bold)
            lbl.grid(row=0, column=col, padx=5, pady=5, sticky="ew")
        for i, (user_id, username, can_view, can_pay, can_issue, can_manage_users, can_manage_settings, can_delete_history, can_delete_payment, can_export, can_manage_companies, can_submit, can_approve, can_edit_template) in enumerate(users, start=1):
            row_data = [str(i), username, "✅" if can_view else "❌", "✅" if can_issue else "❌", "✅" if can_pay else "❌", "✅" if can_manage_companies else "❌", "✅" if can_manage_users else "❌", "✅" if can_manage_settings else "❌", "✅" if can_delete_history else "❌", "✅" if can_delete_payment else "❌", "✅" if can_export else "❌", "✅" if can_submit else "❌", "✅" if can_approve else "❌", "✅" if can_edit_template else "❌"]
            for col, val in enumerate(row_data):
                lbl = ctk.CTkLabel(self.users_frame, text=val, font=self.font_farsi)
                lbl.grid(row=i, column=col, padx=5, pady=2, sticky="ew")
            btn_select = ctk.CTkButton(self.users_frame, text="انتخاب", font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=60, command=lambda uid=user_id, u=username, v=can_view, iu=can_issue, p=can_pay, mu=can_manage_users, ms=can_manage_settings, dh=can_delete_history, dp=can_delete_payment, ex=can_export, mc=can_manage_companies, sub=can_submit, appr=can_approve, et=can_edit_template: self.select_user(uid, u, v, iu, p, mu, ms, dh, dp, ex, mc, sub, appr, et))
            btn_select.grid(row=i, column=len(row_data), padx=5, pady=2)
        for col in range(len(headers)):
            self.users_frame.grid_columnconfigure(col, weight=1)

    def select_user(self, user_id, username, can_view, can_issue, can_pay, can_manage_users,
                    can_manage_settings, can_delete_history, can_delete_payment, can_export,
                    can_manage_companies, can_submit, can_approve, can_edit_template):
        self.selected_user_id = user_id
        self.entry_username.delete(0, "end")
        self.entry_username.insert(0, username)
        self.entry_password.delete(0, "end")
        self.var_view.set(can_view)
        self.var_issue.set(can_issue)
        self.var_pay.set(can_pay)
        self.var_manage_users.set(can_manage_users)
        self.var_manage_settings.set(can_manage_settings)
        self.var_delete_history.set(can_delete_history)
        self.var_delete_payment.set(can_delete_payment)
        self.var_export.set(can_export)
        self.var_manage_companies.set(can_manage_companies)
        self.var_submit.set(can_submit)
        self.var_approve.set(can_approve)
        self.var_edit_template.set(can_edit_template)
        self.btn_edit.configure(state="normal")
        self.btn_delete.configure(state="normal")

    def apply_role_preset(self, choice):
        if choice == "ادمین":
            self.var_view.set(True); self.var_issue.set(True); self.var_pay.set(True)
            self.var_manage_users.set(True); self.var_manage_settings.set(True)
            self.var_delete_history.set(True); self.var_delete_payment.set(True)
            self.var_export.set(True); self.var_manage_companies.set(True)
            self.var_submit.set(True); self.var_approve.set(True)
            self.var_edit_template.set(True)
        elif choice == "مدیر":
            self.var_view.set(True); self.var_issue.set(True); self.var_pay.set(True)
            self.var_manage_users.set(False); self.var_manage_settings.set(False)
            self.var_delete_history.set(False); self.var_delete_payment.set(False)
            self.var_export.set(True); self.var_manage_companies.set(True)
            self.var_submit.set(False); self.var_approve.set(False)
            self.var_edit_template.set(False)
        elif choice == "کاربر":
            self.var_view.set(True); self.var_issue.set(True); self.var_pay.set(True)
            self.var_manage_users.set(False); self.var_manage_settings.set(False)
            self.var_delete_history.set(False); self.var_delete_payment.set(False)
            self.var_export.set(True); self.var_manage_companies.set(False)
            self.var_submit.set(False); self.var_approve.set(False)
            self.var_edit_template.set(False)
        elif choice == "کاربر تأیید":
            self.var_view.set(True); self.var_issue.set(False); self.var_pay.set(False)
            self.var_manage_users.set(False); self.var_manage_settings.set(False)
            self.var_delete_history.set(False); self.var_delete_payment.set(False)
            self.var_export.set(False); self.var_manage_companies.set(False)
            self.var_submit.set(True); self.var_approve.set(False)
            self.var_edit_template.set(False)
        elif choice == "مدیر شرکت‌ها":
            self.var_view.set(True); self.var_issue.set(True); self.var_pay.set(True)
            self.var_manage_users.set(False); self.var_manage_settings.set(False)
            self.var_delete_history.set(False); self.var_delete_payment.set(False)
            self.var_export.set(True); self.var_manage_companies.set(True)
            self.var_submit.set(False); self.var_approve.set(False)
            self.var_edit_template.set(False)

    def add_user(self):
        username = self.entry_username.get().strip()
        password = self.entry_password.get()
        if not username or not password:
            messagebox.showerror("خطا", "نام کاربری و رمز عبور نمی‌توانند خالی باشند.")
            return
        success = self.user_manager.add_user(
            username, password,
            self.var_view.get(), self.var_pay.get(), self.var_issue.get(),
            self.var_manage_users.get(), self.var_manage_settings.get(),
            self.var_delete_history.get(), self.var_delete_payment.get(),
            self.var_export.get(), self.var_manage_companies.get(),
            self.var_submit.get(), self.var_approve.get(),
            self.var_edit_template.get()
        )
        if success:
            self.refresh_users()
            self.entry_username.delete(0, "end")
            self.entry_password.delete(0, "end")
            self.var_view.set(True); self.var_issue.set(False); self.var_pay.set(False)
            self.var_manage_users.set(False); self.var_manage_settings.set(False)
            self.var_delete_history.set(False); self.var_delete_payment.set(False)
            self.var_export.set(True); self.var_manage_companies.set(False)
            self.var_submit.set(True); self.var_approve.set(False)
            self.var_edit_template.set(False)
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
            self.selected_user_id = None
        else:
            messagebox.showerror("خطا", "این نام کاربری قبلاً ثبت شده است.")

    def edit_user(self):
        if self.selected_user_id is None:
            return
        username = self.entry_username.get().strip()
        password = self.entry_password.get()
        if not username:
            messagebox.showerror("خطا", "نام کاربری نمی‌تواند خالی باشد.")
            return
        success = self.user_manager.update_user(
            self.selected_user_id, username, password,
            self.var_view.get(), self.var_pay.get(), self.var_issue.get(),
            self.var_manage_users.get(), self.var_manage_settings.get(),
            self.var_delete_history.get(), self.var_delete_payment.get(),
            self.var_export.get(), self.var_manage_companies.get(),
            self.var_submit.get(), self.var_approve.get(),
            self.var_edit_template.get()
        )
        if success:
            self.refresh_users()
            self.entry_username.delete(0, "end")
            self.entry_password.delete(0, "end")
            self.var_view.set(True); self.var_issue.set(False); self.var_pay.set(False)
            self.var_manage_users.set(False); self.var_manage_settings.set(False)
            self.var_delete_history.set(False); self.var_delete_payment.set(False)
            self.var_export.set(True); self.var_manage_companies.set(False)
            self.var_submit.set(True); self.var_approve.set(False)
            self.var_edit_template.set(False)
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
            self.selected_user_id = None
        else:
            messagebox.showerror("خطا", "این نام کاربری قبلاً ثبت شده است.")

    def delete_user(self):
        if self.selected_user_id is None:
            return
        if messagebox.askyesno("تأیید حذف", "آیا از حذف این کاربر اطمینان دارید؟"):
            self.user_manager.delete_user(self.selected_user_id)
            self.refresh_users()
            self.entry_username.delete(0, "end")
            self.entry_password.delete(0, "end")
            self.var_view.set(True); self.var_issue.set(False); self.var_pay.set(False)
            self.var_manage_users.set(False); self.var_manage_settings.set(False)
            self.var_delete_history.set(False); self.var_delete_payment.set(False)
            self.var_export.set(True); self.var_manage_companies.set(False)
            self.var_submit.set(True); self.var_approve.set(False)
            self.var_edit_template.set(False)
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
            self.selected_user_id = None

class ManageServicesDialog(ctk.CTkToplevel):
    def __init__(self, parent, db):
        super().__init__(parent)
        self.db = db
        self.title("مدیریت دسته‌بندی سرویس‌ها")
        self.geometry("900x700")
        self.transient(parent)
        self.grab_set()
        self.font_farsi = ("B Nazanin", 14)
        self.font_farsi_bold = ("B Nazanin", 14, "bold")
        
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # بخش چپ: درخت
        left_frame = ctk.CTkFrame(main_frame, width=400)
        left_frame.pack(side="left", fill="both", expand=True, padx=5, pady=5)
        ctk.CTkLabel(left_frame, text="ساختار درختی سرویس‌ها", font=self.font_farsi_bold).pack(pady=5)
        self.tree_frame = ctk.CTkFrame(left_frame)
        self.tree_frame.pack(fill="both", expand=True)
        
        # تنظیم استایل برای Treeview (رفع مشکل نمایش فارسی)
        style = ttk.Style()
        style.theme_use("default")  # استفاده از تم پیش‌فرض
        style.configure("Persian.Treeview",
                        font=("B Nazanin", 12),
                        rowheight=30,
                        background="#2b2b2b",
                        foreground="white",
                        fieldbackground="#2b2b2b")
        style.configure("Persian.Treeview.Heading",
                        font=("B Nazanin", 13, "bold"),
                        background="#3c3c3c",
                        foreground="white")
        style.map("Persian.Treeview.Heading",
                  background=[("active", "#4c4c4c")])
        
        self.tree = ttk.Treeview(self.tree_frame, columns=("code", "status"), show="tree headings", height=20, style="Persian.Treeview")
        self.tree.heading("#0", text="نام سرویس")
        self.tree.heading("code", text="کد")
        self.tree.heading("status", text="وضعیت")
        self.tree.column("#0", width=250)
        self.tree.column("code", width=100)
        self.tree.column("status", width=80)
        vsb = ttk.Scrollbar(self.tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        self.tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        
        # بخش راست: فرم
        right_frame = ctk.CTkFrame(main_frame, width=400)
        right_frame.pack(side="right", fill="both", expand=True, padx=5, pady=5)
        ctk.CTkLabel(right_frame, text="افزودن/ویرایش سرویس", font=self.font_farsi_bold).pack(pady=5)
        
        form_frame = ctk.CTkFrame(right_frame)
        form_frame.pack(pady=10, padx=10, fill="x")
        
        ctk.CTkLabel(form_frame, text="نام سرویس:", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.name_entry = ctk.CTkEntry(form_frame, width=200, font=self.font_farsi)
        self.name_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        ctk.CTkLabel(form_frame, text="کد سرویس:", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.code_entry = ctk.CTkEntry(form_frame, width=200, font=self.font_farsi)
        self.code_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        ctk.CTkLabel(form_frame, text="سرویس والد:", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.parent_combo = ctk.CTkComboBox(form_frame, values=["(بدون والد)"], font=self.font_farsi, width=200)
        self.parent_combo.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        
        ctk.CTkLabel(form_frame, text="ترتیب نمایش:", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.order_entry = ctk.CTkEntry(form_frame, width=200, font=self.font_farsi)
        self.order_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        
        self.active_var = ctk.BooleanVar(value=True)
        self.active_check = ctk.CTkCheckBox(form_frame, text="فعال", variable=self.active_var, font=self.font_farsi)
        self.active_check.grid(row=4, column=1, padx=5, pady=5, sticky="w")
        
        btn_frame = ctk.CTkFrame(right_frame)
        btn_frame.pack(pady=10)
        self.btn_add = ctk.CTkButton(btn_frame, text="افزودن", command=self.add_service, font=self.font_farsi, fg_color=COLOR_PRIMARY, width=100)
        self.btn_add.pack(side="left", padx=5)
        self.btn_edit = ctk.CTkButton(btn_frame, text="ویرایش", command=self.edit_service, font=self.font_farsi, fg_color=COLOR_PRIMARY, width=100, state="disabled")
        self.btn_edit.pack(side="left", padx=5)
        self.btn_delete = ctk.CTkButton(btn_frame, text="حذف", command=self.delete_service, font=self.font_farsi, fg_color=COLOR_ACCENT, width=100, state="disabled")
        self.btn_delete.pack(side="left", padx=5)
        self.btn_refresh = ctk.CTkButton(btn_frame, text="بازخوانی", command=self.load_tree, font=self.font_farsi, fg_color=COLOR_PRIMARY, width=100)
        self.btn_refresh.pack(side="left", padx=5)
        
        self.load_tree()
        self.tree.bind("<<TreeviewSelect>>", self.on_select)
    
    def load_tree(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        tree_data = self.db.get_service_category_tree()
        all_cats = self.db.get_all_service_categories(include_inactive=True)
        parent_items = ["(بدون والد)"] + [f"{c['id']} - {c['name']}" for c in all_cats]
        self.parent_combo.configure(values=parent_items)
        self.parent_combo.set("(بدون والد)")
        for cat in tree_data:
            self._add_tree_node(cat)
    
    def _add_tree_node(self, cat, parent=""):
        status = "فعال" if cat['is_active'] else "غیرفعال"
        node_id = self.tree.insert(parent, "end", text=cat['name'], values=(cat['code'], status), iid=str(cat['id']))
        for child in sorted(cat.get('children', []), key=lambda x: x.get('display_order', 0)):
            self._add_tree_node(child, node_id)
    
    def on_select(self, event):
        selected = self.tree.selection()
        if not selected:
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
            return
        self.btn_edit.configure(state="normal")
        self.btn_delete.configure(state="normal")
        cat_id = int(selected[0])
        all_cats = self.db.get_all_service_categories(include_inactive=True)
        current = next((c for c in all_cats if c['id'] == cat_id), None)
        if current:
            self.name_entry.delete(0, "end")
            self.name_entry.insert(0, current['name'])
            self.code_entry.delete(0, "end")
            self.code_entry.insert(0, current['code'])
            self.order_entry.delete(0, "end")
            self.order_entry.insert(0, str(current.get('display_order', 0)))
            self.active_var.set(current['is_active'])
            if current['parent_id']:
                parent_text = f"{current['parent_id']} - {next((p['name'] for p in all_cats if p['id'] == current['parent_id']), '')}"
                self.parent_combo.set(parent_text)
            else:
                self.parent_combo.set("(بدون والد)")
            self.current_id = cat_id
    
    def add_service(self):
        name = self.name_entry.get().strip()
        code = self.code_entry.get().strip()
        if not name or not code:
            messagebox.showerror("خطا", "نام و کد سرویس الزامی است")
            return
        parent_text = self.parent_combo.get()
        parent_id = None
        if parent_text != "(بدون والد)":
            try:
                parent_id = int(parent_text.split(" - ")[0])
            except:
                messagebox.showerror("خطا", "والد نامعتبر")
                return
        try:
            display_order = int(self.order_entry.get()) if self.order_entry.get().strip() else 0
        except:
            display_order = 0
        is_active = self.active_var.get()
        try:
            self.db.add_service_category(name, code, parent_id, is_active, display_order)
            self.load_tree()
            self.clear_form()
            messagebox.showinfo("موفقیت", "سرویس با موفقیت اضافه شد")
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در افزودن: {str(e)}")
    
    def edit_service(self):
        if not hasattr(self, 'current_id'):
            return
        name = self.name_entry.get().strip()
        code = self.code_entry.get().strip()
        if not name or not code:
            messagebox.showerror("خطا", "نام و کد سرویس الزامی است")
            return
        parent_text = self.parent_combo.get()
        parent_id = None
        if parent_text != "(بدون والد)":
            try:
                parent_id = int(parent_text.split(" - ")[0])
            except:
                messagebox.showerror("خطا", "والد نامعتبر")
                return
        try:
            display_order = int(self.order_entry.get()) if self.order_entry.get().strip() else 0
        except:
            display_order = 0
        is_active = self.active_var.get()
        try:
            self.db.update_service_category(self.current_id, name, code, parent_id, is_active, display_order)
            self.load_tree()
            self.clear_form()
            self.btn_edit.configure(state="disabled")
            self.btn_delete.configure(state="disabled")
            messagebox.showinfo("موفقیت", "سرویس با موفقیت ویرایش شد")
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در ویرایش: {str(e)}")
    
    def delete_service(self):
        if not hasattr(self, 'current_id'):
            return
        if messagebox.askyesno("تأیید حذف", "آیا از حذف این سرویس اطمینان دارید؟\nتوجه: زیرمجموعه‌ها نیز حذف خواهند شد."):
            try:
                self.db.delete_service_category(self.current_id)
                self.load_tree()
                self.clear_form()
                self.btn_edit.configure(state="disabled")
                self.btn_delete.configure(state="disabled")
                messagebox.showinfo("موفقیت", "سرویس حذف شد")
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در حذف: {str(e)}")
    
    def clear_form(self):
        self.name_entry.delete(0, "end")
        self.code_entry.delete(0, "end")
        self.order_entry.delete(0, "end")
        self.active_var.set(True)
        self.parent_combo.set("(بدون والد)")
        if hasattr(self, 'current_id'):
            delattr(self, 'current_id')

class SettingsDialog(ctk.CTkToplevel):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.title("تنظیمات سیستم")
        self.geometry("900x800")
        self.transient(parent)
        self.grab_set()
        self.font_farsi = ("B Nazanin", 16)
        self.tabview = ctk.CTkTabview(self, segmented_button_fg_color=COLOR_PRIMARY)
        self.tabview._segmented_button.configure(font=self.font_farsi)
        self.tabview.pack(fill="both", expand=True, padx=15, pady=15)
        self.tab_format = self.tabview.add("فرمت شماره نامه")
        self.create_format_tab()
        self.tab_users = self.tabview.add("مدیریت کاربران")
        self.create_users_tab()
        self.tab_database = self.tabview.add("پایگاه داده")
        self.create_database_tab()
        self.tab_company = self.tabview.add("اطلاعات شرکت")
        self.create_company_tab()
        self.tab_services = self.tabview.add("مدیریت سرویس‌ها")
        self.create_services_tab()
        self.btn_close = ctk.CTkButton(self, text="بستن", command=self.destroy, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120)
        self.btn_close.pack(pady=10)
    def create_format_tab(self):
        frame = ctk.CTkFrame(self.tab_format)
        frame.pack(pady=20, padx=20, fill="both", expand=True)
        ctk.CTkLabel(frame, text="فرمت شماره نامه", font=self.font_farsi).pack(anchor="w", pady=5)
        self.format_entry = ctk.CTkEntry(frame, width=350, font=self.font_farsi)
        self.format_entry.pack(anchor="w", pady=5)
        self.format_entry.insert(0, self.app.settings.get("letter_format", "YYYY/CODE/NNNNN"))
        ctk.CTkLabel(frame, text="توضیحات:\nYYYY: سال چهار رقمی\nYY: دو رقم آخر سال\nCODE: کد شرکت (بدون تغییر)\nCCC: کد شرکت سه‌رقمی با صفر\nNUM: شماره پیاپی (بدون صفر)\nNNNNN: شماره پیاپی پنج‌رقمی با صفر\nمی‌توانید از جداکننده‌های / یا - استفاده کنید.", font=self.font_farsi, text_color="gray", justify="right").pack(anchor="w", pady=5)
        self.btn_save_format = ctk.CTkButton(frame, text="ذخیره", command=self.save_format, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT)
        self.btn_save_format.pack(pady=10)
    def create_users_tab(self):
        if not self.app.can_manage_users:
            frame = ctk.CTkFrame(self.tab_users)
            frame.pack(fill="both", expand=True)
            ctk.CTkLabel(frame, text="شما مجوز مدیریت کاربران را ندارید.", font=self.font_farsi, text_color="red").pack(pady=50)
            return
        frame = ctk.CTkFrame(self.tab_users)
        frame.pack(pady=20, padx=20, fill="both", expand=True)
        ctk.CTkLabel(frame, text="برای مدیریت کاربران روی دکمه زیر کلیک کنید:", font=self.font_farsi).pack(pady=20)
        btn_open = ctk.CTkButton(frame, text="مدیریت کاربران", command=self.open_users_dialog, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT)
        btn_open.pack(pady=10)
    def create_database_tab(self):
        frame = ctk.CTkFrame(self.tab_database)
        frame.pack(pady=20, padx=20, fill="both", expand=True)
        ctk.CTkLabel(frame, text="برای مدیریت پایگاه داده روی دکمه زیر کلیک کنید:", font=self.font_farsi).pack(pady=20)
        btn_open = ctk.CTkButton(frame, text="مدیریت پایگاه داده", command=self.open_database_dialog, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT)
        btn_open.pack(pady=10)
    def create_company_tab(self):
        frame = ctk.CTkFrame(self.tab_company)
        frame.pack(pady=20, padx=20, fill="both", expand=True)
        scroll = ctk.CTkScrollableFrame(frame)
        scroll.pack(fill="both", expand=True)

        # تعریف همه فیلدها
        self.entry_company_name = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_company_national_id = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_company_economic_code = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_company_reg_number = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_company_address = ctk.CTkEntry(scroll, width=500, font=self.font_farsi)
        self.entry_company_postal = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_company_phone = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_bank1_name = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_bank1_account = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_bank1_shaba = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_bank2_name = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_bank2_account = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_bank2_shaba = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_manager_name = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_manager_position = ctk.CTkEntry(scroll, width=300, font=self.font_farsi)
        self.entry_approved_path = ctk.CTkEntry(scroll, width=500, font=self.font_farsi)

        # چیدمان (grid)
        labels = [
            ("نام شرکت", self.entry_company_name, 0),
            ("شناسه ملی", self.entry_company_national_id, 1),
            ("کد اقتصادی", self.entry_company_economic_code, 2),
            ("شماره ثبت", self.entry_company_reg_number, 3),
            ("آدرس", self.entry_company_address, 4),
            ("کد پستی", self.entry_company_postal, 5),
            ("تلفن", self.entry_company_phone, 6),
            ("نام بانک اول", self.entry_bank1_name, 7),
            ("شماره حساب اول", self.entry_bank1_account, 8),
            ("شماره شبا اول", self.entry_bank1_shaba, 9),
            ("نام بانک دوم", self.entry_bank2_name, 10),
            ("شماره حساب دوم", self.entry_bank2_account, 11),
            ("شماره شبا دوم", self.entry_bank2_shaba, 12),
            ("نام مدیر", self.entry_manager_name, 13),
            ("سمت مدیر", self.entry_manager_position, 14),
            ("مسیر ذخیره فایل‌های تأیید شده", self.entry_approved_path, 15)
        ]
        for label, entry, row in labels:
            ctk.CTkLabel(scroll, text=label, font=self.font_farsi).grid(row=row, column=0, padx=5, pady=5, sticky="e")
            entry.grid(row=row, column=1, padx=5, pady=5, sticky="w")

        # دکمه انتخاب مسیر برای approved_path
        btn_select_approved = ctk.CTkButton(scroll, text="انتخاب", command=self.select_approved_path, font=self.font_farsi, fg_color=COLOR_PRIMARY)
        btn_select_approved.grid(row=15, column=2, padx=5, pady=5)

        # متغیرهای مسیر فایل‌ها
        self.logo_path = ctk.StringVar()
        self.signature_path = ctk.StringVar()
        self.stamp_path = ctk.StringVar()

        def select_logo():
            filename = filedialog.askopenfilename(filetypes=[("PNG files", "*.png")])
            if filename:
                dest_dir = os.path.join(get_base_path(), "assets")
                os.makedirs(dest_dir, exist_ok=True)
                dest_path = os.path.join(dest_dir, "logo.png")
                shutil.copy2(filename, dest_path)
                self.logo_path.set(dest_path)
                messagebox.showinfo("اطلاع", "لوگو با موفقیت بارگذاری شد.")

        def select_signature():
            filename = filedialog.askopenfilename(filetypes=[("PNG files", "*.png")])
            if filename:
                dest_dir = os.path.join(get_base_path(), "assets")
                os.makedirs(dest_dir, exist_ok=True)
                dest_path = os.path.join(dest_dir, "signature.png")
                shutil.copy2(filename, dest_path)
                self.signature_path.set(dest_path)
                messagebox.showinfo("اطلاع", "امضا با موفقیت بارگذاری شد.")

        def select_stamp():
            filename = filedialog.askopenfilename(filetypes=[("PNG files", "*.png")])
            if filename:
                dest_dir = os.path.join(get_base_path(), "assets")
                os.makedirs(dest_dir, exist_ok=True)
                dest_path = os.path.join(dest_dir, "stamp.png")
                shutil.copy2(filename, dest_path)
                self.stamp_path.set(dest_path)
                messagebox.showinfo("اطلاع", "مهر با موفقیت بارگذاری شد.")

        ctk.CTkButton(scroll, text="انتخاب لوگو", command=select_logo, font=self.font_farsi, fg_color=COLOR_PRIMARY).grid(row=16, column=0, padx=5, pady=5)
        ctk.CTkLabel(scroll, textvariable=self.logo_path, font=self.font_farsi).grid(row=16, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkButton(scroll, text="انتخاب امضا", command=select_signature, font=self.font_farsi, fg_color=COLOR_PRIMARY).grid(row=17, column=0, padx=5, pady=5)
        ctk.CTkLabel(scroll, textvariable=self.signature_path, font=self.font_farsi).grid(row=17, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkButton(scroll, text="انتخاب مهر", command=select_stamp, font=self.font_farsi, fg_color=COLOR_PRIMARY).grid(row=18, column=0, padx=5, pady=5)
        ctk.CTkLabel(scroll, textvariable=self.stamp_path, font=self.font_farsi).grid(row=18, column=1, padx=5, pady=5, sticky="w")

        # بارگذاری تنظیمات از دیتابیس و پر کردن فیلدها
        settings = self.app.db.get_company_settings()
        if settings:
            # تابع کمکی برای درج ایمن
            def safe_insert(entry, value):
                try:
                    entry.delete(0, 'end')
                    entry.insert(0, str(value))
                except Exception:
                    pass  # در صورت بروز خطا نادیده گرفته شود

            safe_insert(self.entry_company_name, settings.get('company_name', ''))
            safe_insert(self.entry_company_national_id, settings.get('national_id', ''))
            safe_insert(self.entry_company_economic_code, settings.get('economic_code', ''))
            safe_insert(self.entry_company_reg_number, settings.get('registration_number', ''))
            safe_insert(self.entry_company_address, settings.get('address', ''))
            safe_insert(self.entry_company_postal, settings.get('postal_code', ''))
            safe_insert(self.entry_company_phone, settings.get('phone', ''))
            safe_insert(self.entry_bank1_name, settings.get('bank_name1', ''))
            safe_insert(self.entry_bank1_account, settings.get('account_number1', ''))
            safe_insert(self.entry_bank1_shaba, settings.get('shaba_number1', ''))
            safe_insert(self.entry_bank2_name, settings.get('bank_name2', ''))
            safe_insert(self.entry_bank2_account, settings.get('account_number2', ''))
            safe_insert(self.entry_bank2_shaba, settings.get('shaba_number2', ''))
            safe_insert(self.entry_manager_name, settings.get('manager_name', ''))
            safe_insert(self.entry_manager_position, settings.get('manager_position', ''))
            safe_insert(self.entry_approved_path, settings.get('approved_output_path', ''))
            self.logo_path.set(settings.get('logo_path', ''))
            self.signature_path.set(settings.get('signature_path', ''))
            self.stamp_path.set(settings.get('stamp_path', ''))

        ctk.CTkButton(scroll, text="ذخیره تنظیمات شرکت", command=self.save_company_settings, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT).grid(row=19, column=0, columnspan=2, pady=20)
    def select_approved_path(self):
        folder = filedialog.askdirectory()
        if folder:
            self.entry_approved_path.delete(0, "end")
            self.entry_approved_path.insert(0, folder)
    def save_format(self):
        new_format = self.format_entry.get().strip()
        if new_format:
            self.app.settings["letter_format"] = new_format
            self.app.save_settings(self.app.settings)
            messagebox.showinfo("موفقیت", "فرمت شماره نامه با موفقیت ذخیره شد.")
        else:
            messagebox.showerror("خطا", "فرمت نمی‌تواند خالی باشد.")
    def save_company_settings(self):
        settings = {
            'company_name': self.entry_company_name.get().strip(),
            'national_id': self.entry_company_national_id.get().strip(),
            'economic_code': self.entry_company_economic_code.get().strip(),
            'registration_number': self.entry_company_reg_number.get().strip(),
            'address': self.entry_company_address.get().strip(),
            'postal_code': self.entry_company_postal.get().strip(),
            'phone': self.entry_company_phone.get().strip(),
            'bank_name1': self.entry_bank1_name.get().strip(),
            'account_number1': self.entry_bank1_account.get().strip(),
            'shaba_number1': self.entry_bank1_shaba.get().strip(),
            'bank_name2': self.entry_bank2_name.get().strip(),
            'account_number2': self.entry_bank2_account.get().strip(),
            'shaba_number2': self.entry_bank2_shaba.get().strip(),
            'manager_name': self.entry_manager_name.get().strip(),
            'manager_position': self.entry_manager_position.get().strip(),
            'logo_path': self.logo_path.get(),
            'signature_path': self.signature_path.get(),
            'stamp_path': self.stamp_path.get(),
            'approved_output_path': self.entry_approved_path.get().strip()
        }
        try:
            self.app.db.save_company_settings(settings)
            messagebox.showinfo("موفقیت", "تنظیمات شرکت با موفقیت ذخیره شد.")
        except Exception as e:
            logging.error(f"خطا در ذخیره تنظیمات شرکت: {e}")
            messagebox.showerror("خطا", f"خطا در ذخیره تنظیمات شرکت:\n{str(e)}")
    def open_users_dialog(self):
        if self.app.can_manage_users:
            dialog = ManageUsersDialog(self, self.app.user_manager)
            self.wait_window(dialog)
        else:
            messagebox.showerror("خطا", "شما مجوز مدیریت کاربران را ندارید.")
    def open_database_dialog(self):
        dialog = DatabaseSettingsDialog(self, self.app.db_manager)
        self.wait_window(dialog)
        
    def create_services_tab(self):
        frame = ctk.CTkFrame(self.tab_services)
        frame.pack(pady=20, padx=20, fill="both", expand=True)
        ctk.CTkLabel(frame, text="مدیریت دسته‌بندی سرویس‌ها (پرواز، هتل، ویزا، ...)", font=self.font_farsi).pack(pady=20)
        btn_open = ctk.CTkButton(frame, text="مدیریت سرویس‌ها", command=self.open_services_dialog, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT)
        btn_open.pack(pady=10)

    def open_services_dialog(self):
        dialog = ManageServicesDialog(self, self.app.db)
        self.wait_window(dialog)
        # پس از بستن، در صورت نیاز لیست سرویس‌های برنامه را به‌روز کنید
        self.app.service_names = self.app.get_service_names()  # اگر در App نگهداری می‌کنید

class DatabaseSettingsDialog(ctk.CTkToplevel):
    def __init__(self, parent, db_manager):
        super().__init__(parent)
        self.title("مدیریت پایگاه داده")
        self.geometry("500x450")
        self.transient(parent)
        self.grab_set()
        self.db_manager = db_manager
        self.font_farsi = ("B Nazanin", 16)
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(pady=15, padx=15, fill="both", expand=True)
        ctk.CTkLabel(main_frame, text="تنظیمات اتصال SQL Server", font=("B Nazanin", 18, "bold")).pack(pady=(5, 15))
        conn_frame = ctk.CTkFrame(main_frame)
        conn_frame.pack(fill="x", pady=5, padx=5)
        ctk.CTkLabel(conn_frame, text="آدرس سرور", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.server_entry = ctk.CTkEntry(conn_frame, width=250, font=self.font_farsi)
        self.server_entry.grid(row=0, column=1, padx=5, pady=5)
        self.server_entry.insert(0, self.db_manager.config.get('server', ''))
        ctk.CTkLabel(conn_frame, text="نام دیتابیس", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.db_entry = ctk.CTkEntry(conn_frame, width=250, font=self.font_farsi)
        self.db_entry.grid(row=1, column=1, padx=5, pady=5)
        self.db_entry.insert(0, self.db_manager.config.get('database', ''))
        ctk.CTkLabel(conn_frame, text="نام کاربری", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.user_entry = ctk.CTkEntry(conn_frame, width=250, font=self.font_farsi)
        self.user_entry.grid(row=2, column=1, padx=5, pady=5)
        self.user_entry.insert(0, self.db_manager.config.get('username', ''))
        ctk.CTkLabel(conn_frame, text="رمز عبور", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.pass_entry = ctk.CTkEntry(conn_frame, width=250, font=self.font_farsi, show="*")
        self.pass_entry.grid(row=3, column=1, padx=5, pady=5)
        self.pass_entry.insert(0, self.db_manager.config.get('password', ''))
        ctk.CTkButton(conn_frame, text="ذخیره تنظیمات", command=self.save_db_config, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT).grid(row=4, column=0, columnspan=2, pady=15)
        backup_frame = ctk.CTkFrame(main_frame)
        backup_frame.pack(fill="x", pady=10, padx=5)
        ctk.CTkLabel(backup_frame, text="پشتیبان‌گیری و بازیابی", font=("B Nazanin", 17, "bold")).pack(anchor="w", pady=(10, 5), padx=5)
        btn_backup = ctk.CTkButton(backup_frame, text="📀 تهیه پشتیبان", command=self.backup_database, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=170)
        btn_backup.pack(side="left", padx=5, pady=5)
        btn_restore = ctk.CTkButton(backup_frame, text="♻️ بازیابی از پشتیبان", command=self.restore_database, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=170)
        btn_restore.pack(side="left", padx=5, pady=5)
        ctk.CTkButton(main_frame, text="بستن", command=self.destroy, font=self.font_farsi, fg_color="gray", hover_color="#555555", width=120).pack(pady=10)
    def save_db_config(self):
        config = {'server': self.server_entry.get().strip(), 'database': self.db_entry.get().strip(), 'username': self.user_entry.get().strip(), 'password': self.pass_entry.get().strip(), 'api_server': self.api_entry.get().strip()}
        try:
            conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={config['server']};DATABASE={config['database']};UID={config['username']};PWD={config['password']}"
            pyodbc.connect(conn_str, timeout=5)
        except Exception as e:
            messagebox.showerror("خطا", f"اتصال به سرور ناموفق بود:\n{str(e)}")
            return
        base_path = get_base_path()
        config_path = os.path.join(base_path, 'db_config.json')
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        self.db_manager.config = config
        messagebox.showinfo("موفقیت", "تنظیمات با موفقیت ذخیره شد.")
    def backup_database(self):
        ext = ".bak"
        initial_name = generate_backup_filename() + ext
        file_path = filedialog.asksaveasfilename(title="ذخیره فایل پشتیبان", initialfile=initial_name, defaultextension=ext, filetypes=[("Backup files", f"*{ext}"), ("All files", "*.*")])
        if not file_path:
            return
        try:
            self.db_manager.backup_database(file_path)
            if messagebox.askyesno("موفقیت", f"پشتیبان با موفقیت در مسیر زیر ذخیره شد:\n{file_path}\n\nآیا می‌خواهید پوشه حاوی فایل باز شود؟"):
                try:
                    os.startfile(os.path.dirname(file_path))
                except:
                    pass
            else:
                messagebox.showinfo("موفقیت", f"فایل با موفقیت ذخیره شد:\n{file_path}")
        except Exception as e:
            messagebox.showerror("خطا", f"پشتیبان‌گیری ناموفق بود:\n{str(e)}")
    def restore_database(self):
        if not messagebox.askyesno("تأیید", "بازیابی دیتابیس تمام اطلاعات فعلی را بازنویسی خواهد کرد. آیا اطمینان دارید؟"):
            return
        ext = ".bak"
        file_path = filedialog.askopenfilename(title="انتخاب فایل پشتیبان", filetypes=[("Backup files", f"*{ext}"), ("All files", "*.*")])
        if not file_path:
            return
        try:
            self.db_manager.restore_database(file_path)
            messagebox.showinfo("موفقیت", "بازیابی با موفقیت انجام شد.")
        except Exception as e:
            messagebox.showerror("خطا", f"بازیابی ناموفق بود:\n{str(e)}")

class AddPaymentDialog(ctk.CTkToplevel):
    def __init__(self, parent, payer_code, payer_name, on_save):
        super().__init__(parent)
        self.title("ثبت پرداخت جدید")
        self.geometry("450x450")
        self.transient(parent)
        self.grab_set()
        self.payer_code = payer_code
        self.payer_name = payer_name
        self.on_save = on_save
        self.font_farsi = ("B Nazanin", 18)

        ctk.CTkLabel(self, text=f"ثبت پرداخت برای {payer_name} (کد: {payer_code})", font=("B Nazanin", 20, "bold")).pack(pady=15)

        main_frame = ctk.CTkFrame(self)
        main_frame.pack(pady=15, padx=20, fill="both", expand=True)

        # تاریخ
        ctk.CTkLabel(main_frame, text="تاریخ پرداخت", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        date_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        date_frame.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        today = jdatetime.date.today()
        self.year_var = ctk.StringVar(value=persian_number(str(today.year)))
        self.month_var = ctk.StringVar(value=persian_number(str(today.month).zfill(2)))
        self.day_var = ctk.StringVar(value=persian_number(str(today.day).zfill(2)))
        year_combo = ctk.CTkComboBox(date_frame, values=[persian_number(str(y)) for y in range(1400, 1431)], variable=self.year_var, width=90, font=self.font_farsi, dropdown_font=self.font_farsi)
        year_combo.pack(side="left", padx=2)
        month_combo = ctk.CTkComboBox(date_frame, values=[persian_number(str(m).zfill(2)) for m in range(1, 13)], variable=self.month_var, width=70, font=self.font_farsi, dropdown_font=self.font_farsi)
        month_combo.pack(side="left", padx=2)
        day_combo = ctk.CTkComboBox(date_frame, values=[persian_number(str(d).zfill(2)) for d in range(1, 32)], variable=self.day_var, width=70, font=self.font_farsi, dropdown_font=self.font_farsi)
        day_combo.pack(side="left", padx=2)

        # مبلغ
        ctk.CTkLabel(main_frame, text="مبلغ (ریال)", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.amount_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.amount_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")

        # شماره پیگیری
        ctk.CTkLabel(main_frame, text="شماره پیگیری", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.tracking_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.tracking_entry.grid(row=2, column=1, padx=5, pady=5, sticky="w")

        # توضیحات
        ctk.CTkLabel(main_frame, text="توضیحات", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.desc_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.desc_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")

        # دکمه‌ها
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(side="bottom", pady=15, fill="x")
        self.save_btn = ctk.CTkButton(btn_frame, text="💾 ذخیره", command=self.save, font=("B Nazanin", 18, "bold"), fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=140)
        self.save_btn.pack(side="left", padx=5, expand=True)
        ctk.CTkButton(btn_frame, text="انصراف", command=self.destroy, font=("B Nazanin", 18, "bold"), fg_color="gray", hover_color="#555555", width=140).pack(side="left", padx=5, expand=True)

    # ⚠️ مهم: این متد حتماً با self تعریف شده است
    def save(self):
        year = self.year_var.get()
        month = self.month_var.get()
        day = self.day_var.get()
        date = f"{year}/{month}/{day}"
        try:
            amount = int(english_number(self.amount_entry.get()))
        except:
            messagebox.showerror("خطا", "مبلغ را به درستی وارد کنید.")
            return
        tracking = self.tracking_entry.get().strip()
        desc = self.desc_entry.get().strip()
        if not tracking:
            messagebox.showerror("خطا", "شماره پیگیری نمی‌تواند خالی باشد.")
            return
        self.on_save(self.payer_code, self.payer_name, date, amount, tracking, desc)
        self.destroy()

class CompanyTransactionsDialog(ctk.CTkToplevel):
    def __init__(self, parent, app, payer_code, payer_name, can_pay, can_delete_payment):
        super().__init__(parent)
        self.app = app
        self.title(f"تراکنش‌های {payer_name}")
        self.geometry("1400x700")
        self.transient(parent)
        self.grab_set()
        self.db = app.db
        self.payer_code = payer_code
        self.payer_name = payer_name
        self.can_pay = can_pay
        self.can_delete_payment = can_delete_payment
        self.font_farsi = ("B Nazanin", 16)
        self.font_farsi_bold = ("B Nazanin", 16, "bold")
        self.manual_debt_ids = []
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(fill="both", expand=True, padx=15, pady=15)
        ctk.CTkLabel(main_frame, text=f"لیست تراکنش‌های {payer_name}", font=("B Nazanin", 20, "bold")).pack(pady=5)
        btn_frame = ctk.CTkFrame(main_frame)
        btn_frame.pack(pady=5, fill="x")
        ctk.CTkButton(btn_frame, text="📥 خروجی اکسل", command=self.export_excel, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=140).pack(side="right", padx=5)
        ctk.CTkButton(btn_frame, text="➕ افزودن بدهی دستی", command=self.add_manual_debt, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=140).pack(side="right", padx=5)
        self.tree_frame = ctk.CTkScrollableFrame(main_frame)
        self.tree_frame.pack(fill="both", expand=True, padx=5, pady=5)
        headers = ["ردیف", "تاریخ", "نوع", "بدهکار (ریال)", "بستانکار (ریال)", "شماره مرجع", "توضیحات", "عملیات"]
        for col, header in enumerate(headers):
            lbl = ctk.CTkLabel(self.tree_frame, text=header, font=self.font_farsi, anchor="center")
            lbl.grid(row=0, column=col, padx=5, pady=5, sticky="ew")
        self.load_transactions()

    def load_transactions(self):
        # پاک کردن کامل محتویات tree_frame
        for widget in self.tree_frame.winfo_children():
            widget.destroy()

        try:
            invoices = self.db.get_invoices(filters={'payer_code': self.payer_code})
            payments = self.db.get_payments(self.payer_code)
            credits = self.db.get_credits(self.payer_code)
            manual_debts = self.db.get_manual_debts(self.payer_code)
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در خواندن داده‌ها:\n{str(e)}")
            return

        # تابع کمکی برای تبدیل میلادی به شمسی
        def to_persian_date(date_val):
            if date_val is None:
                return ""
            try:
                if hasattr(date_val, 'year'):
                    greg = date_val
                else:
                    parts = str(date_val).split('-')
                    if len(parts) == 3:
                        greg = datetime.date(int(parts[0]), int(parts[1]), int(parts[2]))
                    else:
                        return str(date_val)
                persian = jdatetime.date.fromgregorian(date=greg)
                return f"{persian.year}/{persian.month:02d}/{persian.day:02d}"
            except:
                return str(date_val)

        # ساخت لیست تراکنش‌ها با ذخیره شناسه
        transactions = []  # هر عنصر: (date, type, amount, ref, desc, record_id, record_type)
        for inv in invoices:
            persian_date = to_persian_date(inv[6])
            transactions.append((persian_date, 'بدهی', inv[5], inv[1], '', inv[0], 'invoice'))
        for pay in payments:
            persian_date = to_persian_date(pay[1])
            transactions.append((persian_date, 'پرداخت', pay[2], pay[3], pay[4], pay[0], 'payment'))
        for cr in credits:
            persian_date = to_persian_date(cr[2])
            transactions.append((persian_date, 'بستانکاری', cr[1], cr[4], cr[3], cr[0], 'credit'))
        for debt in manual_debts:
            persian_date = to_persian_date(debt[3])
            transactions.append((persian_date, 'بدهی دستی', debt[2], debt[1], debt[4], debt[0], 'manual_debt'))

        if not transactions:
            ctk.CTkLabel(self.tree_frame, text="هیچ تراکنشی برای این شرکت یافت نشد.",
                        font=self.font_farsi, text_color="gray").pack(pady=20)
            return

        # مرتب‌سازی بر اساس تاریخ (نزولی)
        transactions.sort(key=lambda x: x[0], reverse=True)

        # هدرهای جدول
        headers = ["ردیف", "تاریخ", "نوع", "بدهکار (ریال)", "بستانکار (ریال)", "شماره مرجع", "توضیحات", "عملیات"]
        for col, header in enumerate(headers):
            lbl = ctk.CTkLabel(self.tree_frame, text=header, font=self.font_farsi_bold, anchor="center")
            lbl.grid(row=0, column=col, padx=5, pady=5, sticky="ew")

        total_debt = 0
        total_credit = 0

        for i, (date_str, typ, amount, ref, desc, record_id, record_type) in enumerate(transactions, start=1):
            if typ in ('بدهی', 'بدهی دستی'):
                debt = amount
                credit = 0
                total_debt += amount
            else:
                debt = 0
                credit = amount
                total_credit += amount

            debt_str = to_persian_digits(f"{debt:,}") if debt > 0 else ""
            credit_str = to_persian_digits(f"{credit:,}") if credit > 0 else ""
            data = [str(i), date_str, typ, debt_str, credit_str, ref, desc]

            for col, val in enumerate(data):
                lbl = ctk.CTkLabel(self.tree_frame, text=val, font=self.font_farsi, anchor="center")
                lbl.grid(row=i, column=col, padx=5, pady=2, sticky="ew")

            # دکمه حذف (فقط برای انواع قابل حذف و در صورت داشتن مجوز)
            if self.can_delete_payment:
                if record_type in ('payment', 'manual_debt', 'credit'):
                    # تعیین متد حذف مناسب
                    if record_type == 'payment':
                        cmd = lambda rid=record_id: self.delete_payment(rid)
                    elif record_type == 'manual_debt':
                        cmd = lambda rid=record_id: self.delete_manual_debt(rid)
                    else:  # credit
                        cmd = lambda rid=record_id: self.delete_credit(rid)
                    
                    btn_del = ctk.CTkButton(self.tree_frame, text="🗑", font=self.font_farsi,
                                            fg_color=COLOR_ACCENT, hover_color="#d45a1c",
                                            width=40, command=cmd)
                    btn_del.grid(row=i, column=len(data), padx=5, pady=2)
                else:
                    # برای بدهی‌های ناشی از فاکتور (غیرقابل حذف)
                    ctk.CTkLabel(self.tree_frame, text="", font=self.font_farsi).grid(row=i, column=len(data), padx=5, pady=2)
            else:
                ctk.CTkLabel(self.tree_frame, text="", font=self.font_farsi).grid(row=i, column=len(data), padx=5, pady=2)

        # جمع کل
        row = len(transactions) + 1
        ctk.CTkLabel(self.tree_frame, text="جمع کل", font=self.font_farsi_bold, anchor="center").grid(row=row, column=2, padx=5, pady=5)
        ctk.CTkLabel(self.tree_frame, text=to_persian_digits(f"{total_debt:,}"), font=self.font_farsi_bold, anchor="center").grid(row=row, column=3, padx=5, pady=5)
        ctk.CTkLabel(self.tree_frame, text=to_persian_digits(f"{total_credit:,}"), font=self.font_farsi_bold, anchor="center").grid(row=row, column=4, padx=5, pady=5)

        for col in range(len(headers)):
            self.tree_frame.grid_columnconfigure(col, weight=1)

    def delete_payment(self, payment_id):
        if not self.can_pay:
            messagebox.showerror("خطا", "شما مجوز حذف پرداخت ندارید.")
            return
        if messagebox.askyesno("تأیید حذف", "آیا از حذف این پرداخت اطمینان دارید؟"):
            try:
                self.db.delete_payment(payment_id)
                self.load_transactions()
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در حذف پرداخت:\n{str(e)}")

    def delete_manual_debt(self, debt_id):
        if not self.can_pay:
            messagebox.showerror("خطا", "شما مجوز حذف بدهی دستی ندارید.")
            return
        if messagebox.askyesno("تأیید حذف", "آیا از حذف این بدهی دستی اطمینان دارید؟"):
            try:
                self.db.delete_manual_debt(debt_id)
                self.load_transactions()
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در حذف بدهی دستی:\n{str(e)}")

    def delete_credit(self, credit_id):
        if not self.can_pay:
            messagebox.showerror("خطا", "شما مجوز حذف بستانکاری ندارید.")
            return
        if messagebox.askyesno("تأیید حذف", "آیا از حذف این بستانکاری اطمینان دارید؟"):
            try:
                self.db.delete_credit(credit_id)
                self.load_transactions()
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در حذف بستانکاری:\n{str(e)}")

    def export_excel(self):
        """خروجی گرفتن از لیست تراکنش‌ها به صورت فایل اکسل"""
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        import datetime
        
        file_path = filedialog.asksaveasfilename(
            title="ذخیره فایل اکسل",
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")]
        )
        if not file_path:
            return
        
        # تابع کمکی برای تبدیل میلادی به شمسی (در صورت نیاز)
        def to_persian_date(date_val):
            if date_val is None:
                return ""
            try:
                if hasattr(date_val, 'year'):  # datetime.date
                    greg = date_val
                else:
                    # رشته مانند "2025-04-04"
                    parts = str(date_val).split('-')
                    if len(parts) == 3:
                        greg = datetime.date(int(parts[0]), int(parts[1]), int(parts[2]))
                    else:
                        return str(date_val)
                persian = jdatetime.date.fromgregorian(date=greg)
                return f"{persian.year}/{persian.month:02d}/{persian.day:02d}"
            except:
                return str(date_val)
        
        # دریافت دوباره داده‌ها از دیتابیس
        try:
            invoices = self.db.get_invoices(filters={'payer_code': self.payer_code})
            payments = self.db.get_payments(self.payer_code)
            credits = self.db.get_credits(self.payer_code)
            manual_debts = self.db.get_manual_debts(self.payer_code)
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در خواندن داده‌ها:\n{str(e)}")
            return
        
        # ساخت لیست تراکنش‌ها با تبدیل تاریخ به شمسی
        transactions = []
        for inv in invoices:
            persian_date = to_persian_date(inv[6])   # inv[6]=issue_date
            transactions.append((persian_date, 'بدهی', inv[5], inv[1], '', inv[7]))
        for pay in payments:
            persian_date = to_persian_date(pay[1])   # pay[1]=payment_date
            transactions.append((persian_date, 'پرداخت', pay[2], pay[3], pay[4], ''))
        for cr in credits:
            persian_date = to_persian_date(cr[2])    # cr[2]=credit_date
            transactions.append((persian_date, 'بستانکاری', cr[1], cr[4], cr[3], ''))
        for debt in manual_debts:
            persian_date = to_persian_date(debt[3])  # debt[3]=date
            transactions.append((persian_date, 'بدهی دستی', debt[2], debt[1], debt[4], ''))
        
        if not transactions:
            messagebox.showwarning("اخطار", "هیچ تراکنشی برای خروجی گرفتن وجود ندارد.")
            return
        
        # مرتب‌سازی بر اساس تاریخ شمسی (به عنوان رشته)
        transactions.sort(key=lambda x: x[0], reverse=True)
        
        # ایجاد کتاب کار اکسل
        wb = Workbook()
        ws = wb.active
        ws.title = "تراکنش‌ها"
        ws.sheet_view.rightToLeft = True
        
        # هدرها
        headers = ["ردیف", "تاریخ", "نوع", "بدهکار (ریال)", "بستانکار (ریال)", "شماره مرجع", "توضیحات"]
        header_font = Font(bold=True, name="B Nazanin", size=12)
        header_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        # پر کردن داده‌ها
        total_debt = 0
        total_credit = 0
        for i, (date_str, typ, amount, ref, desc, _) in enumerate(transactions, start=1):
            if typ in ('بدهی', 'بدهی دستی'):
                debt = amount
                credit = 0
                total_debt += amount
            else:   # پرداخت یا بستانکاری
                debt = 0
                credit = amount
                total_credit += amount
            
            row_data = [
                i,
                date_str,   # تاریخ شمسی
                typ,
                f"{debt:,}" if debt > 0 else "",
                f"{credit:,}" if credit > 0 else "",
                ref,
                desc
            ]
            for col, value in enumerate(row_data, start=1):
                cell = ws.cell(row=i+1, column=col, value=value)
                cell.font = Font(name="B Nazanin", size=11)
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = thin_border
        
        # اضافه کردن ردیف جمع
        total_row = len(transactions) + 2
        ws.cell(row=total_row, column=3, value="جمع کل").font = Font(bold=True, name="B Nazanin")
        ws.cell(row=total_row, column=4, value=f"{total_debt:,}").font = Font(bold=True, name="B Nazanin")
        ws.cell(row=total_row, column=5, value=f"{total_credit:,}").font = Font(bold=True, name="B Nazanin")
        
        # تنظیم عرض ستون‌ها
        column_widths = [8, 15, 15, 18, 18, 20, 30]
        for i, width in enumerate(column_widths, start=1):
            ws.column_dimensions[chr(64+i)].width = width
        
        # ذخیره فایل
        wb.save(file_path)
        messagebox.showinfo("موفقیت", f"فایل اکسل با موفقیت در مسیر زیر ذخیره شد:\n{file_path}")

    def add_manual_debt(self):
        """باز کردن دیالوگ برای افزودن بدهی دستی"""
        dialog = ctk.CTkToplevel(self)
        dialog.title("ثبت بدهی دستی")
        dialog.geometry("550x500")
        dialog.transient(self)
        dialog.grab_set()
        font_farsi = ("B Nazanin", 16)
        
        ctk.CTkLabel(dialog, text="افزودن بدهی دستی", font=("B Nazanin", 20, "bold")).pack(pady=10)
        frame = ctk.CTkFrame(dialog)
        frame.pack(pady=10, padx=20, fill="both", expand=True)
        
        # ورودی توضیحات
        ctk.CTkLabel(frame, text="توضیحات", font=font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        entry_desc = ctk.CTkEntry(frame, width=300, font=font_farsi)
        entry_desc.grid(row=0, column=1, padx=5, pady=5)
        
        # ورودی مبلغ
        ctk.CTkLabel(frame, text="مبلغ (ریال)", font=font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        entry_amount = ctk.CTkEntry(frame, width=300, font=font_farsi)
        entry_amount.grid(row=1, column=1, padx=5, pady=5)
        
        # انتخاب تاریخ
        ctk.CTkLabel(frame, text="تاریخ", font=font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        date_frame = ctk.CTkFrame(frame, fg_color="transparent")
        date_frame.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        today = jdatetime.date.today()
        year_var = ctk.StringVar(value=persian_number(str(today.year)))
        month_var = ctk.StringVar(value=persian_number(str(today.month).zfill(2)))
        day_var = ctk.StringVar(value=persian_number(str(today.day).zfill(2)))
        year_combo = ctk.CTkComboBox(date_frame, values=[persian_number(str(y)) for y in range(1400, 1431)], variable=year_var, width=90, font=font_farsi, dropdown_font=font_farsi)
        year_combo.pack(side="left", padx=2)
        month_combo = ctk.CTkComboBox(date_frame, values=[persian_number(str(m).zfill(2)) for m in range(1, 13)], variable=month_var, width=70, font=font_farsi, dropdown_font=font_farsi)
        month_combo.pack(side="left", padx=2)
        day_combo = ctk.CTkComboBox(date_frame, values=[persian_number(str(d).zfill(2)) for d in range(1, 32)], variable=day_var, width=70, font=font_farsi, dropdown_font=font_farsi)
        day_combo.pack(side="left", padx=2)
        
        # یادداشت
        ctk.CTkLabel(frame, text="یادداشت", font=font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        entry_notes = ctk.CTkEntry(frame, width=300, font=font_farsi)
        entry_notes.grid(row=3, column=1, padx=5, pady=5)
        
        # تابع ذخیره (محلی)
        def save():
            desc = entry_desc.get().strip()
            if not desc:
                messagebox.showerror("خطا", "توضیحات نمی‌تواند خالی باشد.")
                return
            try:
                amount = int(english_number(entry_amount.get()))
            except:
                messagebox.showerror("خطا", "مبلغ را به درستی وارد کنید.")
                return
            
            # تبدیل تاریخ شمسی به میلادی
            year_persian = int(english_number(year_var.get()))
            month_persian = int(english_number(month_var.get()))
            day_persian = int(english_number(day_var.get()))
            try:
                gregorian_date = jdatetime.date(year_persian, month_persian, day_persian).togregorian()
                date_str = gregorian_date.strftime("%Y-%m-%d")
            except Exception as e:
                messagebox.showerror("خطا", f"تاریخ وارد شده معتبر نیست:\n{str(e)}")
                return
            
            notes = entry_notes.get().strip()
            self.db.add_manual_debt(self.payer_code, desc, amount, date_str, notes)
            self.load_transactions()   # بازخوانی جدول تراکنش‌ها
            dialog.destroy()
        
        # دکمه‌ها
        btn_frame = ctk.CTkFrame(dialog)
        btn_frame.pack(pady=10)
        ctk.CTkButton(btn_frame, text="ذخیره", command=save, font=font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="انصراف", command=dialog.destroy, font=font_farsi, fg_color="gray", hover_color="#555555", width=100).pack(side="left", padx=5)

class AddPaymentDialog(ctk.CTkToplevel):
    def __init__(self, parent, payer_code, payer_name, on_save):
        super().__init__(parent)
        self.title("ثبت پرداخت جدید")
        self.geometry("450x450")
        self.transient(parent)
        self.grab_set()
        self.payer_code = payer_code
        self.payer_name = payer_name
        self.on_save = on_save
        self.font_farsi = ("B Nazanin", 18)

        ctk.CTkLabel(self, text=f"ثبت پرداخت برای {payer_name} (کد: {payer_code})", font=("B Nazanin", 20, "bold")).pack(pady=15)

        main_frame = ctk.CTkFrame(self)
        main_frame.pack(pady=15, padx=20, fill="both", expand=True)

        ctk.CTkLabel(main_frame, text="تاریخ پرداخت", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        date_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        date_frame.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        today = jdatetime.date.today()
        self.year_var = ctk.StringVar(value=persian_number(str(today.year)))
        self.month_var = ctk.StringVar(value=persian_number(str(today.month).zfill(2)))
        self.day_var = ctk.StringVar(value=persian_number(str(today.day).zfill(2)))
        ctk.CTkComboBox(date_frame, values=[persian_number(str(y)) for y in range(1400, 1431)], variable=self.year_var, width=90, font=self.font_farsi).pack(side="left", padx=2)
        ctk.CTkComboBox(date_frame, values=[persian_number(str(m).zfill(2)) for m in range(1, 13)], variable=self.month_var, width=70, font=self.font_farsi).pack(side="left", padx=2)
        ctk.CTkComboBox(date_frame, values=[persian_number(str(d).zfill(2)) for d in range(1, 32)], variable=self.day_var, width=70, font=self.font_farsi).pack(side="left", padx=2)

        ctk.CTkLabel(main_frame, text="مبلغ (ریال)", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.amount_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.amount_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(main_frame, text="شماره پیگیری", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.tracking_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.tracking_entry.grid(row=2, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(main_frame, text="توضیحات", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.desc_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.desc_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")

        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(side="bottom", pady=15, fill="x")
        ctk.CTkButton(btn_frame, text="💾 ذخیره", command=self.save, font=("B Nazanin", 18, "bold"), fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=140).pack(side="left", padx=5, expand=True)
        ctk.CTkButton(btn_frame, text="انصراف", command=self.destroy, font=("B Nazanin", 18, "bold"), fg_color="gray", hover_color="#555555", width=140).pack(side="left", padx=5, expand=True)

    def save(self):
        year = self.year_var.get()
        month = self.month_var.get()
        day = self.day_var.get()
        date = f"{year}/{month}/{day}"
        try:
            amount = int(english_number(self.amount_entry.get()))
        except:
            messagebox.showerror("خطا", "مبلغ را به درستی وارد کنید.")
            return
        tracking = self.tracking_entry.get().strip()
        desc = self.desc_entry.get().strip()
        if not tracking:
            messagebox.showerror("خطا", "شماره پیگیری نمی‌تواند خالی باشد.")
            return
        self.on_save(self.payer_code, self.payer_name, date, amount, traload_transactionscking, desc)
        self.destroy()

class RejectReasonDialog(ctk.CTkToplevel):
    def __init__(self, parent, font_farsi):
        super().__init__(parent)
        self.title("رد درخواست")
        self.geometry("400x250")
        self.transient(parent)
        self.grab_set()
        self.font_farsi = font_farsi
        self.result = None
        
        ctk.CTkLabel(self, text="دلیل رد را وارد کنید:", font=self.font_farsi).pack(pady=20)
        self.text_entry = ctk.CTkTextbox(self, height=100, font=self.font_farsi, wrap="word")
        self.text_entry.pack(pady=10, padx=20, fill="both", expand=True)
        
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(pady=20)
        ctk.CTkButton(btn_frame, text="تأیید", command=self.confirm, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100).pack(side="left", padx=10)
        ctk.CTkButton(btn_frame, text="انصراف", command=self.destroy, font=self.font_farsi, fg_color="gray", hover_color="#555555", width=100).pack(side="left", padx=10)
    
    def confirm(self):
        self.result = self.text_entry.get("0.0", "end").strip()
        if not self.result:
            messagebox.showerror("خطا", "لطفاً دلیل رد را وارد کنید.")
            return
        self.destroy()

class ManagementDialog(ctk.CTkToplevel):
    def __init__(self, parent, app, payers_list, can_view, can_pay, can_issue, can_delete_history, default_tab="تاریخچه صورتحساب‌ها"):
        super().__init__(parent)
        self.app = app
        self.title("مدیریت صورتحساب‌ها و حساب شرکت‌ها")
        self.geometry("1300x700")
        self.transient(parent)
        self.grab_set()
        self.db = app.db
        self.payers_list = payers_list
        self.can_view = can_view
        self.can_pay = can_pay
        self.can_issue = can_issue
        self.can_delete_history = can_delete_history
        self.font_farsi = ("B Nazanin", 16)
        self.font_farsi_bold = ("B Nazanin", 16, "bold")
        self.tabview = ctk.CTkTabview(self, segmented_button_fg_color=COLOR_PRIMARY)
        self.tabview._segmented_button.configure(font=self.font_farsi)
        self.tabview.pack(fill="both", expand=True, padx=15, pady=15)
        self.tab_history = self.tabview.add("تاریخچه صورتحساب‌ها")
        self.create_history_tab()
        self.tab_performance = self.tabview.add("عملکرد شرکت‌ها")
        self.create_performance_tab()
        self.tab_workflow = self.tabview.add("گردش کار")
        self.create_workflow_tab()
        self.tabview.set(default_tab)
    
    def create_history_tab(self):
        self.filter_payer = ctk.StringVar()
        self.filter_start_date = ctk.StringVar(value="")
        self.filter_end_date = ctk.StringVar(value="")
        main_frame = ctk.CTkFrame(self.tab_history)
        main_frame.pack(fill="both", expand=True, padx=5, pady=5)
        ctk.CTkLabel(main_frame, text="لیست صورتحساب‌های صادر شده", font=("B Nazanin", 18, "bold")).pack(pady=5)
        headers = ["ردیف", "شماره نامه", "طرف حساب", "کد", "نوع", "مبلغ (ریال)", "تاریخ صدور", "دانلود"]
        if self.can_delete_history:
            headers.append("حذف")
        filter_frame = ctk.CTkFrame(main_frame)
        filter_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(filter_frame, text="طرف حساب", font=self.font_farsi).pack(side="right", padx=5)
        payer_entry = ctk.CTkEntry(filter_frame, textvariable=self.filter_payer, width=180, font=self.font_farsi)
        payer_entry.pack(side="right", padx=5)
        ctk.CTkLabel(filter_frame, text="از تاریخ", font=self.font_farsi).pack(side="right", padx=(10,2))
        self.start_date_label = ctk.CTkLabel(filter_frame, textvariable=self.filter_start_date, width=100, font=self.font_farsi, fg_color="#333333", corner_radius=5)
        self.start_date_label.pack(side="right", padx=2)
        ctk.CTkButton(filter_frame, text="📅", command=lambda: self.select_date("start"), width=40, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT).pack(side="right", padx=2)
        ctk.CTkLabel(filter_frame, text="تا تاریخ", font=self.font_farsi).pack(side="right", padx=(10,2))
        self.end_date_label = ctk.CTkLabel(filter_frame, textvariable=self.filter_end_date, width=100, font=self.font_farsi, fg_color="#333333", corner_radius=5)
        self.end_date_label.pack(side="right", padx=2)
        ctk.CTkButton(filter_frame, text="📅", command=lambda: self.select_date("end"), width=40, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT).pack(side="right", padx=2)
        ctk.CTkButton(filter_frame, text="اعمال فیلتر", command=self.apply_history_filter, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="right", padx=10)
        ctk.CTkButton(filter_frame, text="پاک کردن", command=self.clear_history_filter, font=self.font_farsi, fg_color="gray", hover_color="#555555", width=100).pack(side="right", padx=5)
        #ctk.CTkButton(filter_frame, text="حذف همه", command=self.delete_all_history, font=self.font_farsi, fg_color=COLOR_ACCENT, hover_color="#d45a1c", width=120).pack(side="left", padx=5)
        self.history_table = ctk.CTkScrollableFrame(main_frame)
        self.history_table.pack(fill="both", expand=True, padx=5, pady=5)
        self.load_history()

    def select_date(self, target):
        initial = self.filter_start_date.get() if target=="start" else self.filter_end_date.get()
        dialog = DatePickerDialog(self, initial if initial else None)
        self.wait_window(dialog)
        if dialog.result:
            if target == "start":
                self.filter_start_date.set(dialog.result)
            else:
                self.filter_end_date.set(dialog.result)
    def apply_history_filter(self):
        self.load_history()
    def clear_history_filter(self):
        self.filter_payer.set("")
        self.filter_start_date.set("")
        self.filter_end_date.set("")
        self.load_history()
    def delete_all_history(self):
        if messagebox.askyesno("تأیید حذف", "آیا از حذف تمام صورتحساب‌ها اطمینان دارید؟"):
            try:
                self.db.delete_all_invoices()
                self.load_history()
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در حذف همه: {str(e)}")
    
    def load_history(self):
        for widget in self.history_table.winfo_children():
            widget.destroy()
        filters = {}
        if self.filter_payer.get():
            filters['payer'] = self.filter_payer.get()
        if self.filter_start_date.get():
            filters['start_date'] = self.filter_start_date.get()
        if self.filter_end_date.get():
            filters['end_date'] = self.filter_end_date.get()
        rows = self.db.get_invoices(filters)
        if not rows:
            ctk.CTkLabel(self.history_table, text="هیچ صورتحسابی با این فیلترها یافت نشد.", font=self.font_farsi, text_color="gray").pack(pady=20)
            return

        # دریافت map شماره نامه به file_id (برای دانلود)
        file_id_map = {}
        try:
            with self.db.connect() as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT letter_number, id FROM final_files")
                for ln, fid in cursor.fetchall():
                    file_id_map[ln] = fid
        except Exception:
            pass

        # تعریف هدرها با توجه به مجوز حذف
        headers = ["ردیف", "شماره نامه", "طرف حساب", "کد", "نوع", "مبلغ (ریال)", "تاریخ صدور", "دانلود"]
        if self.can_delete_history:
            headers.append("حذف")
        
        for col, header in enumerate(headers):
            lbl = ctk.CTkLabel(self.history_table, text=header, font=self.font_farsi, anchor="center")
            lbl.grid(row=0, column=col, padx=5, pady=5, sticky="ew")

        for i, row in enumerate(rows, start=1):
            inv_id, letter_no, payer, payer_code, inv_type, amount_val, issue_date, file_path = row
            amount_str = f"{amount_val:,}"
            data = [str(i), letter_no, payer, payer_code or "", inv_type, amount_str, issue_date]
            
            # ستون‌های اطلاعاتی
            for col, value in enumerate(data):
                lbl = ctk.CTkLabel(self.history_table, text=value, font=self.font_farsi, anchor="center")
                lbl.grid(row=i, column=col, padx=5, pady=2, sticky="ew")
            
            # ستون دانلود
            btn_download = ctk.CTkButton(self.history_table, text="⬇ دانلود", font=self.font_farsi,
                            width=80, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT,
                            command=lambda ln=letter_no: self.app.regenerate_and_download(ln))
            btn_download.grid(row=i, column=len(data), padx=2, pady=2)
            
            # ستون حذف (در صورت داشتن مجوز)
            if self.can_delete_history:
                btn_delete = ctk.CTkButton(self.history_table, text="🗑 حذف", font=self.font_farsi,
                                width=80, fg_color=COLOR_ACCENT, hover_color="#d45a1c",
                                command=lambda inv_id=inv_id, ln=letter_no, pc=payer_code, idate=issue_date: self.delete_invoice(inv_id, ln, pc, idate))
                btn_delete.grid(row=i, column=len(data)+1, padx=2, pady=2)

        for col in range(len(headers)):
            self.history_table.grid_columnconfigure(col, weight=1)

    def delete_record(self, record_id):
        if messagebox.askyesno("تأیید حذف", "آیا از حذف این آیتم اطمینان دارید؟"):
            try:
                self.db.delete_invoice(record_id)
                self.load_history()
                self.load_performance()
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در حذف: {str(e)}")
    def open_file(self, file_path):
        try:
            os.startfile(file_path)
        except Exception as e:
            messagebox.showerror("خطا", f"فایل یافت نشد یا قابل باز کردن نیست:\n{str(e)}")
    def create_performance_tab(self):
        main_frame = ctk.CTkFrame(self.tab_performance)
        main_frame.pack(fill="both", expand=True, padx=5, pady=5)
        ctk.CTkLabel(main_frame, text="لیست شرکت‌ها و مانده حساب", font=("B Nazanin", 18, "bold")).pack(pady=5)
        self.performance_table = ctk.CTkScrollableFrame(main_frame)
        self.performance_table.pack(fill="both", expand=True, padx=5, pady=5)
        self.load_performance()
    def load_performance(self):
        for widget in self.performance_table.winfo_children():
            widget.destroy()
        invoices = self.db.get_invoices()
        invoice_totals = {}
        for inv in invoices:
            code = inv[3]
            amount = inv[5]
            invoice_totals[code] = invoice_totals.get(code, 0) + amount
        credits = self.db.get_credits()
        credit_totals = {}
        for cr in credits:
            if len(cr) == 6:
                code = cr[5]
                amount = cr[1]
                credit_totals[code] = credit_totals.get(code, 0) + amount
            if self.can_delete_payment:
                btn_del = ctk.CTkButton(self.tree_frame, text="🗑", font=self.font_farsi, fg_color=COLOR_ACCENT, hover_color="#d45a1c", width=40, command=lambda cid=credit_id: self.delete_credit(cid))
        manual_debts = self.db.get_manual_debts()
        manual_totals = {}
        for debt in manual_debts:
            if len(debt) == 6:
                code = debt[5]
                amount = debt[2]
                manual_totals[code] = manual_totals.get(code, 0) + amount
        payments = self.db.get_payments()
        payment_totals = {}
        for pay in payments:
            payer_code = pay[5]
            amount = pay[2]
            payment_totals[payer_code] = payment_totals.get(payer_code, 0) + amount
        all_codes = set(invoice_totals.keys()) | set(payment_totals.keys()) | set(manual_totals.keys()) | set(credit_totals.keys())
        for p in self.payers_list:
            all_codes.add(p['code'])
        name_by_code = {p['code']: p['name'] for p in self.payers_list}
        headers = ["ردیف", "کد", "نام شرکت", "کل بدهی (ریال)", "کل پرداختی (ریال)", "مانده (ریال)", "عملیات"]
        col_widths = [5, 8, 25, 18, 18, 18, 15]
        for col, (header, width) in enumerate(zip(headers, col_widths)):
            lbl = ctk.CTkLabel(self.performance_table, text=header, font=self.font_farsi, anchor="center", width=width*10)
            lbl.grid(row=0, column=col, padx=5, pady=5, sticky="ew")
        companies_data = []
        for code in sorted(all_codes):
            name = name_by_code.get(code, code)
            total_debt = invoice_totals.get(code, 0) + manual_totals.get(code, 0) - credit_totals.get(code, 0)
            total_paid = payment_totals.get(code, 0)
            balance = total_debt - total_paid
            companies_data.append((code, name, total_debt, total_paid, balance))
        for i, (code, name, debt, paid, balance) in enumerate(companies_data, start=1):
            balance_display = to_persian_digits(f"{balance:,}") if balance >= 0 else f"({to_persian_digits(str(abs(balance)))} بستانکار)"
            data = [str(i), code, name, to_persian_digits(f"{debt:,}") if debt != 0 else "۰", to_persian_digits(f"{paid:,}") if paid != 0 else "۰", balance_display]
            for col, val in enumerate(data):
                lbl = ctk.CTkLabel(self.performance_table, text=val, font=self.font_farsi, anchor="center", width=col_widths[col]*10)
                lbl.grid(row=i, column=col, padx=5, pady=2, sticky="ew")
            if self.can_pay:
                btn_pay = ctk.CTkButton(self.performance_table, text="💳 ثبت پرداخت", font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, command=lambda c=code, n=name: self.add_payment(c, n))
                btn_pay.grid(row=i, column=len(data), padx=5, pady=2)
            else:
                ctk.CTkLabel(self.performance_table, text="", font=self.font_farsi).grid(row=i, column=len(data), padx=5, pady=2)
            btn_detail = ctk.CTkButton(self.performance_table, text="📋 جزئیات", font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, command=lambda c=code, n=name: self.show_company_details(c, n))
            btn_detail.grid(row=i, column=len(data)+1, padx=5, pady=2)
        for col in range(len(headers)+2):
            self.performance_table.grid_columnconfigure(col, weight=1)
    def add_payment(self, payer_code, payer_name):
        if not self.can_pay:
            messagebox.showerror("خطا", "شما مجوز ثبت پرداخت ندارید.")
            return
        dialog = AddPaymentDialog(self, payer_code, payer_name, self.save_payment)
        self.wait_window(dialog)
    def save_payment(self, payer_code, payer_name, date, amount, tracking, desc):
        try:
            year = date.split('/')[0]
            payments = self.db.get_payments(payer_code)
            count = len(payments) + 1
            payment_code = f"PAY-{year}-{count:04d}"
            self.db.add_payment(payment_code, payer_code, payer_name, date, amount, tracking, desc)
            messagebox.showinfo("موفقیت", f"پرداخت با کد {payment_code} ثبت شد.")
            self.load_performance()
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در ثبت پرداخت: {str(e)}")
    def show_company_details(self, payer_code, payer_name):
        dialog = CompanyTransactionsDialog(self, self.app, payer_code, payer_name, self.can_pay, self.app.can_delete_payment)
        self.wait_window(dialog)
    def create_workflow_tab(self):
        main_frame = ctk.CTkFrame(self.tab_workflow, fg_color="#2b2b2b")
        main_frame.pack(fill="both", expand=True, padx=5, pady=5)
        ctk.CTkLabel(main_frame, text="گردش کار درخواست‌های تأیید", font=("B Nazanin", 18, "bold")).pack(pady=5)
        filter_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        filter_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(filter_frame, text="فیلتر وضعیت:", font=self.font_farsi).pack(side="right", padx=5)
        self.status_filter = ctk.StringVar(value="همه")
        status_combo = ctk.CTkComboBox(filter_frame, values=["همه", "pending", "approved", "rejected"], variable=self.status_filter, font=self.font_farsi, width=120)
        status_combo.pack(side="right", padx=5)
        ctk.CTkButton(filter_frame, text="اعمال فیلتر", command=self.load_workflow_requests, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100).pack(side="right", padx=5)
        ctk.CTkButton(filter_frame, text="بروزرسانی", command=self.load_workflow_requests, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=100).pack(side="right", padx=5)
        self.workflow_tree_frame = ctk.CTkScrollableFrame(main_frame, fg_color="#333333")
        self.workflow_tree_frame.pack(fill="both", expand=True, padx=5, pady=5)
        self.load_workflow_requests()
    def load_workflow_requests(self):
        for widget in self.workflow_tree_frame.winfo_children():
            widget.destroy()
        if self.app.can_approve:
            requests = self.db.get_approval_requests()
        else:
            requests = self.db.get_approval_requests(user_id=self.app.user_id)
        if not requests:
            ctk.CTkLabel(self.workflow_tree_frame, text="هیچ درخواستی یافت نشد.", font=self.font_farsi, text_color="gray").pack(pady=20)
            return
        status = self.status_filter.get()
        if status != "همه":
            requests = [r for r in requests if r[5] == status]
        # هدرها - اضافه کردن ستون حذف
        headers = ["شناسه", "متقاضی", "طرف حساب", "شماره نامه", "تاریخ نامه", "وضعیت", "تاریخ درخواست", "دلیل رد", "عملیات", "صدور فایل نهایی", "حذف"]
        for col, header in enumerate(headers):
            lbl = ctk.CTkLabel(self.workflow_tree_frame, text=header, font=self.font_farsi_bold, anchor="center")
            lbl.grid(row=0, column=col, padx=5, pady=5, sticky="ew")
        for i, req in enumerate(requests, start=1):
            req_id, req_owner_id, req_name, payer, letter_no, letter_date, stat, created, reason = req
            stat_text = {"pending": "در انتظار", "approved": "تأیید شده", "rejected": "رد شده"}.get(stat, stat)
            reason_display = reason if reason else ""
            data = [str(i), req_name, payer, letter_no, letter_date, stat_text, created, reason_display]
            for col, val in enumerate(data):
                lbl = ctk.CTkLabel(self.workflow_tree_frame, text=val, font=self.font_farsi, anchor="center")
                lbl.grid(row=i, column=col, padx=5, pady=2, sticky="ew")
            col_offset = len(data)  # 8
            # دکمه‌های تأیید/رد/حذف فقط برای کاربران دارای مجوز تأیید
            if self.app.can_approve or (self.app.user_id == req_owner_id and self.app.can_delete_history):
                if stat == "pending":
                    btn_approve = ctk.CTkButton(self.workflow_tree_frame, text="✅ تأیید", font=self.font_farsi, fg_color="green", hover_color="#2E7D32", width=70, command=lambda rid=req_id: self.approve_workflow_request(rid))
                    btn_approve.grid(row=i, column=col_offset, padx=5, pady=2)
                    btn_reject = ctk.CTkButton(self.workflow_tree_frame, text="رد", font=self.font_farsi, fg_color=COLOR_ACCENT, hover_color="#d45a1c", width=70, command=lambda rid=req_id: self.reject_workflow_request(rid))
                    btn_reject.grid(row=i, column=col_offset+1, padx=5, pady=2)
                else:
                    # برای درخواست‌های تأیید یا رد شده، جای خالی نگه دار
                    ctk.CTkLabel(self.workflow_tree_frame, text="").grid(row=i, column=col_offset, padx=5, pady=2)
                    ctk.CTkLabel(self.workflow_tree_frame, text="").grid(row=i, column=col_offset+1, padx=5, pady=2)
                # دکمه حذف برای همه درخواست‌ها (برای مدیران تأیید)
                btn_delete = ctk.CTkButton(self.workflow_tree_frame, text="🗑 حذف", font=self.font_farsi, fg_color=COLOR_ACCENT, hover_color="#d45a1c", width=60, command=lambda rid=req_id: self.delete_workflow_request(rid))
                btn_delete.grid(row=i, column=col_offset+2, padx=5, pady=2)
            else:
                # کاربران عادی (بدون مجوز تأیید) جای خالی می‌گذارند
                ctk.CTkLabel(self.workflow_tree_frame, text="").grid(row=i, column=col_offset, padx=5, pady=2)
                ctk.CTkLabel(self.workflow_tree_frame, text="").grid(row=i, column=col_offset+1, padx=5, pady=2)
                ctk.CTkLabel(self.workflow_tree_frame, text="").grid(row=i, column=col_offset+2, padx=5, pady=2)
            # دکمه جزئیات برای همه
            btn_details = ctk.CTkButton(self.workflow_tree_frame, text="جزئیات", font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=70, command=lambda rid=req_id: self.show_workflow_details(rid))
            btn_details.grid(row=i, column=col_offset+3, padx=5, pady=2)
            # دکمه صدور فایل نهایی فقط برای درخواست‌های تأیید شده
            if stat == "approved":
                if self.app.user_id == req_owner_id or self.app.can_approve:
                    btn_final = ctk.CTkButton(self.workflow_tree_frame, text="صدور فایل نهایی", font=self.font_farsi, fg_color=COLOR_COMPANY, hover_color="#388E3C", width=130, command=lambda rid=req_id: self.export_final_files_by_user(rid))
                    btn_final.grid(row=i, column=col_offset+4, padx=5, pady=2)
                else:
                    ctk.CTkLabel(self.workflow_tree_frame, text="").grid(row=i, column=col_offset+4, padx=5, pady=2)
            else:
                ctk.CTkLabel(self.workflow_tree_frame, text="").grid(row=i, column=col_offset+4, padx=5, pady=2)
            # دکمه حذف (برای مدیران تأیید) قبلاً اضافه شد
        for col in range(len(headers)):
            self.workflow_tree_frame.grid_columnconfigure(col, weight=1)
    
    def show_workflow_details(self, req_id):
        data = self.db.get_approval_request_by_id(req_id)
        if not data:
            messagebox.showerror("خطا", "درخواست یافت نشد.")
            return

        detail_dialog = ctk.CTkToplevel(self)
        detail_dialog.title(f"جزئیات درخواست #{req_id} - کارتابل گردش کار")
        detail_dialog.geometry("1100x800")
        detail_dialog.transient(self)
        detail_dialog.grab_set()

        notebook = ctk.CTkTabview(detail_dialog)
        notebook._segmented_button.configure(font=self.font_farsi)  # تنظیم فونت تب‌ها
        notebook.pack(fill="both", expand=True, padx=10, pady=10)

        # ---------- تب خلاصه ----------
        summary_tab = notebook.add("خلاصه")
        text_summary = ctk.CTkTextbox(summary_tab, font=self.font_farsi, wrap="word")
        text_summary.pack(fill="both", expand=True)

        # محاسبه جمع مبالغ
        import json
        services = json.loads(data['service_data_json'])
        flights = json.loads(data['flight_data_json'])
        hotels = json.loads(data['hotel_data_json'])

        total_service = sum(item.get('balance', 0) for item in services)
        total_flight = sum(item.get('balance', 0) for item in flights)
        total_hotel = sum(item.get('balance', 0) for item in hotels)
        total_all = total_service + total_flight + total_hotel

        summary = f"""
    شناسه درخواست: {data['id']}
    متقاضی: {data['requester_name']}
    طرف حساب: {data['payer_name']} (کد: {data['payer_code']})
    شماره نامه موقت: {data['letter_number']}
    تاریخ نامه: {data['letter_date']}
    بازه صورتحساب: {data['period_range']}
    پوشه خروجی موقت: {data['output_dir']}
    تاریخ ارسال: {data['created_at']}
    وضعیت: {data['status']}
    دلیل رد: {data.get('rejection_reason', '')}

    --- خلاصه هزینه‌ها ---
    خدمات دیگر : {to_persian_digits(f'{total_service:,}')} ریال
    پروازها    : {to_persian_digits(f'{total_flight:,}')} ریال
    هتل‌ها     : {to_persian_digits(f'{total_hotel:,}')} ریال
    جمع کل     : {to_persian_digits(f'{total_all:,}')} ریال
    """
        text_summary.insert("0.0", summary)
        text_summary.configure(state="disabled")

        # ---------- تب خدمات ----------
        service_tab = notebook.add("خدمات")
        self._create_detail_table(service_tab, services, "خدمات")

        # ---------- تب پروازها ----------
        flight_tab = notebook.add("پروازها")
        self._create_detail_table(flight_tab, flights, "پرواز")

        # ---------- تب هتل‌ها ----------
        hotel_tab = notebook.add("هتل‌ها")
        self._create_detail_table(hotel_tab, hotels, "هتل")

        btn_close = ctk.CTkButton(detail_dialog, text="بستن", command=detail_dialog.destroy,
                              font=self.font_farsi, fg_color=COLOR_PRIMARY)
        btn_close.pack(pady=10)

    def _create_detail_table(self, parent, data_list, title):
        if not data_list:
            ctk.CTkLabel(parent, text=f"هیچ داده‌ای برای {title} وجود ندارد.", font=self.font_farsi, text_color="gray").pack(pady=20)
            return
        if title == "خدمات":
            columns = ("نوع خدمت", "نام مسافر", "شماره قرارداد", "تاریخ", "توضیحات", "بدهکار", "بستانکار", "مانده")
            col_widths = (120, 150, 100, 100, 150, 120, 120, 120)
        elif title == "پرواز":
            columns = ("شماره قرارداد", "نام مسافر", "مسیر", "شماره بلیط", "تاریخ پرواز", "توضیحات", "بدهکار", "بستانکار", "مانده")
            col_widths = (100, 150, 150, 120, 100, 150, 120, 120, 120)
        else:
            columns = ("شماره قرارداد", "نام مسافر", "هتل", "نوع اتاق", "تعداد نفرات", "تاریخ ورود/خروج", "توضیحات", "بدهکار", "بستانکار", "مانده")
            col_widths = (100, 150, 150, 120, 90, 150, 150, 120, 120, 120)
        tree_frame = ctk.CTkFrame(parent, fg_color="transparent")
        tree_frame.pack(fill="both", expand=True, padx=10, pady=10)
        style = ttk.Style()
        style.configure("Custom.Treeview", font=("B Nazanin", 12), rowheight=30)
        style.configure("Custom.Treeview.Heading", font=("B Nazanin", 13, "bold"))
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15, style="Custom.Treeview")
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=vsb.set)
        for col, width in zip(columns, col_widths):
            tree.heading(col, text=col)
            tree.column(col, width=width, anchor="center")
        tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        for idx, item in enumerate(data_list, start=1):
            if title == "خدمات":
                values = (item.get('type', ''), item.get('passenger', ''), item.get('contract', ''), item.get('date', ''), item.get('notes', ''), to_persian_digits(f"{item.get('debt', 0):,}"), to_persian_digits(f"{item.get('credit', 0):,}"), to_persian_digits(f"{item.get('balance', 0):,}"))
            elif title == "پرواز":
                values = (item.get('contract', ''), item.get('passenger', ''), item.get('route', ''), item.get('ticket', ''), item.get('date', ''), item.get('notes', ''), to_persian_digits(f"{item.get('debt', 0):,}"), to_persian_digits(f"{item.get('credit', 0):,}"), to_persian_digits(f"{item.get('balance', 0):,}"))
            else:
                values = (item.get('contract', ''), item.get('passenger', ''), item.get('hotel', ''), item.get('room', ''), str(item.get('pax', 1)), item.get('date', ''), item.get('notes', ''), to_persian_digits(f"{item.get('debt', 0):,}"), to_persian_digits(f"{item.get('credit', 0):,}"), to_persian_digits(f"{item.get('balance', 0):,}"))
            tree.insert("", "end", values=values)
    def approve_workflow_request(self, req_id):
        if messagebox.askyesno("تأیید", "آیا از تأیید این درخواست اطمینان دارید؟"):
            try:
                success, result = self.app.call_api_generate(req_id)
                if success:
                    self.db.approve_request(req_id, self.app.user_id)
                    messagebox.showinfo("موفقیت", "درخواست تأیید شد و فایل روی سرور ساخته شد.")
                else:
                    messagebox.showerror("خطا", f"خطا در ساخت فایل: {result}")
                self.load_workflow_requests()
            except Exception as e:
                messagebox.showerror("خطا", str(e))

    def reject_workflow_request(self, req_id):
        dialog = RejectReasonDialog(self, self.font_farsi)
        self.wait_window(dialog)
        reason = dialog.result
        if not reason:
            return
        try:
            # گرفتن شماره نامه موقت (که قبلاً رزرو شده) از دیتابیس
            req = self.db.get_approval_request_by_id(req_id)
            if not req:
                messagebox.showerror("خطا", "درخواست یافت نشد")
                return
            letter_number = req.get('letter_number')  # شماره موقت (رزرو شده)
            
            # آزادسازی شماره در سرور
            import requests
            api_url = f"http://{self.app.api_server}:5000/release-letter"
            requests.post(api_url, json={'letter_number': letter_number}, timeout=5)
            
            # رد درخواست در دیتابیس محلی
            self.db.reject_request(req_id, self.app.user_id, reason)
            messagebox.showinfo("موفقیت", "درخواست رد شد و شماره نامه آزاد گردید.")
            self.load_workflow_requests()
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در رد درخواست: {str(e)}")

    def delete_workflow_request(self, req_id):
        if messagebox.askyesno("تأیید", "آیا از حذف این درخواست اطمینان دارید؟"):
            try:
                self.db.delete_approval_request(req_id)
                messagebox.showinfo("موفقیت", "درخواست حذف شد.")
                self.load_workflow_requests()
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در حذف: {str(e)}")

    def export_final_files_by_user(self, req_id):
        data = self.db.get_approval_request_by_id(req_id)
        if not data:
            messagebox.showerror("خطا", "درخواست یافت نشد.")
            return
        if self.app.user_id != data['requester_id'] and not self.app.can_approve:
            messagebox.showerror("خطا", "شما مجوز صدور فایل نهایی این درخواست را ندارید.")
            return
        output_dir = filedialog.askdirectory(title="انتخاب پوشه برای ذخیره فایل‌های نهایی")
        if not output_dir:
            return
        try:
            final_files = self.app.generate_from_approval(data, output_dir, self.app.user_id)
            messagebox.showinfo("موفقیت", f"فایل‌های نهایی با موفقیت در مسیر زیر ذخیره شدند:\n{output_dir}")
            self.load_workflow_requests()
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در تولید فایل‌های نهایی:\n{str(e)}")

    def delete_invoice(self, invoice_id, letter_number, payer_code, issue_date):
        if not self.can_delete_history:
            messagebox.showerror("خطا", "شما مجوز حذف تاریخچه را ندارید.")
            return
        if messagebox.askyesno("تأیید حذف", f"آیا از حذف صورتحساب {letter_number} اطمینان دارید؟"):
            try:
                self.db.delete_invoice(invoice_id)   # این متد شماره نامه را آزاد می‌کند
                self.load_history()                 # بروزرسانی جدول تاریخچه
                self.load_performance()             # بروزرسانی جدول عملکرد شرکت‌ها
                messagebox.showinfo("موفقیت", "صورتحساب با موفقیت حذف شد.")
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در حذف: {str(e)}")

class DatePickerDialog(ctk.CTkToplevel):
    def __init__(self, parent, initial_date=None):
        super().__init__(parent)
        self.title("انتخاب تاریخ")
        self.geometry("350x300")
        self.transient(parent)
        self.grab_set()
        self.result = None
        self.font_farsi = ("B Nazanin", 18)
        today = jdatetime.date.today()
        if initial_date:
            try:
                parts = initial_date.split('/')
                self.year = int(parts[0])
                self.month = int(parts[1])
                self.day = int(parts[2])
            except:
                self.year = today.year
                self.month = today.month
                self.day = today.day
        else:
            self.year = today.year
            self.month = today.month
            self.day = today.day
        year_frame = ctk.CTkFrame(self)
        year_frame.pack(pady=5, padx=10, fill="x")
        ctk.CTkLabel(year_frame, text="سال", font=self.font_farsi).pack(side="right", padx=5)
        self.year_var = ctk.StringVar(value=str(self.year))
        year_combo = ctk.CTkComboBox(year_frame, values=[str(y) for y in range(1400, 1431)], variable=self.year_var, width=100, font=self.font_farsi, dropdown_font=self.font_farsi)
        year_combo.pack(side="right", padx=5)
        month_frame = ctk.CTkFrame(self)
        month_frame.pack(pady=5, padx=10, fill="x")
        ctk.CTkLabel(month_frame, text="ماه", font=self.font_farsi).pack(side="right", padx=5)
        month_values = [str(m).zfill(2) for m in range(1, 13)]
        self.month_var = ctk.StringVar(value=str(self.month).zfill(2))
        month_combo = ctk.CTkComboBox(month_frame, values=month_values, variable=self.month_var, width=80, font=self.font_farsi, dropdown_font=self.font_farsi)
        month_combo.pack(side="right", padx=5)
        day_frame = ctk.CTkFrame(self)
        day_frame.pack(pady=5, padx=10, fill="x")
        ctk.CTkLabel(day_frame, text="روز", font=self.font_farsi).pack(side="right", padx=5)
        day_values = [str(d).zfill(2) for d in range(1, 32)]
        self.day_var = ctk.StringVar(value=str(self.day).zfill(2))
        day_combo = ctk.CTkComboBox(day_frame, values=day_values, variable=self.day_var, width=80, font=self.font_farsi, dropdown_font=self.font_farsi)
        day_combo.pack(side="right", padx=5)
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(pady=15)
        ctk.CTkButton(btn_frame, text="تأیید", command=self.confirm, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT).pack()
    def confirm(self):
        year = self.year_var.get()
        month = self.month_var.get()
        day = self.day_var.get()
        self.result = f"{year}/{month}/{day}"
        self.destroy()

class DateRangePickerDialog(ctk.CTkToplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("انتخاب بازه صورتحساب")
        self.geometry("600x650")
        self.transient(parent)
        self.grab_set()
        self.result = None
        self.font_farsi = ("B Nazanin", 16)
        self.today = jdatetime.date.today()
        self.current_year = self.today.year
        self.current_month = self.today.month
        self.current_month_name = MONTH_LIST[self.current_month - 1]
        self.start_date = None
        self.end_date = None
        self.selection_mode = "start"
        top_frame = ctk.CTkFrame(self)
        top_frame.pack(pady=10, padx=10, fill="x")
        ctk.CTkLabel(top_frame, text="سال", font=self.font_farsi).pack(side="right", padx=5)
        self.year_var = ctk.StringVar(value=str(self.current_year))
        self.year_combo = ctk.CTkComboBox(top_frame, values=[str(y) for y in range(1400, 1431)], variable=self.year_var, width=100, font=self.font_farsi, dropdown_font=self.font_farsi, command=self.on_year_month_change)
        self.year_combo.pack(side="right", padx=5)
        ctk.CTkLabel(top_frame, text="ماه", font=self.font_farsi).pack(side="right", padx=5)
        self.month_var = ctk.StringVar(value=self.current_month_name)
        self.month_combo = ctk.CTkComboBox(top_frame, values=MONTH_LIST, variable=self.month_var, width=120, font=self.font_farsi, dropdown_font=self.font_farsi, command=self.on_year_month_change)
        self.month_combo.pack(side="right", padx=5)
        self.today_btn = ctk.CTkButton(top_frame, text="برو به امروز", command=self.goto_today, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT)
        self.today_btn.pack(side="left", padx=5)
        days_frame = ctk.CTkFrame(self)
        days_frame.pack(pady=10, padx=10, fill="x")
        weekdays = ["ش", "ی", "د", "س", "چ", "پ", "ج"]
        for i, day in enumerate(weekdays):
            lbl = ctk.CTkLabel(days_frame, text=day, font=self.font_farsi, width=70, anchor="center")
            lbl.grid(row=0, column=i, padx=2, pady=2)
        self.calendar_frame = ctk.CTkFrame(self)
        self.calendar_frame.pack(pady=10, padx=10, fill="both", expand=True)
        self.selection_label = ctk.CTkLabel(self, text="شروع: انتخاب نشده | پایان: انتخاب نشده", font=self.font_farsi, text_color="gray")
        self.selection_label.pack(pady=5)
        mode_frame = ctk.CTkFrame(self)
        mode_frame.pack(pady=5, padx=10, fill="x")
        self.start_mode_btn = ctk.CTkButton(mode_frame, text="انتخاب شروع", command=self.set_start_mode, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=140)
        self.start_mode_btn.pack(side="right", padx=5)
        self.end_mode_btn = ctk.CTkButton(mode_frame, text="انتخاب پایان", command=self.set_end_mode, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=140)
        self.end_mode_btn.pack(side="right", padx=5)
        self.confirm_btn = ctk.CTkButton(self, text="تأیید بازه", command=self.confirm, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=140, state="disabled")
        self.confirm_btn.pack(pady=10)
        self.update_calendar()
    def on_year_month_change(self, choice=None):
        try:
            year = int(self.year_var.get())
            month_name = self.month_var.get()
            month = MONTH_LIST.index(month_name) + 1
            self.current_year = year
            self.current_month = month
            self.update_calendar()
        except:
            pass
    def goto_today(self):
        self.year_var.set(str(self.today.year))
        self.month_var.set(MONTH_LIST[self.today.month - 1])
        self.current_year = self.today.year
        self.current_month = self.today.month
        self.update_calendar()
    def get_month_days(self, year, month):
        if month <= 6:
            return 31
        elif month <= 11:
            return 30
        else:
            return 29
    def update_calendar(self):
        for widget in self.calendar_frame.winfo_children():
            widget.destroy()
        try:
            first_day = jdatetime.date(self.current_year, self.current_month, 1)
            start_weekday = first_day.weekday()
            days_in_month = self.get_month_days(self.current_year, self.current_month)
        except:
            return
        row = 0
        for i in range(start_weekday):
            lbl = ctk.CTkLabel(self.calendar_frame, text="", width=70, height=40)
            lbl.grid(row=row, column=i, padx=2, pady=2)
        col = start_weekday
        for day in range(1, days_in_month + 1):
            date = jdatetime.date(self.current_year, self.current_month, day)
            btn = ctk.CTkButton(self.calendar_frame, text=str(day), width=70, height=40, font=self.font_farsi, fg_color="transparent", border_width=1, command=lambda d=date: self.on_day_click(d))
            btn.grid(row=row, column=col, padx=2, pady=2)
            if self.start_date == date:
                btn.configure(fg_color=COLOR_PRIMARY, text_color="white")
            elif self.end_date == date:
                btn.configure(fg_color=COLOR_ACCENT, text_color="white")
            col += 1
            if col >= 7:
                col = 0
                row += 1
        self.update_selection_label()
    def on_day_click(self, date):
        if self.selection_mode == "start":
            self.start_date = date
            self.selection_mode = "end"
            self.start_mode_btn.configure(fg_color="gray", text_color="white")
            self.end_mode_btn.configure(fg_color=COLOR_PRIMARY, text_color="white")
        else:
            if self.start_date is None:
                messagebox.showerror("خطا", "ابتدا تاریخ شروع را انتخاب کنید.")
                return
            if date < self.start_date:
                messagebox.showerror("خطا", "تاریخ پایان نمی‌تواند قبل از تاریخ شروع باشد.")
                return
            self.end_date = date
            self.selection_mode = "start"
            self.start_mode_btn.configure(fg_color=COLOR_PRIMARY, text_color="white")
            self.end_mode_btn.configure(fg_color="gray", text_color="white")
        self.update_calendar()
        if self.start_date and self.end_date:
            self.confirm_btn.configure(state="normal")
        else:
            self.confirm_btn.configure(state="disabled")
    def set_start_mode(self):
        self.selection_mode = "start"
        self.start_mode_btn.configure(fg_color=COLOR_PRIMARY, text_color="white")
        self.end_mode_btn.configure(fg_color="gray", text_color="white")
    def set_end_mode(self):
        if self.start_date is None:
            messagebox.showerror("خطا", "ابتدا تاریخ شروع را انتخاب کنید.")
            return
        self.selection_mode = "end"
        self.start_mode_btn.configure(fg_color="gray", text_color="white")
        self.end_mode_btn.configure(fg_color=COLOR_PRIMARY, text_color="white")
    def update_selection_label(self):
        start_str = self.start_date.strftime("%Y/%m/%d") if self.start_date else "انتخاب نشده"
        end_str = self.end_date.strftime("%Y/%m/%d") if self.end_date else "انتخاب نشده"
        self.selection_label.configure(text=f"شروع: {start_str} | پایان: {end_str}")
    def confirm(self):
        if self.start_date and self.end_date:
            self.result = {
                'start_day': self.start_date.day,
                'end_day': self.end_date.day,
                'start_month': MONTH_LIST[self.start_date.month - 1],
                'start_year': self.start_date.year,
                'end_month': MONTH_LIST[self.end_date.month - 1],
                'end_year': self.end_date.year,
                'period_range': f"از {self.start_date.day} {MONTH_LIST[self.start_date.month-1]} {self.start_date.year} تا {self.end_date.day} {MONTH_LIST[self.end_date.month-1]} {self.end_date.year}"
            }
            self.destroy()
        else:
            messagebox.showerror("خطا", "لطفاً هر دو تاریخ را انتخاب کنید.")

class ServiceEditDialog(ctk.CTkToplevel):
    def __init__(self, parent, service_data=None, service_categories=None):
        """
        service_categories: لیست نام سرویس‌های فعال (مثلاً ['ویزا', 'اتوبوس', ...]) که از دیتابیس گرفته می‌شود
        """
        super().__init__(parent)
        self.title("ویرایش خدمت")
        self.geometry("550x500")
        self.transient(parent)
        self.grab_set()
        self.result = None
        self.font_farsi = ("B Nazanin", 16)

        # استفاده از سرویس‌های پویای ارسالی، در غیر این صورت لیست خالی
        if service_categories is None:
            service_categories = []
        self.service_categories = service_categories

        if service_data is None:
            service_data = {"type": "", "passenger": "", "contract": "", "date": "", "notes": "", "debt": 0, "credit": 0}

        ctk.CTkLabel(self, text="ویرایش اطلاعات خدمت", font=("B Nazanin", 20, "bold")).pack(pady=15)
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(pady=15, padx=20, fill="both", expand=True)

        # انتخاب نوع خدمت (داینامیک)
        ctk.CTkLabel(main_frame, text="نوع خدمت", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.type_combo = ctk.CTkComboBox(main_frame, values=self.service_categories, font=self.font_farsi, dropdown_font=self.font_farsi, width=250)
        self.type_combo.set(service_data.get("type", ""))
        self.type_combo.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        # سایر فیلدها بدون تغییر
        ctk.CTkLabel(main_frame, text="نام مسافر", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.passenger_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.passenger_entry.insert(0, service_data.get("passenger", ""))
        self.passenger_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(main_frame, text="شماره قرارداد", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.contract_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.contract_entry.insert(0, service_data.get("contract", ""))
        self.contract_entry.grid(row=2, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(main_frame, text="تاریخ", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.date_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.date_entry.insert(0, service_data.get("date", ""))
        self.date_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(main_frame, text="توضیحات", font=self.font_farsi).grid(row=4, column=0, padx=5, pady=5, sticky="e")
        self.notes_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.notes_entry.insert(0, service_data.get("notes", ""))
        self.notes_entry.grid(row=4, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(main_frame, text="بدهکار (ریال)", font=self.font_farsi).grid(row=5, column=0, padx=5, pady=5, sticky="e")
        self.debt_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.debt_entry.insert(0, to_persian_digits(str(service_data.get("debt", 0))))
        self.debt_entry.grid(row=5, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(main_frame, text="بستانکار (ریال)", font=self.font_farsi).grid(row=6, column=0, padx=5, pady=5, sticky="e")
        self.credit_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.credit_entry.insert(0, to_persian_digits(str(service_data.get("credit", 0))))
        self.credit_entry.grid(row=6, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkLabel(main_frame, text="مانده", font=self.font_farsi).grid(row=7, column=0, padx=5, pady=5, sticky="e")
        self.balance_label = ctk.CTkLabel(main_frame, text="0", font=self.font_farsi, fg_color="#333333", corner_radius=5, width=20)
        self.balance_label.grid(row=7, column=1, padx=5, pady=5, sticky="w")

        self.update_balance()
        self.debt_entry.bind("<KeyRelease>", lambda e: self.update_balance())
        self.credit_entry.bind("<KeyRelease>", lambda e: self.update_balance())

        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(pady=15)
        ctk.CTkButton(btn_frame, text="ذخیره", command=self.save, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="انصراف", command=self.destroy, font=self.font_farsi, fg_color="gray", hover_color="#555555", width=120).pack(side="left", padx=5)

    def update_balance(self):
        try:
            debt = persian_to_english_number(self.debt_entry.get())
            credit = persian_to_english_number(self.credit_entry.get())
            balance = debt - credit
            self.balance_label.configure(text=to_persian_digits(f"{balance:,}"))
        except:
            self.balance_label.configure(text="خطا")

    def save(self):
        try:
            debt = persian_to_english_number(self.debt_entry.get())
            credit = persian_to_english_number(self.credit_entry.get())
        except:
            messagebox.showerror("خطا", "مقادیر بدهکار و بستانکار باید عدد باشند.")
            return
        service_type = self.type_combo.get()
        if not service_type:
            messagebox.showerror("خطا", "نوع خدمت را انتخاب کنید.")
            return
        passenger = self.passenger_entry.get().strip()
        if not passenger:
            messagebox.showerror("خطا", "نام مسافر نمی‌تواند خالی باشد.")
            return
        contract = self.contract_entry.get().strip()
        date = self.date_entry.get().strip()
        notes = self.notes_entry.get().strip()
        self.result = {"type": service_type, "passenger": passenger, "contract": contract, "date": date, "notes": notes, "debt": debt, "credit": credit, "balance": debt - credit}
        self.destroy()

class FlightEditDialog(ctk.CTkToplevel):
    def __init__(self, parent, flight_data=None):
        super().__init__(parent)
        self.title("ویرایش پرواز")
        self.geometry("600x500")
        self.transient(parent)
        self.grab_set()
        self.result = None
        self.font_farsi = ("B Nazanin", 16)
        if flight_data is None:
            flight_data = {"contract": "", "passenger": "", "route": "", "ticket": "", "date": "", "notes": "", "debt": 0, "credit": 0}
        ctk.CTkLabel(self, text="ویرایش اطلاعات پرواز", font=("B Nazanin", 20, "bold")).pack(pady=15)
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(pady=15, padx=20, fill="both", expand=True)
        ctk.CTkLabel(main_frame, text="شماره قرارداد", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.contract_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.contract_entry.insert(0, flight_data.get("contract", ""))
        self.contract_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="نام مسافر", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.passenger_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.passenger_entry.insert(0, flight_data.get("passenger", ""))
        self.passenger_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="مسیر", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.route_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.route_entry.insert(0, flight_data.get("route", ""))
        self.route_entry.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="شماره بلیط", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.ticket_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.ticket_entry.insert(0, flight_data.get("ticket", ""))
        self.ticket_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="تاریخ پرواز", font=self.font_farsi).grid(row=4, column=0, padx=5, pady=5, sticky="e")
        self.date_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.date_entry.insert(0, flight_data.get("date", ""))
        self.date_entry.grid(row=4, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="توضیحات", font=self.font_farsi).grid(row=5, column=0, padx=5, pady=5, sticky="e")
        self.notes_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.notes_entry.insert(0, flight_data.get("notes", ""))
        self.notes_entry.grid(row=5, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="بدهکار (ریال)", font=self.font_farsi).grid(row=6, column=0, padx=5, pady=5, sticky="e")
        self.debt_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.debt_entry.insert(0, to_persian_digits(str(flight_data.get("debt", 0))))
        self.debt_entry.grid(row=6, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="بستانکار (ریال)", font=self.font_farsi).grid(row=7, column=0, padx=5, pady=5, sticky="e")
        self.credit_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.credit_entry.insert(0, to_persian_digits(str(flight_data.get("credit", 0))))
        self.credit_entry.grid(row=7, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="مانده", font=self.font_farsi).grid(row=8, column=0, padx=5, pady=5, sticky="e")
        self.balance_label = ctk.CTkLabel(main_frame, text="0", font=self.font_farsi, fg_color="#333333", corner_radius=5, width=20)
        self.balance_label.grid(row=8, column=1, padx=5, pady=5, sticky="w")
        self.update_balance()
        self.debt_entry.bind("<KeyRelease>", lambda e: self.update_balance())
        self.credit_entry.bind("<KeyRelease>", lambda e: self.update_balance())
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(pady=15)
        ctk.CTkButton(btn_frame, text="ذخیره", command=self.save, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="انصراف", command=self.destroy, font=self.font_farsi, fg_color="gray", hover_color="#555555", width=120).pack(side="left", padx=5)
    def update_balance(self):
        try:
            debt = persian_to_english_number(self.debt_entry.get())
            credit = persian_to_english_number(self.credit_entry.get())
            balance = debt - credit
            self.balance_label.configure(text=to_persian_digits(f"{balance:,}"))
        except:
            self.balance_label.configure(text="خطا")
    def save(self):
        try:
            debt = persian_to_english_number(self.debt_entry.get())
            credit = persian_to_english_number(self.credit_entry.get())
        except:
            messagebox.showerror("خطا", "مقادیر بدهکار و بستانکار باید عدد باشند.")
            return
        contract = self.contract_entry.get().strip()
        passenger = self.passenger_entry.get().strip()
        if not passenger:
            messagebox.showerror("خطا", "نام مسافر نمی‌تواند خالی باشد.")
            return
        route = self.route_entry.get().strip()
        ticket = self.ticket_entry.get().strip()
        date = self.date_entry.get().strip()
        notes = self.notes_entry.get().strip()
        self.result = {"contract": contract, "passenger": passenger, "route": route, "ticket": ticket, "date": date, "notes": notes, "debt": debt, "credit": credit, "balance": debt - credit}
        self.destroy()

class HotelEditDialog(ctk.CTkToplevel):
    def __init__(self, parent, hotel_data=None):
        super().__init__(parent)
        self.title("ویرایش هتل")
        self.geometry("650x550")
        self.transient(parent)
        self.grab_set()
        self.result = None
        self.font_farsi = ("B Nazanin", 16)
        if hotel_data is None:
            hotel_data = {"contract": "", "passenger": "", "hotel": "", "room": "", "pax": 1, "date": "", "notes": "", "debt": 0, "credit": 0}
        ctk.CTkLabel(self, text="ویرایش اطلاعات هتل", font=("B Nazanin", 20, "bold")).pack(pady=15)
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(pady=15, padx=20, fill="both", expand=True)
        ctk.CTkLabel(main_frame, text="شماره قرارداد", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.contract_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.contract_entry.insert(0, hotel_data.get("contract", ""))
        self.contract_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="نام مسافر", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.passenger_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.passenger_entry.insert(0, hotel_data.get("passenger", ""))
        self.passenger_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="نام هتل", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.hotel_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.hotel_entry.insert(0, hotel_data.get("hotel", ""))
        self.hotel_entry.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="نوع اتاق", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.room_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.room_entry.insert(0, hotel_data.get("room", ""))
        self.room_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="تعداد نفرات", font=self.font_farsi).grid(row=4, column=0, padx=5, pady=5, sticky="e")
        self.pax_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.pax_entry.insert(0, str(hotel_data.get("pax", 1)))
        self.pax_entry.grid(row=4, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="تاریخ ورود/خروج", font=self.font_farsi).grid(row=5, column=0, padx=5, pady=5, sticky="e")
        self.date_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.date_entry.insert(0, hotel_data.get("date", ""))
        self.date_entry.grid(row=5, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="توضیحات", font=self.font_farsi).grid(row=6, column=0, padx=5, pady=5, sticky="e")
        self.notes_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.notes_entry.insert(0, hotel_data.get("notes", ""))
        self.notes_entry.grid(row=6, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="بدهکار (ریال)", font=self.font_farsi).grid(row=7, column=0, padx=5, pady=5, sticky="e")
        self.debt_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.debt_entry.insert(0, to_persian_digits(str(hotel_data.get("debt", 0))))
        self.debt_entry.grid(row=7, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="بستانکار (ریال)", font=self.font_farsi).grid(row=8, column=0, padx=5, pady=5, sticky="e")
        self.credit_entry = ctk.CTkEntry(main_frame, width=250, font=self.font_farsi)
        self.credit_entry.insert(0, to_persian_digits(str(hotel_data.get("credit", 0))))
        self.credit_entry.grid(row=8, column=1, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(main_frame, text="مانده", font=self.font_farsi).grid(row=9, column=0, padx=5, pady=5, sticky="e")
        self.balance_label = ctk.CTkLabel(main_frame, text="0", font=self.font_farsi, fg_color="#333333", corner_radius=5, width=20)
        self.balance_label.grid(row=9, column=1, padx=5, pady=5, sticky="w")
        self.update_balance()
        self.debt_entry.bind("<KeyRelease>", lambda e: self.update_balance())
        self.credit_entry.bind("<KeyRelease>", lambda e: self.update_balance())
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(pady=15)
        ctk.CTkButton(btn_frame, text="ذخیره", command=self.save, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="انصراف", command=self.destroy, font=self.font_farsi, fg_color="gray", hover_color="#555555", width=120).pack(side="left", padx=5)
    def update_balance(self):
        try:
            debt = persian_to_english_number(self.debt_entry.get())
            credit = persian_to_english_number(self.credit_entry.get())
            balance = debt - credit
            self.balance_label.configure(text=to_persian_digits(f"{balance:,}"))
        except:
            self.balance_label.configure(text="خطا")
    def save(self):
        try:
            debt = persian_to_english_number(self.debt_entry.get())
            credit = persian_to_english_number(self.credit_entry.get())
            pax = int(english_number(self.pax_entry.get())) if self.pax_entry.get().strip() else 1
        except:
            messagebox.showerror("خطا", "مقادیر عددی را درست وارد کنید.")
            return
        contract = self.contract_entry.get().strip()
        passenger = self.passenger_entry.get().strip()
        if not passenger:
            messagebox.showerror("خطا", "نام مسافر نمی‌تواند خالی باشد.")
            return
        hotel = self.hotel_entry.get().strip()
        room = self.room_entry.get().strip()
        date = self.date_entry.get().strip()
        notes = self.notes_entry.get().strip()
        self.result = {"contract": contract, "passenger": passenger, "hotel": hotel, "room": room, "pax": pax, "date": date, "notes": notes, "debt": debt, "credit": credit, "balance": debt - credit}
        self.destroy()

class LoginDialog(ctk.CTkToplevel):
    def __init__(self, parent, user_manager, db_manager):
        super().__init__(parent)
        self.title("ورود به سیستم")
        self.geometry("400x350")
        self.transient(parent)
        self.grab_set()
        self.user_manager = user_manager
        self.db_manager = db_manager
        self.result = None
        self.font_farsi = ("B Nazanin", 18)
        ctk.CTkLabel(self, text="ورود به نرم افزار", font=("B Nazanin", 22, "bold")).pack(pady=20)
        frame = ctk.CTkFrame(self)
        frame.pack(pady=15, padx=20, fill="both", expand=True)
        ctk.CTkLabel(frame, text="نام کاربری", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.username_entry = ctk.CTkEntry(frame, width=220, font=self.font_farsi)
        self.username_entry.grid(row=0, column=1, padx=5, pady=5)
        ctk.CTkLabel(frame, text="رمز عبور", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.password_entry = ctk.CTkEntry(frame, width=220, font=self.font_farsi, show="*")
        self.password_entry.grid(row=1, column=1, padx=5, pady=5)
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(pady=15)
        ctk.CTkButton(btn_frame, text="ورود", command=self.login, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="انصراف", command=self.destroy, font=self.font_farsi, fg_color="gray", hover_color="#555555", width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="⚙ تنظیمات اتصال", command=self.open_server_config, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=150).pack(side="left", padx=5)
        self.bind("<Return>", lambda e: self.login())
    def login(self):
        username = self.username_entry.get().strip()
        password = self.password_entry.get()
        user = self.user_manager.authenticate(username, password)
        if user:
            self.result = user
            self.destroy()
        else:
            messagebox.showerror("خطا", "نام کاربری یا رمز عبور اشتباه است.")
    def open_server_config(self):
        dialog = ServerConfigDialog(self, self.db_manager)
        self.wait_window(dialog)

class ServerConfigDialog(ctk.CTkToplevel):
    def __init__(self, parent, db_manager):
        super().__init__(parent)
        self.title("تنظیمات اتصال به سرور")
        self.geometry("500x400")
        self.transient(parent)
        self.grab_set()
        self.db_manager = db_manager
        self.font_farsi = ("B Nazanin", 16)
        ctk.CTkLabel(self, text="تنظیمات اتصال به SQL Server", font=("B Nazanin", 20, "bold")).pack(pady=15)
        frame = ctk.CTkFrame(self)
        frame.pack(pady=15, padx=20, fill="both", expand=True)
        ctk.CTkLabel(frame, text="آدرس سرور", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.server_entry = ctk.CTkEntry(frame, width=250, font=self.font_farsi)
        self.server_entry.grid(row=0, column=1, padx=5, pady=5)
        self.server_entry.insert(0, self.db_manager.config.get('server', ''))
        ctk.CTkLabel(frame, text="نام دیتابیس", font=self.font_farsi).grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.db_entry = ctk.CTkEntry(frame, width=250, font=self.font_farsi)
        self.db_entry.grid(row=1, column=1, padx=5, pady=5)
        self.db_entry.insert(0, self.db_manager.config.get('database', ''))
        ctk.CTkLabel(frame, text="نام کاربری", font=self.font_farsi).grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.user_entry = ctk.CTkEntry(frame, width=250, font=self.font_farsi)
        self.user_entry.grid(row=2, column=1, padx=5, pady=5)
        self.user_entry.insert(0, self.db_manager.config.get('username', ''))
        ctk.CTkLabel(frame, text="رمز عبور", font=self.font_farsi).grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.pass_entry = ctk.CTkEntry(frame, width=250, font=self.font_farsi, show="*")
        self.pass_entry.grid(row=3, column=1, padx=5, pady=5)
        self.pass_entry.insert(0, self.db_manager.config.get('password', ''))
        ctk.CTkLabel(conn_frame, text="آدرس سرور API", font=self.font_farsi).grid(row=4, column=0, padx=5, pady=5, sticky="e")
        self.api_entry = ctk.CTkEntry(conn_frame, width=250, font=self.font_farsi)
        self.api_entry.grid(row=4, column=1, padx=5, pady=5)
        self.api_entry.insert(0, self.db_manager.config.get('api_server', ''))
        btn_frame = ctk.CTkFrame(self)
        btn_frame.pack(pady=15)
        ctk.CTkButton(btn_frame, text="ذخیره", command=self.save_config, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="انصراف", command=self.destroy, font=self.font_farsi, fg_color="gray", hover_color="#555555", width=120).pack(side="left", padx=5)
    
    def save_config(self):
        config = {'server': self.server_entry.get().strip(), 'database': self.db_entry.get().strip(), 'username': self.user_entry.get().strip(), 'password': self.pass_entry.get().strip()}
        try:
            conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={config['server']};DATABASE={config['database']};UID={config['username']};PWD={config['password']}"
            pyodbc.connect(conn_str, timeout=5)
        except Exception as e:
            messagebox.showerror("خطا", f"اتصال به سرور ناموفق بود:\n{str(e)}")
            return
        base_path = get_base_path()
        config_path = os.path.join(base_path, 'db_config.json')
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        self.db_manager.config = config
        messagebox.showinfo("موفقیت", "تنظیمات ذخیره شد.")
        self.destroy()

class App(ctk.CTk):
    
    DEFAULT_HTML_TEMPLATE = """
    <div class="invoice-header">صورتحساب {{ service_type }} &nbsp;&nbsp; شماره: {{ invoice_number }} &nbsp;&nbsp; تاریخ: {{ invoice_date }}</div>
    <div class="box-title">مشخصات فروشنده</div>
    <table class="seller-box">
        <tr><td class="label">شناسه ملی:ERCERD<td colspan="3">{{ seller.national_id }}ERCERD</td></tr>
        <tr><td class="label">نام شخص حقوقی:ERCERD<td colspan="3">{{ seller.company_name }}ERCERD</td></tr>
        <tr><td class="label">نشانی:ERCERD<td colspan="3">{{ seller.address }}ERCERD</td></tr>
        <tr><td class="label">کدپستی:ERCERD<td colspan="3">{{ seller.postal_code }}ERCERD<td class="label">تلفن:ERCERD<td colspan="3">{{ seller.phone }}ERCERD</td></tr>
    </table>
    <div class="box-title">مشخصات خریدار</div>
    <table class="buyer-box">
        <tr><td class="label">شناسه ملی:ERCERD<td colspan="3">{{ buyer.national_id }}ERCERD</td></tr>
        <tr><td class="label">نام شخص حقوقی:ERCERD<td colspan="3">{{ buyer.name }}ERCERD</td></tr>
        <tr><td class="label">نشانی:ERCERD<td colspan="3">{{ buyer.address }}ERCERD</td></tr>
        <tr><td class="label">کدپستی:ERCERD<td colspan="3">{{ buyer.postal_code }}ERCERD<td class="label">تلفن:ERCERD<td colspan="3">{{ buyer.phone }}ERCERD</td></tr>
    </table>
    <table class="data-table"><thead><tr><th>ردیف</th><th>شماره قرارداد</th><th>نام مسافر</th><th>شرح خدمات</th><th>تاریخ</th><th>توضیحات</th><th>بدهکار (ریال)</th><th>بستانکار (ریال)</th><th>مانده (ریال)</th></tr></thead>
    <tbody>{% for row in page.rows %}<tr>
        <td>{{ loop.index + (page.page_num-1)*8 }}ERCERD
        <td>{{ row.contract }}ERCERD
        <td>{{ row.passenger }}ERCERD
        <td>{{ row.description }}ERCERD
        <td>{{ row.date }}ERCERD
        <td>{{ row.notes }}ERCERD
        <td>{{ row.debt|persian_number }}ERCERD
        <td>{{ row.credit|persian_number }}ERCERD
        <td>{{ row.balance|persian_number }}ERCERD
    </tr>{% endfor %}</tbody></table>
    <div class="totals-row"><div>جمع بدهکار این صفحه: {{ page.total_debt|persian_number }} ریال</div><div>جمع بستانکار این صفحه: {{ page.total_credit|persian_number }} ریال</div><div>مانده این صفحه: {{ page.total_balance|persian_number }} ریال</div></div>
    {% if page.is_last %}<div class="totals-row"><div>جمع کل نهایی بدهکار: {{ grand_total_debt|persian_number }} ریال</div><div>جمع کل نهایی بستانکار: {{ grand_total_credit|persian_number }} ریال</div><div>مانده نهایی: {{ grand_total_balance|persian_number }} ریال</div></div>
    <div class="footer"><div class="signature">{% if seller.signature_path %}<img src="file:///{{ seller.signature_path }}">{% else %}<p>................................</p>{% endif %}<p>{{ seller.manager_name }}<br>{{ seller.manager_position }}</p></div>
    <div class="stamp">{% if seller.stamp_path %}<img src="file:///{{ seller.stamp_path }}">{% else %}<p>................................</p>{% endif %}<p>مهر شرکت</p></div></div>{% endif %}
    """

    DEFAULT_CSS_STYLE = """
        @page { size: A4 landscape; margin: 1cm; }
        body { font-family: 'B Nazanin', Tahoma, Arial, sans-serif; font-size: 9px; line-height: 1.2; }
        .page { page-break-after: avoid; }
        .seller-box, .buyer-box { border: 1px solid #000; border-collapse: collapse; width: 100%; margin-bottom: 6px; }
        .seller-box td, .buyer-box td { border: 1px solid #000; padding: 4px; vertical-align: top; }
        .seller-box .label, .buyer-box .label { font-weight: bold; background-color: #f0f0f0; width: 25%; }
        .box-title { font-weight: bold; margin: 3px 0; text-align: center; background-color: #e0e0e0; padding: 3px; font-size: 10px; }
        .invoice-header { text-align: center; font-size: 14px; font-weight: bold; margin-bottom: 10px; }
        .data-table { width: 100%; border-collapse: collapse; margin: 8px 0; border: 1px solid #000; }
        .data-table th, .data-table td { border: 1px solid #000; padding: 3px; text-align: center; }
        .data-table th { background-color: #f0f0f0; font-weight: bold; font-size: 9px; }
        .totals-row { margin-top: 6px; display: flex; justify-content: space-between; border-top: 1px solid #000; padding-top: 4px; font-size: 9px; }
        .footer { margin-top: 20px; display: flex; justify-content: space-between; align-items: flex-end; }
    """
    def __init__(self):
        super().__init__()
        self.after(0, self._set_icon_after_load)
        self.title("مدیریت صورتحساب سفرهای کرمان خودرو")
        self.geometry("1200x950")
        self.resizable(True, True)
        
        # مقداردهی اولیه مجوزها (پیش‌فرض False)
        self.can_view = False
        self.can_pay = False
        self.can_issue = False
        self.can_manage_users = False
        self.can_manage_settings = False
        self.can_delete_history = False
        self.can_delete_payment = False
        self.can_export = False
        self.can_manage_companies = False
        self.can_submit = False
        self.can_approve = False
        
        self.base_path = get_base_path()
        self.config_path = os.path.join(self.base_path, 'db_config.json')
        with open(self.config_path, 'r', encoding='utf-8') as f:
            db_config = json.load(f)
        self.api_server = db_config.get('api_server', '127.0.0.1')
        self.db = SQLServerDatabase(db_config)
        self.db_manager = DatabaseManager(db_config)
        self.user_manager = UserManager(self.db)
        self.settings = self.load_settings()
        self.backup_dir = os.path.join(os.path.expanduser("~"), "Documents", "KKTCOIG_Backups")
        os.makedirs(self.backup_dir, exist_ok=True)
        
        self.current_user = None
        self.show_login()   # دیالوگ ورود (Modal)
        if not self.current_user:
            self.destroy()
            return
        
        # به‌روزرسانی مجوزها از کاربر وارد شده
        (self.user_id, self.username, self.can_view, self.can_pay, self.can_issue,
         self.can_manage_users, self.can_manage_settings, self.can_delete_history,
         self.can_delete_payment, self.can_export, self.can_manage_companies,
         self.can_submit, self.can_approve, self.can_edit_template) = self.current_user
        
        # ساخت رابط کاربری
        self._build_ui()
        self.update_buttons_by_permission()
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
    
    def _build_ui(self):
        """ساخت تمام اجزای رابط کاربری بعد از لاگین موفق"""
        self.excel_path = ctk.StringVar()
        self.payer = ctk.StringVar()
        self.generated_letter_number = ctk.StringVar(value="")
        self.year_letter = ctk.StringVar()
        self.month_letter = ctk.StringVar()
        self.day_letter = ctk.StringVar()
        self.letter_date = ctk.StringVar()
        self.start_day = ctk.StringVar(value="")
        self.end_day = ctk.StringVar(value="")
        self.start_month = ctk.StringVar(value="")
        self.start_year = ctk.StringVar(value="")
        self.end_month = ctk.StringVar(value="")
        self.end_year = ctk.StringVar(value="")
        self.period_range = ctk.StringVar(value="")
        self.period_display = ctk.StringVar(value="هنوز بازه‌ای انتخاب نشده")
        self.payers_list = self.db.get_all_payers()
        self.font_farsi = ctk.CTkFont(family="B Nazanin", size=16)
        self.font_farsi_bold = ctk.CTkFont(family="B Nazanin", size=16, weight="bold")
        self.font_farsi_header = ctk.CTkFont(family="B Nazanin", size=22, weight="bold")
        self.current_step = 0
        self.frames = []
        self.selected_files = []
        self.file_items = []
        self.service_data = []
        self.flight_data = []
        self.hotel_data = []
        
        # منوی بالایی
        self.menu_frame = ctk.CTkFrame(self, height=70, corner_radius=0, fg_color="#2b2b2b")
        self.menu_frame.pack(fill="x", padx=0, pady=0)
        self.menu_frame.pack_propagate(False)
        self.btn_logout = ctk.CTkButton(self.menu_frame, text="🚪 خروج از حساب", width=140, command=self.logout, font=self.font_farsi, fg_color="#d32f2f", hover_color="#b71c1c", text_color="white")
        self.btn_logout.pack(side="left", padx=10, pady=5)
        
        # بازسازی منو (ساخت دکمه‌های سمت راست بر اساس مجوزها)
        self.rebuild_menu()
        
        # لوگو و عنوان
        self.logo_frame = ctk.CTkFrame(self, height=140, fg_color="transparent", corner_radius=20)
        self.logo_frame.pack(pady=(20, 10), padx=20, fill="x")
        ctk.CTkLabel(self.logo_frame, text="نرم افزار مدیریت صورتحساب سفرهای کرمان خودرو", font=self.font_farsi_header).pack(expand=True)
        
        # کانتینر اصلی مراحل
        self.main_container = ctk.CTkFrame(self, corner_radius=20)
        self.main_container.pack(pady=15, padx=20, fill="both", expand=True)
        
        # نوار دکمه‌های پایین صفحه
        self.nav_frame = ctk.CTkFrame(self, corner_radius=20)
        self.nav_frame.pack(pady=15, padx=20, fill="x")
        self.btn_home = ctk.CTkButton(self.nav_frame, text="خانه", command=self.go_home, font=self.font_farsi, width=100, corner_radius=12, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT)
        self.btn_home.pack(side="left", padx=5)
        self.btn_prev = ctk.CTkButton(self.nav_frame, text="قبلی", command=self.prev_step, font=self.font_farsi, width=120, corner_radius=12, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, state="disabled")
        self.btn_prev.pack(side="left", padx=5)
        self.btn_next = ctk.CTkButton(self.nav_frame, text="بعدی", command=self.next_step, font=self.font_farsi, width=120, corner_radius=12, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, state="disabled")
        self.btn_next.pack(side="left", padx=5)
        self.btn_action = ctk.CTkButton(self.nav_frame, text="", font=self.font_farsi_bold, width=220, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, corner_radius=12, state="disabled")
        self.btn_action.pack(side="right", padx=5)
        
        # نوار پیشرفت
        self.progress_frame = ctk.CTkFrame(self, corner_radius=20, height=40)
        self.progress_frame.pack(pady=(5, 10), padx=20, fill="x")
        self.progress_bar = ctk.CTkProgressBar(self.progress_frame, width=800, height=25, corner_radius=10, progress_color=COLOR_PRIMARY)
        self.progress_bar.pack(pady=5, padx=10)
        self.progress_bar.set(0)
        
        # ساخت فریم‌های مراحل
        self.create_step_0()
        self.create_step_1()
        self.create_step_2()
        self.create_step_3()
        self.create_step_4()
        self.create_step_5()
        
        # نمایش مرحله ابتدایی
        self.show_step(0)
    
    #def set_window_icon(self):
        #try:
            #icon_path = resource_path("KKTCO.ico")
           # if os.path.exists(icon_path):
           #     self.after(200, self._apply_icon)
       # except Exception as e:
          #  print(f"icon error: {e}")

    def _set_icon_after_load(self):
        """تنظیم آیکون بعد از بارگذاری کامل پنجره"""
        self.update_idletasks()
        self.after(100, self._force_set_icon)

    def _force_set_icon(self):
        """اجبار به تنظیم آیکون با روش ویندوز"""
        try:
            import ctypes
            icon_path = resource_path("KKTCO.ico")
            if not os.path.exists(icon_path):
                return
            # روش مستقیم Win32 API
            hwnd = ctypes.windll.user32.GetParent(self.winfo_id())
            if not hwnd:
                hwnd = self.winfo_id()
            # بارگذاری آیکون با سایزهای مختلف
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID('KKTCO.Invoice.v1')
            # آیکون کوچک و بزرگ
            hicon_small = ctypes.windll.user32.LoadImageW(0, icon_path, 1, 16, 16, 0x00000010 | 0x00002000)
            hicon_big = ctypes.windll.user32.LoadImageW(0, icon_path, 1, 32, 32, 0x00000010 | 0x00002000)
            if hicon_small:
                ctypes.windll.user32.SendMessageW(hwnd, 0x0080, 0, hicon_small)   # WM_SETICON, ICON_SMALL
            if hicon_big:
                ctypes.windll.user32.SendMessageW(hwnd, 0x0080, 1, hicon_big)     # WM_SETICON, ICON_BIG
            # همچنین متد tkinter را هم صدا بزن
            self.iconbitmap(icon_path)
            self.wm_iconbitmap(icon_path)
            self.update()
        except Exception as e:
            print("Icon error:", e)
            
    def _apply_icon(self):
        try:
            icon_path = resource_path("KKTCO.ico")
            if os.path.exists(icon_path):
                # مرحله ۱: تنظیم Application ID مخصوص ویندوز
                import ctypes
                ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID('KKTCO.IG.v1')
                
                # مرحله ۲: تنظیم آیکون با تاخیر با دو روش مختلف
                self.iconbitmap(icon_path)  # روش استاندارد
                
                # مرحله ۳: با متد wm_iconbitmap هم امتحان کن (مخصوص ویندوز)
                self.wm_iconbitmap(icon_path)
                
                # مرحله ۴: پنجره را مجبور به آپدیت کن
                self.update_idletasks()
        except Exception as e:
            print(f"icon error: {e}")

    def update_buttons_by_permission(self):
        if hasattr(self, 'btn_start'):
            self.btn_start.configure(state="normal" if (self.can_issue or self.can_submit) else "disabled")
        if hasattr(self, 'btn_management'):
            self.btn_management.configure(state="normal" if self.can_view else "disabled")
        
        # برای دکمه‌های منو، اگر وجود داشته باشند آن‌ها را مخفی/نمایش بده
        if hasattr(self, 'btn_companies'):
            if self.can_manage_companies:
                self.btn_companies.pack(side="right", padx=5)
            else:
                self.btn_companies.pack_forget()
        if hasattr(self, 'btn_settings'):
            if self.can_manage_settings:
                self.btn_settings.pack(side="right", padx=5)
            else:
                self.btn_settings.pack_forget()
        if hasattr(self, 'btn_workflow'):
            if self.can_approve or self.can_submit:
                self.btn_workflow.pack(side="right", padx=5)
            else:
                self.btn_workflow.pack_forget()
        
        if not (self.can_issue or self.can_submit) and self.current_step > 0:
            self.show_step(0)

    def update_action_button(self):
        if self.current_step == len(self.frames)-1:
            if self.can_approve:
                self.btn_action.configure(text="تولید مستقیم (مدیر)", command=self.generate_direct, state="normal")
            elif self.can_submit:
                self.btn_action.configure(text="ارسال برای تأیید", command=self.submit_for_approval, state="normal")
            else:
                self.btn_action.configure(text="بدون دسترسی", state="disabled")
        else:
            self.btn_action.configure(state="disabled")
    def show_login(self):
        login = LoginDialog(self, self.user_manager, self.db_manager)
        self.wait_window(login)
        if login.result:
            self.current_user = login.result
        else:
            self.quit()
    def load_settings(self):
        settings_path = os.path.join(self.base_path, "settings.json")
        default_settings = {"letter_format": "YYYY/CODE/NNNNN"}
        try:
            with open(settings_path, "r", encoding="utf-8") as f:
                settings = json.load(f)
                return {**default_settings, **settings}
        except FileNotFoundError:
            return default_settings
    def save_settings(self, settings):
        settings_path = os.path.join(self.base_path, "settings.json")
        with open(settings_path, "w", encoding="utf-8") as f:
            json.dump(settings, f, ensure_ascii=False, indent=2)
    def open_settings(self):
        if self.can_manage_settings:
            dialog = SettingsDialog(self, self)
            self.wait_window(dialog)
        else:
            messagebox.showerror("خطا", "شما مجوز دسترسی به تنظیمات را ندارید.")
    def show_help(self):
        help_window = ctk.CTkToplevel(self)
        help_window.title("راهنما")
        help_window.geometry("600x500")
        help_window.transient(self)
        help_window.grab_set()
        text_widget = ctk.CTkTextbox(help_window, wrap="word", font=self.font_farsi)
        text_widget.pack(fill="both", expand=True, padx=15, pady=15)
        help_text = """
راهنمای استفاده از نرم‌افزار تولید صورتحساب

1. از صفحه اصلی، گزینه "شروع فرآیند جدید" را انتخاب کنید.
2. در مرحله اول، طرف حساب، شماره نامه، تاریخ را مشخص کنید.
3. در مرحله دوم، فایل‌های اکسل مربوط به صورتحساب‌ها را انتخاب کنید.
4. در مرحله سوم، اطلاعات خدمات، پروازها و هتل‌ها را ویرایش یا تکمیل کنید.
5. در مرحله چهارم، بازه صورتحساب را با استفاده از تقویم شمسی انتخاب کنید.
6. در مرحله آخر، اطلاعات را بررسی کرده و دکمه "ارسال برای تأیید" را بزنید (کاربر عادی) یا "تولید مستقیم" (مدیر).
   کاربران دارای مجوز تأیید می‌توانند از دکمه "کارتابل گردش کار" در نوار بالا درخواست‌ها را دیده و پس از تأیید، کاربر متقاضی می‌تواند فایل نهایی را صادر کند.
"""
        text_widget.insert("0.0", help_text)
        text_widget.configure(state="disabled")
    def show_about(self):
        about_window = ctk.CTkToplevel(self)
        about_window.title("درباره ما")
        about_window.geometry("450x350")
        about_window.transient(self)
        about_window.grab_set()
        text_widget = ctk.CTkTextbox(about_window, wrap="word", font=self.font_farsi)
        text_widget.pack(fill="both", expand=True, padx=15, pady=15)
        about_text = """
به نرم افزار مدیریت صورتحساب سفرهای کرمان خودرو خوش آمدید

نسخه برنامه: 3.0 (با گردش کار تأیید و صدور فایل نهایی توسط کاربر)
تاریخ بروزرسانی: 15 اردیبهشت ماه 1405
راه ارتباطی:
+989124585880
myazdanpanahfadaee@gmail.com
"""
        text_widget.insert("0.0", about_text)
        text_widget.configure(state="disabled")
    def go_home(self):
        self.show_step(0)
    def add_payer_db(self, name, code, parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email):
        try:
            self.db.add_payer(code, name, parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email)
            self.payers_list = self.db.get_all_payers()
            self.refresh_payer_combo()
            messagebox.showinfo("موفقیت", "شرکت با موفقیت اضافه شد.")
            return True
        except Exception as e:
            logging.error(f"خطا در افزودن پرداخت‌کننده: {e}")
            messagebox.showerror("خطا", f"خطا در افزودن شرکت:\n{str(e)}")
            return False
    def update_payer_db(self, old_code, new_name, new_code, new_parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email):
        try:
            self.db.update_payer(old_code, new_name, new_code, new_parent_code, national_id, economic_code, registration_number, address, postal_code, phone, customer_type, mobile, email)
            self.payers_list = self.db.get_all_payers()
            self.refresh_payer_combo()
            messagebox.showinfo("موفقیت", "شرکت با موفقیت ویرایش شد.")
            return True
        except Exception as e:
            logging.error(f"خطا در به‌روزرسانی پرداخت‌کننده: {e}")
            messagebox.showerror("خطا", f"خطا در ویرایش شرکت:\n{str(e)}")
            return False
    def delete_payer_db(self, code):
        self.db.delete_payer(code)
        self.payers_list = self.db.get_all_payers()
        self.refresh_payer_combo()
        messagebox.showinfo("موفقیت", "شرکت با موفقیت حذف شد.")
    def refresh_payer_combo(self):
        if hasattr(self, 'payer_combo'):
            display_values = [f"{p['code']} - {p['name']}" for p in self.payers_list]
            self.payer_combo.configure(values=display_values)
            if self.payer.get() not in display_values:
                self.payer.set("")
    def generate_letter_number(self, payer_code):
        if not payer_code:
            return ""
        try:
            response = requests.post(
                f"http://{self.api_server}:5000/reserve-letter",
                json={'payer_code': payer_code},
                timeout=10
            )
            if response.status_code == 200:
                return response.json().get('letter_number', '')
            else:
                messagebox.showerror("خطا", f"خطا در رزرو شماره: {response.text}")
                return ""
        except Exception as e:
            messagebox.showerror("خطا", f"ارتباط با سرور: {str(e)}")
            return ""
        
    def generate_multiple_letter_numbers(self, payer_code, count):
        year_two_digit = str(jdatetime.date.today().year)[-2:]
        numbers = self.db.reserve_letter_numbers(payer_code, year_two_digit, count)
        final_numbers = []
        for num in numbers:
            format_str = self.settings.get("letter_format", "YYYY/CODE/NNNNN")
            letter_number = format_str.replace("YYYY", str(jdatetime.date.today().year)).replace("YY", year_two_digit).replace("CODE", payer_code).replace("CCC", payer_code.zfill(3)).replace("NUM", str(num)).replace("NNNNN", str(num).zfill(5))
            final_numbers.append(letter_number)
        return final_numbers, numbers
    
    def save_to_history(self, letter_number, payer_name, payer_code, invoice_type, amount, issue_date, file_path=""):
        try:
            amount_int = int(amount)
            self.db.add_invoice(letter_number, payer_name, payer_code, invoice_type, amount_int, issue_date, file_path)
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در ذخیره تاریخچه:\n{str(e)}")
    def manage_payers(self):
        if not self.can_manage_companies:
            messagebox.showerror("خطا", "شما مجوز مدیریت شرکت‌ها را ندارید.")
            return
        dialog = ManagePayersDialog(self, self)
        self.wait_window(dialog)
    
    def generate_letter_for_current_payer(self):
        selected = self.payer.get()
        if not selected:
            self.generated_letter_number.set("")
            return
        parts = selected.split(" - ", 1)
        if len(parts) != 2:
            self.generated_letter_number.set("")
            return
        payer_code = re.sub(r'[^0-9]', '', parts[0])
        if not payer_code:
            self.generated_letter_number.set("")
            return
        today = jdatetime.date.today()
        year_two_digit = str(today.year % 100).zfill(2)
        # گرفتن شماره بعدی (بدون مصرف) با استفاده از یک متد کمکی
        next_num = self._peek_next_letter_number(payer_code, year_two_digit)
        if next_num is None:
            self.generated_letter_number.set("")
            return
        format_str = self.settings.get("letter_format", "YYYY/CODE/NNNNN")
        letter_number = format_str.replace("YYYY", str(today.year)) \
                                .replace("YY", year_two_digit) \
                                .replace("CODE", payer_code) \
                                .replace("CCC", payer_code.zfill(3)) \
                                .replace("NUM", str(next_num)) \
                                .replace("NNNNN", str(next_num).zfill(5))
        self.generated_letter_number.set(letter_number)

    def _peek_next_letter_number(self, payer_code, year):
        """متد کمکی برای دیدن شماره بعدی بدون مصرف"""
        with self.db.connect() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT MIN(number) FROM available_letter_numbers WHERE payer_code = ? AND year = ?", (payer_code, year))
            row = cursor.fetchone()
            if row and row[0] is not None:
                return row[0]
            cursor.execute("SELECT last_number FROM letter_sequences WHERE payer_code = ? AND year = ?", (payer_code, year))
            row = cursor.fetchone()
            return row[0] + 1 if row else 1

    def set_default_date(self):
        today = jdatetime.date.today()
        self.year_letter.set(persian_number(str(today.year)))
        self.month_letter.set(persian_number(str(today.month).zfill(2)))
        self.day_letter.set(persian_number(str(today.day).zfill(2)))
        self.letter_date.set(f"{today.year}/{today.month:02d}/{today.day:02d}")
    def create_step_0(self):
        frame = ctk.CTkFrame(self.main_container, corner_radius=20)
        self.frames.append(frame)
        ctk.CTkLabel(frame, text="به نرم‌افزار تولید صورتحساب خوش آمدید", font=self.font_farsi_header).pack(pady=50)
        self.btn_start = ctk.CTkButton(frame, text="شروع فرآیند جدید", command=lambda: self.show_step(1), font=self.font_farsi_bold, width=280, height=60, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, corner_radius=12)
        self.btn_start.pack(pady=20)
        self.btn_management = ctk.CTkButton(frame, text="مدیریت صورتحساب‌ها و حساب‌ها", command=self.open_management, font=self.font_farsi_bold, width=280, height=60, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, corner_radius=12)
        self.btn_management.pack(pady=20)
        self.update_buttons_by_permission()
    def open_management(self):
        if not self.can_view:
            messagebox.showerror("خطا", "شما مجوز مشاهده این بخش را ندارید.")
            return
        try:
            dialog = ManagementDialog(self, self, self.payers_list, self.can_view, self.can_pay,
                                    self.can_issue, self.can_delete_history)   # اضافه کردن can_delete_history
            self.wait_window(dialog)
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در باز کردن صفحه مدیریت:\n{str(e)}")

    def create_step_1(self):
        frame = ctk.CTkFrame(self.main_container, corner_radius=20)
        self.frames.append(frame)
        title = to_persian_digits("مرحله ۱ از ۵: اطلاعات نامه")
        ctk.CTkLabel(frame, text=title, font=self.font_farsi_header).pack(pady=20)
        fields_frame = ctk.CTkFrame(frame, corner_radius=15)
        fields_frame.pack(pady=20, padx=20, fill="both", expand=True)
        fields_frame.grid_columnconfigure(0, weight=0)
        fields_frame.grid_columnconfigure(1, weight=1)
        payer_container = ctk.CTkFrame(fields_frame, fg_color="transparent")
        payer_container.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(5, 10))
        payer_container.grid_columnconfigure(0, weight=0)
        payer_container.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(payer_container, text="طرف حساب", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        display_values = [f"{p['code']} - {p['name']}" for p in self.payers_list]
        self.payer_combo = ctk.CTkOptionMenu(payer_container, values=display_values, variable=self.payer, font=self.font_farsi, dropdown_font=self.font_farsi, fg_color=COLOR_PRIMARY, button_color=COLOR_ACCENT, button_hover_color="#d45a1c", width=380, dynamic_resizing=False, command=lambda _: self.generate_letter_for_current_payer())
        self.payer_combo.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        #self.payer_combo.bind("<<ComboboxSelected>>", lambda e: self.generate_letter_for_current_payer())
        self.btn_manage_payers = ctk.CTkButton(payer_container, text="⚙", command=self.manage_payers, width=50, font=self.font_farsi, corner_radius=10, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT)
        self.btn_manage_payers.grid(row=0, column=2, padx=5, pady=5)
        self.payer_error = ctk.CTkLabel(payer_container, text="", font=self.font_farsi, text_color="red")
        self.payer_error.grid(row=1, column=1, padx=5, pady=0, sticky="w")
        self.payer_error.grid_remove()
        letter_container = ctk.CTkFrame(fields_frame, fg_color="transparent")
        letter_container.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(5, 10))
        letter_container.grid_columnconfigure(0, weight=0)
        letter_container.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(letter_container, text="شماره نامه", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.letter_number_entry = ctk.CTkEntry(letter_container, textvariable=self.generated_letter_number, font=self.font_farsi, width=380, corner_radius=5)
        self.letter_number_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        #self.btn_generate_letter = ctk.CTkButton(letter_container, text="تولید شماره", command=self.generate_letter_for_current_payer, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120)
        #self.btn_generate_letter.grid(row=0, column=2, padx=5, pady=5)
        date_container = ctk.CTkFrame(fields_frame, fg_color="transparent")
        date_container.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(5, 10))
        date_container.grid_columnconfigure(0, weight=0)
        date_container.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(date_container, text="تاریخ نامه", font=self.font_farsi).grid(row=0, column=0, padx=5, pady=5, sticky="e")
        date_frame = ctk.CTkFrame(date_container, fg_color="transparent")
        date_frame.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        year_values = [persian_number(str(y)) for y in range(1400, 1431)]
        self.year_letter_combo = ctk.CTkComboBox(date_frame, values=year_values, variable=self.year_letter, width=100, font=self.font_farsi, dropdown_font=self.font_farsi, corner_radius=10)
        self.year_letter_combo.pack(side="left", padx=2)
        month_values = [persian_number(str(m).zfill(2)) for m in range(1, 13)]
        self.month_letter_combo = ctk.CTkComboBox(date_frame, values=month_values, variable=self.month_letter, width=80, font=self.font_farsi, dropdown_font=self.font_farsi, corner_radius=10)
        self.month_letter_combo.pack(side="left", padx=2)
        day_values = [persian_number(str(d).zfill(2)) for d in range(1, 32)]
        self.day_letter_combo = ctk.CTkComboBox(date_frame, values=day_values, variable=self.day_letter, width=80, font=self.font_farsi, dropdown_font=self.font_farsi, corner_radius=10)
        self.day_letter_combo.pack(side="left", padx=2)
        self.date_error = ctk.CTkLabel(date_container, text="", font=self.font_farsi, text_color="red")
        self.date_error.grid(row=1, column=1, padx=5, pady=0, sticky="w")
        self.date_error.grid_remove()
        self.set_default_date()
        ctk.CTkLabel(frame, text="لطفاً اطلاعات طرف حساب و نامه را وارد کنید.", font=self.font_farsi, text_color="gray").pack(pady=20)
    def hide_field_errors(self):
        self.payer_error.grid_remove()
        self.date_error.grid_remove()
    def show_field_error(self, field, message):
        message = to_persian_digits(message)
        if field == "payer":
            label = self.payer_error
            widget = self.payer_combo
        elif field == "date":
            label = self.date_error
            widget = self.year_letter_combo
        else:
            return
        label.configure(text=f"• {message}")
        label.grid()
        self.flash_widget(widget)
    def flash_widget(self, widget, count=3):
        if count <= 0:
            widget.configure(border_color=("gray", "gray"))
            return
        new_color = "red" if count % 2 == 0 else "gray"
        widget.configure(border_color=new_color)
        self.after(100, lambda: self.flash_widget(widget, count-1))
    def create_step_2(self):
        frame = ctk.CTkFrame(self.main_container, corner_radius=20)
        self.frames.append(frame)
        title = to_persian_digits("مرحله ۲ از ۵: انتخاب فایل‌های صورتحساب")
        ctk.CTkLabel(frame, text=title, font=self.font_farsi_header).pack(pady=20)
        add_file_btn = ctk.CTkButton(frame, text="➕ افزودن فایل اکسل", command=self.add_file, font=self.font_farsi_bold, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=220)
        add_file_btn.pack(pady=10)
        self.files_frame = ctk.CTkScrollableFrame(frame, height=220)
        self.files_frame.pack(pady=10, padx=20, fill="both", expand=True)
        self.clear_files()
        ctk.CTkLabel(frame, text="پس از افزودن فایل‌ها، نوع هر یک را بررسی و در صورت لزوم اصلاح کنید.", font=self.font_farsi, text_color="gray").pack(pady=10)
    def add_file(self):
        filenames = filedialog.askopenfilenames(title="انتخاب فایل‌های اکسل", filetypes=[("Excel files", "*.xlsx *.xls")])
        if not filenames:
            return
        for filename in filenames:
            if any(item.file_path == filename for item in self.file_items):
                continue
            try:
                df = pd.read_excel(filename)
                file_type = detect_file_type(df)
                item = FileItem(self.files_frame, filename, file_type, on_remove=self.remove_file, on_type_change=self.change_file_type)
                item.pack(fill="x", pady=2)
                self.file_items.append(item)
                self.selected_files.append({'path': filename, 'type': file_type})
            except Exception as e:
                messagebox.showerror("خطا", f"خطا در خواندن فایل {os.path.basename(filename)}:\n{str(e)}")
    def remove_file(self, item):
        for i, f in enumerate(self.selected_files):
            if f['path'] == item.file_path:
                del self.selected_files[i]
                break
        self.file_items.remove(item)
        item.destroy()
    def change_file_type(self, file_path, new_type):
        for f in self.selected_files:
            if f['path'] == file_path:
                f['type'] = new_type
                break
    def clear_files(self):
        for item in self.file_items:
            item.destroy()
        self.file_items.clear()
        self.selected_files.clear()
    def load_service_data_from_files(self):
        self.service_data = []
        for file_info in self.selected_files:
            if file_info['type'] == "خدمات":
                try:
                    df = pd.read_excel(file_info['path'])
                    out_df = process_service_data(df)
                    for _, row in out_df.iterrows():
                        self.service_data.append({"type": row.get("شرح خدمات", ""), "passenger": row["نام مسافر"], "contract": row["قرارداد"], "date": row.get("تاریخ", ""), "notes": row.get("توضیحات", ""), "debt": row["بدهکار"], "credit": row["بستانکار"], "balance": row["مانده"]})
                except Exception as e:
                    messagebox.showerror("خطا", f"خطا در خواندن فایل خدمات:\n{str(e)}")
    def load_flight_data_from_files(self):
        self.flight_data = []
        for file_info in self.selected_files:
            if file_info['type'] == "پرواز":
                try:
                    df = pd.read_excel(file_info['path'])
                    out_df = process_flight_data(df)
                    for _, row in out_df.iterrows():
                        self.flight_data.append({"contract": row["قرارداد"], "passenger": row["نام مسافر"], "route": row["مسیر"], "ticket": row["شماره بلیط"], "date": row["تاریخ پرواز"], "notes": "", "debt": row["بدهکار"], "credit": row["بستانکار"], "balance": row["مانده"]})
                except Exception as e:
                    messagebox.showerror("خطا", f"خطا در خواندن فایل پرواز:\n{str(e)}")
    def load_hotel_data_from_files(self):
        self.hotel_data = []
        for file_info in self.selected_files:
            if file_info['type'] == "هتل":
                try:
                    df = pd.read_excel(file_info['path'])
                    out_df = process_hotel_data(df)
                    for _, row in out_df.iterrows():
                        self.hotel_data.append({"contract": row["قرارداد"], "passenger": row["نام مسافر"], "hotel": row["هتل"], "room": row["نوع اتاق"], "pax": row["تعداد نفرات"], "date": row["تاریخ ورود و خروج"], "notes": "", "debt": row["بدهکار"], "credit": row["بستانکار"], "balance": row["مانده"]})
                except Exception as e:
                    messagebox.showerror("خطا", f"خطا در خواندن فایل هتل:\n{str(e)}")
    def refresh_service_tree(self):
        for item in self.service_tree.get_children():
            self.service_tree.delete(item)
        if not self.service_data:
            self.empty_label_service.place(relx=0.5, rely=0.5, anchor="center")
        else:
            self.empty_label_service.place_forget()
            for idx, svc in enumerate(self.service_data):
                values = (svc.get("type", ""), svc.get("passenger", ""), svc.get("contract", ""), svc.get("date", ""), svc.get("notes", ""), to_persian_digits(f"{svc.get('debt', 0):,}"), to_persian_digits(f"{svc.get('credit', 0):,}"), to_persian_digits(f"{svc.get('balance', 0):,}"))
                self.service_tree.insert("", "end", iid=idx, values=values)
    
    
    
    def add_service_item(self):
        service_names = self.get_service_names()   # خواندن از دیتابیس
        dialog = ServiceEditDialog(self, service_data=None, service_categories=service_names)
        self.wait_window(dialog)
        if dialog.result:   
            self.service_data.append(dialog.result)
            self.refresh_service_tree()

    def edit_service_item(self, event=None):
        selected = self.service_tree.selection()
        if not selected:
            messagebox.showwarning("انتخاب", "لطفاً یک ردیف را انتخاب کنید.")
            return
        idx = int(selected[0])
        old_data = self.service_data[idx]
        service_names = self.get_service_names()
        dialog = ServiceEditDialog(self, service_data=old_data, service_categories=service_names)
        self.wait_window(dialog)
        if dialog.result:
            self.service_data[idx] = dialog.result
            self.refresh_service_tree()

    def delete_service_item(self):
        selected = self.service_tree.selection()
        if not selected:
            messagebox.showwarning("انتخاب", "لطفاً یک ردیف را انتخاب کنید.")
            return
        idx = int(selected[0])
        if messagebox.askyesno("تأیید", "آیا از حذف این آیتم اطمینان دارید؟"):
            del self.service_data[idx]
            self.refresh_service_tree()
    def refresh_flight_tree(self):
        for item in self.flight_tree.get_children():
            self.flight_tree.delete(item)
        if not self.flight_data:
            self.empty_label_flight.place(relx=0.5, rely=0.5, anchor="center")
        else:
            self.empty_label_flight.place_forget()
            for idx, flt in enumerate(self.flight_data):
                values = (flt.get("contract", ""), flt.get("passenger", ""), flt.get("route", ""), flt.get("ticket", ""), flt.get("date", ""), flt.get("notes", ""), to_persian_digits(f"{flt.get('debt', 0):,}"), to_persian_digits(f"{flt.get('credit', 0):,}"), to_persian_digits(f"{flt.get('balance', 0):,}"))
                self.flight_tree.insert("", "end", iid=idx, values=values)
    def add_flight_item(self):
        dialog = FlightEditDialog(self)
        self.wait_window(dialog)
        if dialog.result:
            self.flight_data.append(dialog.result)
            self.refresh_flight_tree()
    def edit_flight_item(self, event=None):
        selected = self.flight_tree.selection()
        if not selected:
            messagebox.showwarning("انتخاب", "لطفاً یک ردیف را انتخاب کنید.")
            return
        idx = int(selected[0])
        old_data = self.flight_data[idx]
        dialog = FlightEditDialog(self, old_data)
        self.wait_window(dialog)
        if dialog.result:
            self.flight_data[idx] = dialog.result
            self.refresh_flight_tree()
    def delete_flight_item(self):
        selected = self.flight_tree.selection()
        if not selected:
            messagebox.showwarning("انتخاب", "لطفاً یک ردیف را انتخاب کنید.")
            return
        idx = int(selected[0])
        if messagebox.askyesno("تأیید", "آیا از حذف این آیتم اطمینان دارید؟"):
            del self.flight_data[idx]
            self.refresh_flight_tree()
    def refresh_hotel_tree(self):
        for item in self.hotel_tree.get_children():
            self.hotel_tree.delete(item)
        if not self.hotel_data:
            self.empty_label_hotel.place(relx=0.5, rely=0.5, anchor="center")
        else:
            self.empty_label_hotel.place_forget()
            for idx, htl in enumerate(self.hotel_data):
                values = (htl.get("contract", ""), htl.get("passenger", ""), htl.get("hotel", ""), htl.get("room", ""), str(htl.get("pax", 1)), htl.get("date", ""), htl.get("notes", ""), to_persian_digits(f"{htl.get('debt', 0):,}"), to_persian_digits(f"{htl.get('credit', 0):,}"), to_persian_digits(f"{htl.get('balance', 0):,}"))
                self.hotel_tree.insert("", "end", iid=idx, values=values)
    def add_hotel_item(self):
        dialog = HotelEditDialog(self)
        self.wait_window(dialog)
        if dialog.result:
            self.hotel_data.append(dialog.result)
            self.refresh_hotel_tree()
    def edit_hotel_item(self, event=None):
        selected = self.hotel_tree.selection()
        if not selected:
            messagebox.showwarning("انتخاب", "لطفاً یک ردیف را انتخاب کنید.")
            return
        idx = int(selected[0])
        old_data = self.hotel_data[idx]
        dialog = HotelEditDialog(self, old_data)
        self.wait_window(dialog)
        if dialog.result:
            self.hotel_data[idx] = dialog.result
            self.refresh_hotel_tree()
    def delete_hotel_item(self):
        selected = self.hotel_tree.selection()
        if not selected:
            messagebox.showwarning("انتخاب", "لطفاً یک ردیف را انتخاب کنید.")
            return
        idx = int(selected[0])
        if messagebox.askyesno("تأیید", "آیا از حذف این آیتم اطمینان دارید؟"):
            del self.hotel_data[idx]
            self.refresh_hotel_tree()
    def create_step_3(self):
        frame = ctk.CTkFrame(self.main_container, corner_radius=20)
        self.frames.append(frame)
        title = to_persian_digits("مرحله ۳ از ۵: ویرایش اطلاعات (خدمات، پرواز، هتل)")
        ctk.CTkLabel(frame, text=title, font=self.font_farsi_header).pack(pady=20)
        tabview = ctk.CTkTabview(frame, segmented_button_fg_color=COLOR_PRIMARY)
        tabview._segmented_button.configure(font=self.font_farsi)
        tabview.pack(pady=10, padx=20, fill="both", expand=True)
        tab_service = tabview.add("خدمات")
        self.create_service_tab(tab_service)
        tab_flight = tabview.add("پرواز")
        self.create_flight_tab(tab_flight)
        tab_hotel = tabview.add("هتل")
        self.create_hotel_tab(tab_hotel)
        ctk.CTkLabel(frame, text="می‌توانید ردیف‌ها را ویرایش کنید. برای افزودن/ویرایش روی دکمه‌ها کلیک کنید.", font=self.font_farsi, text_color="gray").pack(pady=5)
    def create_service_tab(self, parent):
        tree_frame = ctk.CTkFrame(parent)
        tree_frame.pack(pady=10, padx=10, fill="both", expand=True)
        columns = ("نوع خدمت", "نام مسافر", "شماره قرارداد", "تاریخ", "توضیحات", "بدهکار", "بستانکار", "مانده")
        self.service_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
        for col in columns:
            self.service_tree.heading(col, text=col)
            self.service_tree.column(col, width=100, anchor="center")
        self.service_tree.column("نوع خدمت", width=120)
        self.service_tree.column("نام مسافر", width=150)
        self.service_tree.column("شماره قرارداد", width=100)
        self.service_tree.column("تاریخ", width=100)
        self.service_tree.column("توضیحات", width=150)
        self.service_tree.column("بدهکار", width=120)
        self.service_tree.column("بستانکار", width=120)
        self.service_tree.column("مانده", width=120)
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.service_tree.yview)
        self.service_tree.configure(yscrollcommand=vsb.set)
        self.service_tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        self.empty_label_service = ctk.CTkLabel(tree_frame, text="هیچ داده‌ای یافت نشد. برای افزودن از دکمه '➕ افزودن' استفاده کنید.", font=self.font_farsi, text_color="gray")
        self.service_tree.bind("<Double-1>", self.edit_service_item)
        btn_frame = ctk.CTkFrame(parent)
        btn_frame.pack(pady=10)
        ctk.CTkButton(btn_frame, text="➕ افزودن", command=self.add_service_item, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="✏ ویرایش", command=self.edit_service_item, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="🗑 حذف", command=self.delete_service_item, font=self.font_farsi, fg_color=COLOR_ACCENT, hover_color="#d45a1c", width=120).pack(side="left", padx=5)
        self.refresh_service_tree()
    def create_flight_tab(self, parent):
        tree_frame = ctk.CTkFrame(parent)
        tree_frame.pack(pady=10, padx=10, fill="both", expand=True)
        columns = ("شماره قرارداد", "نام مسافر", "مسیر", "شماره بلیط", "تاریخ پرواز", "توضیحات", "بدهکار", "بستانکار", "مانده")
        self.flight_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
        for col in columns:
            self.flight_tree.heading(col, text=col)
            self.flight_tree.column(col, width=100, anchor="center")
        self.flight_tree.column("شماره قرارداد", width=100)
        self.flight_tree.column("نام مسافر", width=150)
        self.flight_tree.column("مسیر", width=150)
        self.flight_tree.column("شماره بلیط", width=120)
        self.flight_tree.column("تاریخ پرواز", width=100)
        self.flight_tree.column("توضیحات", width=150)
        self.flight_tree.column("بدهکار", width=120)
        self.flight_tree.column("بستانکار", width=120)
        self.flight_tree.column("مانده", width=120)
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.flight_tree.yview)
        self.flight_tree.configure(yscrollcommand=vsb.set)
        self.flight_tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        self.empty_label_flight = ctk.CTkLabel(tree_frame, text="هیچ داده‌ای یافت نشد. برای افزودن از دکمه '➕ افزودن' استفاده کنید.", font=self.font_farsi, text_color="gray")
        self.flight_tree.bind("<Double-1>", self.edit_flight_item)
        btn_frame = ctk.CTkFrame(parent)
        btn_frame.pack(pady=10)
        ctk.CTkButton(btn_frame, text="➕ افزودن", command=self.add_flight_item, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="✏ ویرایش", command=self.edit_flight_item, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="🗑 حذف", command=self.delete_flight_item, font=self.font_farsi, fg_color=COLOR_ACCENT, hover_color="#d45a1c", width=120).pack(side="left", padx=5)
        self.refresh_flight_tree()
    def create_hotel_tab(self, parent):
        tree_frame = ctk.CTkFrame(parent)
        tree_frame.pack(pady=10, padx=10, fill="both", expand=True)
        columns = ("شماره قرارداد", "نام مسافر", "هتل", "نوع اتاق", "تعداد نفرات", "تاریخ", "توضیحات", "بدهکار", "بستانکار", "مانده")
        self.hotel_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
        for col in columns:
            self.hotel_tree.heading(col, text=col)
            self.hotel_tree.column(col, width=100, anchor="center")
        self.hotel_tree.column("شماره قرارداد", width=100)
        self.hotel_tree.column("نام مسافر", width=150)
        self.hotel_tree.column("هتل", width=150)
        self.hotel_tree.column("نوع اتاق", width=120)
        self.hotel_tree.column("تعداد نفرات", width=90)
        self.hotel_tree.column("تاریخ", width=150)
        self.hotel_tree.column("توضیحات", width=150)
        self.hotel_tree.column("بدهکار", width=120)
        self.hotel_tree.column("بستانکار", width=120)
        self.hotel_tree.column("مانده", width=120)
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.hotel_tree.yview)
        self.hotel_tree.configure(yscrollcommand=vsb.set)
        self.hotel_tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        self.empty_label_hotel = ctk.CTkLabel(tree_frame, text="هیچ داده‌ای یافت نشد. برای افزودن از دکمه '➕ افزودن' استفاده کنید.", font=self.font_farsi, text_color="gray")
        self.hotel_tree.bind("<Double-1>", self.edit_hotel_item)
        btn_frame = ctk.CTkFrame(parent)
        btn_frame.pack(pady=10)
        ctk.CTkButton(btn_frame, text="➕ افزودن", command=self.add_hotel_item, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="✏ ویرایش", command=self.edit_hotel_item, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=120).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text="🗑 حذف", command=self.delete_hotel_item, font=self.font_farsi, fg_color=COLOR_ACCENT, hover_color="#d45a1c", width=120).pack(side="left", padx=5)
        self.refresh_hotel_tree()
    def create_step_4(self):
        frame = ctk.CTkFrame(self.main_container, corner_radius=20)
        self.frames.append(frame)
        title = to_persian_digits("مرحله ۴ از ۵: دوره صورتحساب")
        ctk.CTkLabel(frame, text=title, font=self.font_farsi_header).pack(pady=20)
        period_frame = ctk.CTkFrame(frame, corner_radius=15)
        period_frame.pack(pady=20, padx=20, fill="x")
        label_period = ctk.CTkLabel(period_frame, text="دوره صورتحساب", font=self.font_farsi)
        label_period.pack(side="right", padx=5, pady=10)
        self.btn_pick_period = ctk.CTkButton(period_frame, text="انتخاب بازه", command=self.pick_period, font=self.font_farsi, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, width=140)
        self.btn_pick_period.pack(side="right", padx=5, pady=10)
        self.period_label = ctk.CTkLabel(period_frame, textvariable=self.period_display, font=self.font_farsi, text_color="gray")
        self.period_label.pack(side="right", padx=5, pady=10)
        ctk.CTkLabel(frame, text="لطفاً بازه صورتحساب را انتخاب کنید.", font=self.font_farsi, text_color="gray").pack(pady=20)
    def pick_period(self):
        dialog = DateRangePickerDialog(self)
        self.wait_window(dialog)
        if dialog.result:
            res = dialog.result
            self.start_day.set(str(res['start_day']))
            self.end_day.set(str(res['end_day']))
            self.start_month.set(res['start_month'])
            self.start_year.set(str(res['start_year']))
            self.end_month.set(res['end_month'])
            self.end_year.set(str(res['end_year']))
            self.period_range.set(res['period_range'])
            self.period_display.set(res['period_range'])
    def create_step_5(self):
        frame = ctk.CTkFrame(self.main_container, corner_radius=20)
        self.frames.append(frame)
        title = to_persian_digits("مرحله ۵ از ۵: خلاصه و تایید")
        self.step6_title = ctk.CTkLabel(frame, text=title, font=self.font_farsi_header)
        self.step6_title.pack(pady=20)
        self.summary_frame = ctk.CTkFrame(frame, corner_radius=15)
        self.summary_frame.pack(pady=20, padx=20, fill="both", expand=True)
        self.summary_text = ctk.CTkTextbox(self.summary_frame, height=280, font=self.font_farsi, wrap="word", corner_radius=10)
        self.summary_text.pack(pady=10, padx=10, fill="both", expand=True)
        self.status_frame = ctk.CTkFrame(self.summary_frame, fg_color="transparent")
        self.status_frame.pack(pady=10, fill="x")
        self.icon_label = ctk.CTkLabel(self.status_frame, text="", font=("B Nazanin", 60), text_color="green")
        self.icon_label.pack(side="left", padx=10)
        self.message_label = ctk.CTkLabel(self.status_frame, text="", font=self.font_farsi, text_color="white")
        self.message_label.pack(side="left")
        self.status_frame.pack_forget()
        self.success_frame = ctk.CTkFrame(frame, corner_radius=15, fg_color="transparent")
        self.success_frame.pack(pady=20, padx=20, fill="both", expand=True)
        self.success_frame.pack_forget()
        self.success_icon = ctk.CTkLabel(self.success_frame, text="✓", font=("B Nazanin", 120), text_color="green")
        self.success_icon.pack(pady=10)
        self.success_details = ctk.CTkLabel(self.success_frame, text="", font=self.font_farsi, text_color="white")
        self.success_details.pack(pady=10)
        ctk.CTkLabel(self.success_frame, text="فایل‌ها با موفقیت ساخته شدند!", font=self.font_farsi_header, text_color="green").pack(pady=10)
        self.restart_btn = ctk.CTkButton(self.success_frame, text="شروع مجدد", command=self.restart, font=self.font_farsi_bold, width=220, fg_color=COLOR_PRIMARY, hover_color=COLOR_ACCENT, corner_radius=12)
        self.restart_btn.pack(pady=20)
        ctk.CTkLabel(frame, text="لطفاً اطلاعات را بررسی کرده و اقدام نهایی را انجام دهید.", font=self.font_farsi, text_color="gray").pack(pady=10)
    def show_step(self, step):
        for i, frame in enumerate(self.frames):
            if i == step:
                frame.pack(fill="both", expand=True)
            else:
                frame.pack_forget()
        self.current_step = step
        total_steps = len(self.frames) - 1
        if step == 0:
            self.btn_home.configure(state="disabled")
            self.btn_prev.configure(state="disabled")
            self.btn_next.configure(state="disabled")
            self.btn_action.configure(state="disabled")
            self.progress_bar.set(0)
        else:
            self.btn_home.configure(state="normal")
            self.btn_prev.configure(state="normal" if step > 1 else "disabled")
            self.btn_next.configure(state="normal" if step < total_steps else "disabled")
            self.progress_bar.set(step / total_steps)
            if step == total_steps:
                self.update_action_button()
                self.update_summary()
                self.step6_title.pack(pady=20)
                self.summary_frame.pack(pady=20, padx=20, fill="both", expand=True)
                self.success_frame.pack_forget()
            else:
                self.btn_action.configure(state="disabled")
                self.status_frame.pack_forget()
                self.success_frame.pack_forget()
        if not (self.can_issue or self.can_submit) and step > 0:
            self.show_step(0)
    def next_step(self):
        if self.current_step < len(self.frames) - 1:
            if self.current_step == 1:
                self.hide_field_errors()
                valid = True
                #if not self.payer.get():
                    #self.show_field_error("payer", "طرف حساب نمی‌تواند خالی باشد.")
                    #valid = False
                #if not self.generated_letter_number.get():
                    #self.show_field_error("payer", "شماره نامه تولید نشده است. روی دکمه 'تولید شماره' کلیک کنید.")
                    #valid = False
                if not self.year_letter.get() or not self.month_letter.get() or not self.day_letter.get():
                    self.show_field_error("date", "تاریخ را به طور کامل انتخاب کنید.")
                    valid = False
                else:
                    self.letter_date.set(f"{english_number(self.year_letter.get())}/{english_number(self.month_letter.get())}/{english_number(self.day_letter.get())}")
                if not valid:
                    return
                self.show_step(self.current_step + 1)
            elif self.current_step == 2:
                self.selected_files.clear()
                for item in self.file_items:
                    self.selected_files.append({'path': item.file_path, 'type': item.type_combo.get()})
                self.load_service_data_from_files()
                self.load_flight_data_from_files()
                self.load_hotel_data_from_files()
                self.refresh_service_tree()
                self.refresh_flight_tree()
                self.refresh_hotel_tree()
                self.show_step(self.current_step + 1)
            elif self.current_step == 3:
                self.show_step(self.current_step + 1)
            elif self.current_step == 4:
                if not self.start_day.get() or not self.end_day.get():
                    self.show_error("لطفاً بازه صورتحساب را انتخاب کنید.")
                    return
                self.show_step(self.current_step + 1)
            else:
                self.show_step(self.current_step + 1)
    def prev_step(self):
        if self.current_step > 0:
            self.show_step(self.current_step - 1)
    def show_error(self, message):
        messagebox.showerror("خطا", to_persian_digits(message))
    def update_summary(self):
        total_flight = sum(item['balance'] for item in self.flight_data) if self.flight_data else 0
        total_hotel = sum(item['balance'] for item in self.hotel_data) if self.hotel_data else 0
        total_service = sum(item['balance'] for item in self.service_data) if self.service_data else 0
        total_debt = total_flight + total_hotel + total_service
        summary = f"""
خلاصه اطلاعات وارد شده:
تعداد فایل‌های انتخاب شده: {len(self.selected_files)}
طرف حساب: {self.payer.get()}
شماره نامه: {self.generated_letter_number.get()}
تاریخ نامه: {self.year_letter.get()}/{self.month_letter.get()}/{self.day_letter.get()}
دوره صورتحساب: {self.period_range.get()}
تفکیک بدهی:
- پرواز: {to_persian_digits(f'{total_flight:,}')} ریال
- هتل: {to_persian_digits(f'{total_hotel:,}')} ریال
- سایر خدمات: {to_persian_digits(f'{total_service:,}')} ریال
جمع کل بدهی: {to_persian_digits(f'{total_debt:,}')} ریال
"""
        self.summary_text.delete("0.0", "end")
        self.summary_text.insert("0.0", summary)
    def restart(self):
        self.excel_path.set("")
        self.payer.set("")
        self.generated_letter_number.set("")
        self.set_default_date()
        self.start_day.set("")
        self.end_day.set("")
        self.start_month.set("")
        self.start_year.set("")
        self.end_month.set("")
        self.end_year.set("")
        self.period_range.set("")
        self.period_display.set("هنوز بازه‌ای انتخاب نشده")
        self.service_data = []
        self.flight_data = []
        self.hotel_data = []
        self.clear_files()
        self.show_step(0)
        self.hide_field_errors()

    def generate_invoice_pdf(self, output_path, context):
        def persian_number_filter(value):
            return to_persian_digits(f"{value:,}")

        def format_balance(value):
            if value < 0:
                return f"({to_persian_digits(f'{abs(value):,}')})"
            else:
                return to_persian_digits(f"{value:,}")

        from jinja2 import Environment
        env = Environment()
        env.filters['persian_number'] = persian_number_filter
        env.filters['format_balance'] = format_balance

        # دریافت قالب از سرور
        try:
            response = requests.get(f"http://{self.api_server}:5000/get-active-template", timeout=3)
            if response.status_code == 200:
                template_data = response.json()
                html_template = template_data.get('html', '')
                css_style = template_data.get('css', '')
            else:
                raise Exception("سرور پاسخ نداد")
        except Exception:
            # Fallback به قالب داخلی (مشابه کد قبلی خودتان)
            html_template = self.DEFAULT_HTML_TEMPLATE
            css_style = self.DEFAULT_CSS_STYLE

        rows = context['rows']
        chunk_size = 8
        chunks = [rows[i:i+chunk_size] for i in range(0, len(rows), chunk_size)]
        page_data = []
        for idx, chunk in enumerate(chunks):
            page_total_debt = sum(r.get('debt', 0) for r in chunk)
            page_total_credit = sum(r.get('credit', 0) for r in chunk)
            page_total_balance = sum(r.get('balance', 0) for r in chunk)
            page_data.append({
                'page_num': idx+1,
                'rows': chunk,
                'total_debt': page_total_debt,
                'total_credit': page_total_credit,
                'total_balance': page_total_balance,
                'is_last': (idx == len(chunks)-1)
            })

        grand_total_debt = context.get('total_debt', 0)
        grand_total_credit = context.get('total_credit', 0)
        grand_total_balance = context.get('total_balance', 0)

        cumulative_debt = 0
        cumulative_credit = 0
        cumulative_balance = 0
        for page in page_data:
            cumulative_debt += page['total_debt']
            cumulative_credit += page['total_credit']
            cumulative_balance += page['total_balance']
            page['cumulative_debt'] = cumulative_debt
            page['cumulative_credit'] = cumulative_credit
            page['cumulative_balance'] = cumulative_balance

        template_context = {
            'pages': page_data,
            'seller': context.get('seller', {}),
            'buyer': context.get('buyer', {}),
            'service_type': context.get('service_type', ''),
            'invoice_number': context.get('invoice_number', ''),
            'invoice_date': context.get('invoice_date', ''),
            'grand_total_debt': grand_total_debt,
            'grand_total_credit': grand_total_credit,
            'grand_total_balance': grand_total_balance
        }

        full_html = f"""
        <!DOCTYPE html>
        <html dir="rtl" lang="fa">
        <head><meta charset="UTF-8"><style>{css_style}</style></head>
        <body>
        {html_template}
        </body>
        </html>
        """
        from jinja2 import Template
        rendered_html = Template(full_html).render(**template_context)

        html_path = output_path.replace('.pdf', '.html')
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(rendered_html)

        edge_paths = [r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
                      r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"]
        edge_exe = next((p for p in edge_paths if os.path.exists(p)), None)
        if edge_exe:
            subprocess.run([edge_exe, "--headless", f"--print-to-pdf={output_path}", "--no-pdf-header-footer", html_path], check=True, timeout=60)
        else:
            webbrowser.open(html_path)
            raise Exception("Microsoft Edge یافت نشد")
        
    def generate_direct(self):
        if not self.can_issue and not self.can_approve:
            self.show_error("شما مجوز تولید مستقیم را ندارید.")
            return
        if not self.payer.get():
            self.show_error("طرف حساب نمی‌تواند خالی باشد.")
            self.show_step(1)
            return
        if not self.generated_letter_number.get():
            self.show_error("شماره نامه تولید نشده است.")
            self.show_step(1)
            return
        if not self.year_letter.get() or not self.month_letter.get() or not self.day_letter.get():
            self.show_error("تاریخ نامه را کامل انتخاب کنید.")
            self.show_step(1)
            return
        if not self.start_day.get() or not self.end_day.get():
            self.show_error("لطفاً بازه صورتحساب را انتخاب کنید.")
            self.show_step(4)
            return
        self.status_frame.pack(pady=10, fill="x")
        self.icon_label.configure(text="", font=("B Nazanin", 1))
        self.message_label.configure(text="در حال پردازش...", text_color="yellow")
        self.update()
        try:
            output_dir = filedialog.askdirectory(title="انتخاب پوشه برای ذخیره فایل‌ها")
            if not output_dir:
                return
            base_dir = output_dir
            today = jdatetime.date.today()
            date_folder = today.strftime("%Y-%m-%d")
            selected_payer = self.payer.get()
            payer_code = ""
            payer_name = selected_payer
            payer_name_for_address = selected_payer
            subsidiary_name = ""
            payer_info = None
            if " - " in selected_payer:
                parts = selected_payer.split(" - ", 1)
                payer_code = parts[0].strip()
                payer_name = parts[1].strip()
                payer_info = next((p for p in self.payers_list if p['code'] == payer_code), None)
                if payer_info:
                    if payer_info.get('parent_code'):
                        parent_info = next((p for p in self.payers_list if p['code'] == payer_info['parent_code']), None)
                        if parent_info:
                            payer_name_for_address = parent_info['name']
                            subsidiary_name = payer_info['name']
                        else:
                            payer_name_for_address = payer_name
                            subsidiary_name = payer_name
                    else:
                        payer_name_for_address = payer_name
                        subsidiary_name = payer_name
                else:
                    payer_name_for_address = selected_payer
                    subsidiary_name = selected_payer
            else:
                payer_name_for_address = selected_payer
                subsidiary_name = selected_payer
            safe_payer = safe_filename(payer_name)
            final_dir = os.path.join(base_dir, date_folder, safe_payer)
            os.makedirs(final_dir, exist_ok=True)
            company_settings = self.db.get_company_settings() or {}
            buyer_info = {
                'name': payer_name,
                'national_id': payer_info.get('national_id', '') if payer_info else '',
                'economic_code': payer_info.get('economic_code', '') if payer_info else '',
                'address': payer_info.get('address', '') if payer_info else '',
                'postal_code': payer_info.get('postal_code', '') if payer_info else '',
                'phone': payer_info.get('phone', '') if payer_info else ''
            }
            service_packages = []
            if self.flight_data:
                flight_total = sum(item['balance'] for item in self.flight_data)
                flight_rows = [{'contract': flt.get('contract',''), 'passenger': flt.get('passenger',''), 'description': flt.get('route',''), 'date': flt.get('date',''), 'notes': flt.get('notes',''), 'debt': flt.get('debt',0), 'credit': flt.get('credit',0), 'balance': flt.get('balance',0)} for flt in self.flight_data]
                service_packages.append({'type': 'پرواز', 'total_balance': flight_total, 'rows': flight_rows})
            if self.hotel_data:
                hotel_total = sum(item['balance'] for item in self.hotel_data)
                hotel_rows = [{'contract': htl.get('contract',''), 'passenger': htl.get('passenger',''), 'description': htl.get('hotel','') + " - " + htl.get('room',''), 'date': htl.get('date',''), 'notes': htl.get('notes','') + (f" (تعداد: {htl.get('pax',1)})" if htl.get('notes') else f"تعداد: {htl.get('pax',1)}"), 'debt': htl.get('debt',0), 'credit': htl.get('credit',0), 'balance': htl.get('balance',0)} for htl in self.hotel_data]
                service_packages.append({'type': 'هتل', 'total_balance': hotel_total, 'rows': hotel_rows})
            if self.service_data:
                service_by_type = {}
                service_rows_by_type = {}
                for item in self.service_data:
                    svc_type = item.get('type', 'سایر')
                    if svc_type not in service_by_type:
                        service_by_type[svc_type] = 0
                        service_rows_by_type[svc_type] = []
                    service_by_type[svc_type] += item['balance']
                    service_rows_by_type[svc_type].append({'contract': item.get('contract',''), 'passenger': item.get('passenger',''), 'description': svc_type, 'date': item.get('date',''), 'notes': item.get('notes',''), 'debt': item.get('debt',0), 'credit': item.get('credit',0), 'balance': item.get('balance',0)})
                for svc_type, total in service_by_type.items():
                    service_packages.append({'type': svc_type, 'total_balance': total, 'rows': service_rows_by_type[svc_type]})
            if not service_packages:
                self.show_error("هیچ داده‌ای برای تولید وجود ندارد.")
                return

            # ========== اصلاح شماره‌گذاری ==========
            if payer_code:
                letter_numbers, raw_numbers = self.generate_multiple_letter_numbers(payer_code, len(service_packages))
            else:
                letter_numbers = [self.generated_letter_number.get()] * len(service_packages)
                raw_numbers = []
            # ====================================

            try:
                for idx, package in enumerate(service_packages):
                    current_letter = letter_numbers[idx]
                    safe_letter = safe_filename(current_letter)
                    total_balance = package['total_balance']
                    if total_balance > 0:
                        amount_to_record = total_balance
                        self.save_to_history(current_letter, payer_name, payer_code, package['type'], amount_to_record, today.strftime("%Y/%m/%d"), "")
                        is_subsidiary = bool(payer_info and payer_info.get('parent_code'))
                        template_path = resource_path("template - subdevision.docx") if is_subsidiary else resource_path("template.docx")
                        ensure_subsidiary_template_exists(template_path) if is_subsidiary else ensure_simple_template_exists(template_path)
                        context_word = {
                            'letter_date': self.reverse_persian_text(self.letter_date.get()),
                            'letter_number': self.reverse_persian_text(current_letter),
                            'payer': payer_name_for_address,
                            'subsidiary': subsidiary_name,
                            'service_type': package['type'],
                            'period_range': self.period_range.get(),
                            'total_debt': f"{amount_to_record:,}",
                            'total_credit': "",
                            'bank_name1': company_settings.get('bank_name1', ''),
                            'account_number1': company_settings.get('account_number1', ''),
                            'shaba_number1': company_settings.get('shaba_number1', ''),
                            'bank_name2': company_settings.get('bank_name2', ''),
                            'account_number2': company_settings.get('account_number2', ''),
                            'shaba_number2': company_settings.get('shaba_number2', ''),
                            'company_name': company_settings.get('company_name', ''),
                            'manager_name': company_settings.get('manager_name', ''),
                            'manager_position': company_settings.get('manager_position', '')
                        }
                    else:
                        amount_to_record = abs(total_balance)
                        self.db.add_credit(payer_code, amount_to_record, today.strftime("%Y/%m/%d"), f"{package['type']} - {self.period_range.get()}", current_letter)
                        template_path = resource_path("template_creditor.docx")
                        ensure_creditor_template_exists(template_path)
                        context_word = {
                            'letter_date': self.reverse_persian_text(self.letter_date.get()),
                            'letter_number': self.reverse_persian_text(current_letter),
                            'payer': payer_name_for_address,
                            'subsidiary': subsidiary_name,
                            'service_type': package['type'],
                            'period_range': self.period_range.get(),
                            'total_credit': f"{amount_to_record:,}",
                            'total_debt': "",
                            'bank_name1': company_settings.get('bank_name1', ''),
                            'account_number1': company_settings.get('account_number1', ''),
                            'shaba_number1': company_settings.get('shaba_number1', ''),
                            'bank_name2': company_settings.get('bank_name2', ''),
                            'account_number2': company_settings.get('account_number2', ''),
                            'shaba_number2': company_settings.get('shaba_number2', ''),
                            'company_name': company_settings.get('company_name', ''),
                            'manager_name': company_settings.get('manager_name', ''),
                            'manager_position': company_settings.get('manager_position', '')
                        }
                    df_excel = pd.DataFrame(package['rows'])
                    column_mapping = {'contract': 'شماره قرارداد', 'passenger': 'نام مسافر', 'description': 'شرح خدمات', 'date': 'تاریخ', 'notes': 'توضیحات', 'debt': 'بدهکار', 'credit': 'بستانکار', 'balance': 'مانده'}
                    df_excel = df_excel.rename(columns=column_mapping)
                    df_excel.insert(0, 'ردیف', range(1, len(df_excel)+1))
                    excel_path = os.path.join(final_dir, f"صورت حساب {package['type']} {safe_payer} {safe_letter}.xlsx")
                    df_excel.to_excel(excel_path, index=False)
                    word_output = os.path.join(final_dir, f"اعلامیه {package['type']} {safe_payer} {safe_letter}.docx")
                    generate_word(template_path, word_output, context_word)
                    pdf_context = {
                        'seller': {
                            'company_name': company_settings.get('company_name', ''),
                            'national_id': company_settings.get('national_id', ''),
                            'economic_code': company_settings.get('economic_code', ''),
                            'address': company_settings.get('address', ''),
                            'postal_code': company_settings.get('postal_code', ''),
                            'phone': company_settings.get('phone', ''),
                            'manager_name': company_settings.get('manager_name', ''),
                            'manager_position': company_settings.get('manager_position', ''),
                            'signature_path': company_settings.get('signature_path', '') if company_settings.get('signature_path') and os.path.exists(company_settings['signature_path']) else '',
                            'stamp_path': company_settings.get('stamp_path', '') if company_settings.get('stamp_path') and os.path.exists(company_settings['stamp_path']) else ''
                        },
                        'buyer': buyer_info,
                        'invoice_number': current_letter,
                        'invoice_date': self.letter_date.get(),
                        'service_type': package['type'],
                        'period_range': self.period_range.get(),
                        'rows': package['rows'],
                        'total_debt': amount_to_record if total_balance > 0 else 0,
                        'total_credit': amount_to_record if total_balance < 0 else 0,
                        'total_balance': total_balance
                    }
                    pdf_output = os.path.join(final_dir, f"فاکتور {package['type']} {safe_payer} {safe_letter}.pdf")
                    self.generate_invoice_pdf(pdf_output, pdf_context)
                self.success_details.configure(text=f"تعداد نامه‌های تولید شده: {len(service_packages)}")
                self.step6_title.pack_forget()
                self.summary_frame.pack_forget()
                self.status_frame.pack_forget()
                self.success_frame.pack(pady=20, padx=20, fill="both", expand=True)
                self.btn_prev.configure(state="disabled")
                self.btn_next.configure(state="disabled")
                self.btn_action.configure(state="disabled")
            except Exception as e:
                raise e
        except Exception as e:
            logging.error(f"خطا در تولید: {traceback.format_exc()}")
            self.message_label.configure(text=f"خطا: {str(e)[:50]}...", text_color="red")
            self.animate_tick(10, step=10, target=80, icon="✗", color="red")
            messagebox.showerror("خطا", f"مشکلی در تولید پیش آمد:\n{str(e)}")
    def submit_for_approval(self):
        if not self.can_submit:
            self.show_error("شما مجوز ارسال درخواست تأیید را ندارید.")
            return
        if not self.payer.get():
            self.show_error("طرف حساب نمی‌تواند خالی باشد.")
            self.show_step(1)
            return
        if not self.generated_letter_number.get():
            self.show_error("شماره نامه تولید نشده است.")
            self.show_step(1)
            return
        if not self.year_letter.get() or not self.month_letter.get() or not self.day_letter.get():
            self.show_error("تاریخ نامه را کامل انتخاب کنید.")
            self.show_step(1)
            return
        if not self.start_day.get() or not self.end_day.get():
            self.show_error("لطفاً بازه صورتحساب را انتخاب کنید.")
            self.show_step(4)
            return
        selected_payer = self.payer.get()
        payer_code = ""
        payer_name = selected_payer
        if " - " in selected_payer:
            parts = selected_payer.split(" - ", 1)
            payer_code = parts[0].strip()
            payer_name = parts[1].strip()
        else:
            payer_code = ""
            payer_name = selected_payer
        try:
            req_id = self.db.create_approval_request(
                requester_id=self.user_id,
                requester_name=self.username,
                payer_code=payer_code,
                payer_name=payer_name,
                letter_number=self.generated_letter_number.get(),
                letter_date=self.letter_date.get(),
                start_day=int(english_number(self.start_day.get())),
                end_day=int(english_number(self.end_day.get())),
                month=self.start_month.get(),
                year=int(english_number(self.start_year.get())),
                period_range=self.period_range.get(),
                service_data=self.service_data,
                flight_data=self.flight_data,
                hotel_data=self.hotel_data
            )
            messagebox.showinfo("موفقیت", f"درخواست شما با شماره {req_id} برای تأیید ارسال شد.\nپس از تأیید مدیر، می‌توانید از بخش گردش کار در مدیریت، فایل نهایی را صادر کنید.")
            self.restart()
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در ارسال درخواست:\n{str(e)}")

    def generate_from_approval(self, request_data, output_root, approver_id):
        import json
        service_data = json.loads(request_data['service_data_json'])
        flight_data = json.loads(request_data['flight_data_json'])
        hotel_data = json.loads(request_data['hotel_data_json'])
        payer_code = request_data['payer_code']
        payer_name = request_data['payer_name']
        letter_number = request_data['letter_number']
        letter_date = request_data['letter_date']
        start_day = request_data['start_day']
        end_day = request_data['end_day']
        month = request_data['month']
        year = request_data['year']
        period_range = request_data['period_range']
        safe_payer = safe_filename(payer_name)
        final_dir = os.path.join(output_root, safe_payer)
        os.makedirs(final_dir, exist_ok=True)
        company_settings = self.db.get_company_settings() or {}
        buyer_info = {}
        payer_info = next((p for p in self.payers_list if p['code'] == payer_code), None)
        if payer_info:
            buyer_info = {
                'name': payer_name,
                'national_id': payer_info.get('national_id', ''),
                'economic_code': payer_info.get('economic_code', ''),
                'address': payer_info.get('address', ''),
                'postal_code': payer_info.get('postal_code', ''),
                'phone': payer_info.get('phone', '')
            }
        else:
            buyer_info = {'name': payer_name, 'national_id': '', 'economic_code': '', 'address': '', 'postal_code': '', 'phone': ''}
        service_packages = []
        if flight_data:
            flight_total = sum(item['balance'] for item in flight_data)
            flight_rows = [{'contract': flt.get('contract',''), 'passenger': flt.get('passenger',''), 'description': flt.get('route',''), 'date': flt.get('date',''), 'notes': flt.get('notes',''), 'debt': flt.get('debt',0), 'credit': flt.get('credit',0), 'balance': flt.get('balance',0)} for flt in flight_data]
            service_packages.append({'type': 'پرواز', 'total_balance': flight_total, 'rows': flight_rows})
        if hotel_data:
            hotel_total = sum(item['balance'] for item in hotel_data)
            hotel_rows = [{'contract': htl.get('contract',''), 'passenger': htl.get('passenger',''), 'description': htl.get('hotel','') + " - " + htl.get('room',''), 'date': htl.get('date',''), 'notes': htl.get('notes','') + (f" (تعداد: {htl.get('pax',1)})" if htl.get('notes') else f"تعداد: {htl.get('pax',1)}"), 'debt': htl.get('debt',0), 'credit': htl.get('credit',0), 'balance': htl.get('balance',0)} for htl in hotel_data]
            service_packages.append({'type': 'هتل', 'total_balance': hotel_total, 'rows': hotel_rows})
        if service_data:
            service_by_type = {}
            service_rows_by_type = {}
            for item in service_data:
                svc_type = item.get('type', 'سایر')
                if svc_type not in service_by_type:
                    service_by_type[svc_type] = 0
                    service_rows_by_type[svc_type] = []
                service_by_type[svc_type] += item['balance']
                service_rows_by_type[svc_type].append({'contract': item.get('contract',''), 'passenger': item.get('passenger',''), 'description': svc_type, 'date': item.get('date',''), 'notes': item.get('notes',''), 'debt': item.get('debt',0), 'credit': item.get('credit',0), 'balance': item.get('balance',0)})
            for svc_type, total in service_by_type.items():
                service_packages.append({'type': svc_type, 'total_balance': total, 'rows': service_rows_by_type[svc_type]})
        if not service_packages:
            raise Exception("هیچ داده‌ای برای تولید وجود ندارد.")
        if payer_code:
            letter_numbers, raw_numbers = self.generate_multiple_letter_numbers(payer_code, len(service_packages))
        else:
            letter_numbers = [letter_number] * len(service_packages)
            raw_numbers = []
        final_files = []
        today = jdatetime.date.today()
        for idx, package in enumerate(service_packages):
            current_letter = letter_numbers[idx]
            safe_letter = safe_filename(current_letter)
            total_balance = package['total_balance']
            if total_balance > 0:
                amount_to_record = total_balance
                self.save_to_history(current_letter, payer_name, payer_code, package['type'], amount_to_record, today.strftime("%Y/%m/%d"), "")
                is_subsidiary = bool(payer_info and payer_info.get('parent_code'))
                template_path = resource_path("template - subdevision.docx") if is_subsidiary else resource_path("template.docx")
                ensure_subsidiary_template_exists(template_path) if is_subsidiary else ensure_simple_template_exists(template_path)
                context_word = {
                    'letter_date': self.reverse_persian_text(letter_date),
                    'letter_number': self.reverse_persian_text(current_letter),
                    'payer': buyer_info['name'],
                    'subsidiary': payer_info.get('name', '') if payer_info else '',
                    'service_type': package['type'],
                    'period_range': period_range,
                    'total_debt': f"{amount_to_record:,}",
                    'total_credit': "",
                    'bank_name1': company_settings.get('bank_name1', ''),
                    'account_number1': company_settings.get('account_number1', ''),
                    'shaba_number1': company_settings.get('shaba_number1', ''),
                    'bank_name2': company_settings.get('bank_name2', ''),
                    'account_number2': company_settings.get('account_number2', ''),
                    'shaba_number2': company_settings.get('shaba_number2', ''),
                    'company_name': company_settings.get('company_name', ''),
                    'manager_name': company_settings.get('manager_name', ''),
                    'manager_position': company_settings.get('manager_position', '')
                }
            else:
                amount_to_record = abs(total_balance)
                self.db.add_credit(payer_code, amount_to_record, today.strftime("%Y/%m/%d"), f"{package['type']} - {period_range}", current_letter)
                template_path = resource_path("template_creditor.docx")
                ensure_creditor_template_exists(template_path)
                context_word = {
                    'letter_date': self.reverse_persian_text(letter_date),
                    'letter_number': self.reverse_persian_text(current_letter),
                    'payer': buyer_info['name'],
                    'subsidiary': payer_info.get('name', '') if payer_info else '',
                    'service_type': package['type'],
                    'period_range': period_range,
                    'total_credit': f"{amount_to_record:,}",
                    'total_debt': "",
                    'bank_name1': company_settings.get('bank_name1', ''),
                    'account_number1': company_settings.get('account_number1', ''),
                    'shaba_number1': company_settings.get('shaba_number1', ''),
                    'bank_name2': company_settings.get('bank_name2', ''),
                    'account_number2': company_settings.get('account_number2', ''),
                    'shaba_number2': company_settings.get('shaba_number2', ''),
                    'company_name': company_settings.get('company_name', ''),
                    'manager_name': company_settings.get('manager_name', ''),
                    'manager_position': company_settings.get('manager_position', '')
                }
            df_excel = pd.DataFrame(package['rows'])
            column_mapping = {'contract': 'شماره قرارداد', 'passenger': 'نام مسافر', 'description': 'شرح خدمات', 'date': 'تاریخ', 'notes': 'توضیحات', 'debt': 'بدهکار', 'credit': 'بستانکار', 'balance': 'مانده'}
            df_excel = df_excel.rename(columns=column_mapping)
            df_excel.insert(0, 'ردیف', range(1, len(df_excel)+1))
            excel_path = os.path.join(final_dir, f"صورت حساب {package['type']} {safe_payer} {safe_letter}.xlsx")
            df_excel.to_excel(excel_path, index=False)
            word_output = os.path.join(final_dir, f"اعلامیه {package['type']} {safe_payer} {safe_letter}.docx")
            generate_word(template_path, word_output, context_word)
            pdf_context = {
                'seller': {
                    'company_name': company_settings.get('company_name', ''),
                    'national_id': company_settings.get('national_id', ''),
                    'economic_code': company_settings.get('economic_code', ''),
                    'address': company_settings.get('address', ''),
                    'postal_code': company_settings.get('postal_code', ''),
                    'phone': company_settings.get('phone', ''),
                    'manager_name': company_settings.get('manager_name', ''),
                    'manager_position': company_settings.get('manager_position', ''),
                    'signature_path': company_settings.get('signature_path', '') if company_settings.get('signature_path') and os.path.exists(company_settings['signature_path']) else '',
                    'stamp_path': company_settings.get('stamp_path', '') if company_settings.get('stamp_path') and os.path.exists(company_settings['stamp_path']) else ''
                },
                'buyer': buyer_info,
                'invoice_number': current_letter,
                'invoice_date': letter_date,
                'service_type': package['type'],
                'period_range': period_range,
                'rows': package['rows'],
                'total_debt': amount_to_record if total_balance > 0 else 0,
                'total_credit': amount_to_record if total_balance < 0 else 0,
                'total_balance': total_balance
            }
            pdf_output = os.path.join(final_dir, f"فاکتور {package['type']} {safe_payer} {safe_letter}.pdf")
            self.generate_invoice_pdf(pdf_output, pdf_context)
            final_files.append({'excel': excel_path, 'word': word_output, 'pdf': pdf_output})
        return final_files
    
    def open_workflow(self):
        if not (self.can_approve or self.can_submit):
            messagebox.showerror("خطا", "شما مجوز دسترسی به این بخش را ندارید.")
            return
        dialog = ManagementDialog(self, self, self.payers_list, self.can_view, self.can_pay,
                                self.can_issue, self.can_delete_history, default_tab="گردش کار")
        self.wait_window(dialog)

    def reverse_persian_text(self, text):
        return '/'.join(reversed(text.split('/'))) if '/' in text else text
    
    def animate_tick(self, size, step=5, target=100, icon="✓", color="green"):
        if size <= target:
            self.icon_label.configure(text=icon, font=("B Nazanin", size), text_color=color)
            self.after(30, lambda: self.animate_tick(size+step, step, target, icon, color))

    def logout(self):
        if messagebox.askyesno("خروج از حساب", "آیا از خروج از حساب کاربری فعلی اطمینان دارید؟"):
            self.hide_main_ui()
            # بازنشانی مجوزها
            self.can_view = self.can_pay = self.can_issue = self.can_manage_users = \
            self.can_manage_settings = self.can_delete_history = self.can_delete_payment = \
            self.can_export = self.can_manage_companies = self.can_submit = self.can_approve = False
            self.current_user = None
            self.user_id = self.username = None
            self.restart()
            self.show_login()
            if self.current_user:
                # به‌روزرسانی مجوزها از کاربر جدید
                (self.user_id, self.username, self.can_view, self.can_pay, self.can_issue,
                self.can_manage_users, self.can_manage_settings, self.can_delete_history,
                self.can_delete_payment, self.can_export, self.can_manage_companies,
                self.can_submit, self.can_approve) = self.current_user
                # بازسازی منوی بالایی با مجوزهای جدید
                self.rebuild_menu()
                # نمایش مجدد UI اصلی
                self.show_main_ui()
                self.update_buttons_by_permission()
                self.show_step(0)
            else:
                self.quit()

    def on_closing(self):
        if messagebox.askyesno("خروج", "آیا مطمئن هستید؟"):
            if messagebox.askyesno("پشتیبان‌گیری", "پشتیبان تهیه شود؟"):
                ext = ".bak"
                filename = generate_backup_filename() + ext
                full_path = os.path.join(self.backup_dir, filename)
                try:
                    self.db_manager.backup_database(full_path)
                    messagebox.showinfo("موفقیت", f"پشتیبان در {full_path} ذخیره شد")
                except Exception as e:
                    messagebox.showerror("خطا", str(e))
            self.destroy()

    def reserve_letter_number(self, payer_code, letter_number_str):
        """شماره نامه را از روی رشته، استخراج و در دیتابیس رزرو می‌کند"""
        # استخراج عدد از انتهای شماره نامه
        match = re.search(r'(\d+)$', letter_number_str)
        if not match:
            return None
        num = int(match.group(1))
        today = jdatetime.date.today()
        year_two_digit = str(today.year % 100).zfill(2)
            # رزرو کردن آن شماره خاص (با فرض اینکه هنوز گرفته نشده)
        with self.db.connect() as conn:
            cursor = conn.cursor()
            # ابتدا بررسی کنیم آیا این شماره در available_letter_numbers هست یا باید از letter_sequences استفاده کنیم
            cursor.execute("SELECT 1 FROM available_letter_numbers WHERE payer_code = ? AND year = ? AND number = ?", (payer_code, year_two_digit, num))
            if cursor.fetchone():
                cursor.execute("DELETE FROM available_letter_numbers WHERE payer_code = ? AND year = ? AND number = ?", (payer_code, year_two_digit, num))
            else:
                # اگر در available نبود، باید last_number را به این عدد برسانیم
                cursor.execute("MERGE INTO letter_sequences AS target USING (VALUES (?, ?, ?)) AS source (payer_code, year, last_number) ON target.payer_code = source.payer_code AND target.year = source.year WHEN MATCHED THEN UPDATE SET last_number = source.last_number WHEN NOT MATCHED THEN INSERT (payer_code, year, last_number) VALUES (source.payer_code, source.year, source.last_number);", (payer_code, year_two_digit, num))
            conn.commit()
        return num
    
    def download_final_file_from_api(self, file_id):
        api_url = f"http://{self.api_server}:5000/download/{file_id}"
        try:
            response = requests.get(api_url, stream=True, timeout=30)
            if response.status_code == 200:
                content_disp = response.headers.get('content-disposition', '')
                filename = "downloaded_file.pdf"
                if 'filename=' in content_disp:
                    filename = content_disp.split('filename=')[1].strip('"')
                save_path = filedialog.asksaveasfilename(initialfile=filename, defaultextension=".pdf")
                if save_path:
                    with open(save_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    messagebox.showinfo("موفقیت", f"فایل در {save_path} ذخیره شد.")
                return True
            else:
                messagebox.showerror("خطا", f"خطا در دانلود: {response.text}")
                return False
        except Exception as e:
            messagebox.showerror("خطا", f"ارتباط با سرور: {str(e)}")
            return False
        
    def call_api_generate(self, request_id):
        api_url = f"http://{self.api_server}:5000/generate"
        try:
            response = requests.post(api_url, json={'request_id': request_id}, timeout=120)
            if response.status_code == 200:
                return True, response.json()
            else:
                return False, response.text
        except Exception as e:
            return False, str(e)
    def download_file_from_api(self, file_id):
        """دریافت فایل از API و ذخیره در کامپیوتر کاربر"""
        api_url = f"http://SERVER_IP:5000/download/{file_id}"
        try:
            response = requests.get(api_url, stream=True)
            if response.status_code == 200:
                # استخراج نام فایل از هدر
                content_disp = response.headers.get('content-disposition', '')
                filename = "downloaded_file.pdf"
                if 'filename=' in content_disp:
                    filename = content_disp.split('filename=')[1].strip('"')
                save_path = filedialog.asksaveasfilename(initialfile=filename, defaultextension=".pdf")
                if save_path:
                    with open(save_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    messagebox.showinfo("موفقیت", f"فایل در {save_path} ذخیره شد.")
                return True
            else:
                messagebox.showerror("خطا", f"خطا در دانلود: {response.text}")
                return False
        except Exception as e:
            messagebox.showerror("خطا", f"ارتباط با سرور: {str(e)}")
            return False
        
    def hide_main_ui(self):
        # مخفی کردن تمام بخش‌های اصلی برنامه
        self.menu_frame.pack_forget()
        self.logo_frame.pack_forget()
        self.main_container.pack_forget()
        self.nav_frame.pack_forget()
        self.progress_frame.pack_forget()

    def show_main_ui(self):
        self.menu_frame.pack(fill="x", padx=0, pady=0)
        self.logo_frame.pack(pady=(20, 10), padx=20, fill="x")
        self.main_container.pack(pady=15, padx=20, fill="both", expand=True)
        self.nav_frame.pack(pady=15, padx=20, fill="x")
        self.progress_frame.pack(pady=(5, 10), padx=20, fill="x")
    
    def regenerate_and_download(self, letter_number):
        """درخواست بازسازی فایل برای یک شماره نامه خاص از سرور و سپس دانلود"""
        api_url = f"http://{self.api_server}:5000/regenerate"
        try:
            # درخواست بازسازی
            response = requests.post(api_url, json={'letter_number': letter_number}, timeout=60)
            if response.status_code == 200:
                data = response.json()
                file_id = data.get('file_id')
                if file_id:
                    # دانلود فایل
                    self.download_final_file_from_api(file_id)
                else:
                    messagebox.showerror("خطا", "فایل ساخته نشد.")
            else:
                messagebox.showerror("خطا", f"خطا در بازسازی: {response.text}")
        except Exception as e:
            messagebox.showerror("خطا", f"ارتباط با سرور: {str(e)}")

    def refresh_service_categories(self):
        self.service_categories = self.db.get_all_service_categories()
        self.service_category_tree = self.db.get_service_category_tree()
        # برای کامبوباکس، فقط نام سرویس‌های سطح والد یا همه را می‌توانید بگیرید
        self.service_names = [cat['name'] for cat in self.service_categories if cat['parent_id'] is None]

    def manage_services(self):
        dialog = ManageServicesDialog(self, self.db, self.service_categories)
        self.wait_window(dialog)
        self.refresh_service_categories()  # به‌روزرسانی پس از بستن دیالوگ

    def get_service_names(self):
        """دریافت لیست نام سرویس‌های فعال برای استفاده در کامبو باکس"""
        try:
            categories = self.db.get_all_service_categories()  # از SQLServerDatabase
            return [cat['name'] for cat in categories if cat['is_active']]
        except Exception as e:
            # اگر جدول هنوز وجود ندارد یا خطایی رخ داد، لیست پیش‌فرض برگردان
            return ["ویزا", "اتوبوس", "قطار", "گشت", "تور", "CIP"]
        
    def rebuild_menu(self):
        """حذف و بازسازی دکمه‌های منوی بالا بر اساس مجوزهای فعلی"""
        # حذف تمام ویجت‌های موجود در menu_frame به جز دکمه logout (در صورت نیاز)
        for widget in self.menu_frame.winfo_children():
            if widget != self.btn_logout:  # دکمه خروج را نگه دار
                widget.destroy()
        
        # دکمه خروج از قبل وجود دارد، فقط آن را در سمت چپ نگه می‌داریم
        self.btn_logout.pack(side="left", padx=10, pady=5)
        
        # ساخت مجدد دکمه‌های سمت راست
        right_frame = ctk.CTkFrame(self.menu_frame, fg_color="transparent")
        right_frame.pack(side="right", padx=10, pady=5)
        
        self.btn_help = ctk.CTkButton(right_frame, text="📘 راهنما", width=100, command=self.show_help,
                                    font=self.font_farsi, fg_color=COLOR_HELP, hover_color="#1976D2")
        self.btn_help.pack(side="right", padx=5)
        
        self.btn_about = ctk.CTkButton(right_frame, text="ℹ️ درباره ما", width=120, command=self.show_about,
                                    font=self.font_farsi, fg_color=COLOR_ABOUT, hover_color="#7B1FA2")
        self.btn_about.pack(side="right", padx=5)
        
        if self.can_manage_companies:
            self.btn_companies = ctk.CTkButton(right_frame, text="🏢 مدیریت طرف حساب‌ها", width=180,
                                            command=self.manage_payers, font=self.font_farsi,
                                            fg_color=COLOR_COMPANY, hover_color="#388E3C")
            self.btn_companies.pack(side="right", padx=5)
        
        if self.can_manage_settings:
            self.btn_settings = ctk.CTkButton(right_frame, text="⚙️ تنظیمات سیستم", width=140,
                                            command=self.open_settings, font=self.font_farsi,
                                            fg_color=COLOR_SETTINGS, hover_color="#F57C00")
            self.btn_settings.pack(side="right", padx=5)
        
        if self.can_approve or self.can_submit:
            self.btn_workflow = ctk.CTkButton(right_frame, text="📋 کارتابل گردش کار", width=170,
                                            command=self.open_workflow, font=self.font_farsi,
                                            fg_color=COLOR_ACCENT, hover_color="#d45a1c")
            self.btn_workflow.pack(side="right", padx=5)

        if self.can_edit_template:
            self.btn_edit_template = ctk.CTkButton(right_frame, text="✏️ ویرایش قالب فاکتور", width=160,
                                                   command=self.open_template_editor, font=self.font_farsi,
                                                   fg_color="#9C27B0", hover_color="#7B1FA2")
            self.btn_edit_template.pack(side="right", padx=5)

    def open_template_editor(self):
        if not self.can_edit_template:
            messagebox.showerror("خطا", "شما مجوز ویرایش قالب را ندارید.")
            return
        try:
            import webview
            editor_url = f"http://{self.api_server}:5000/template-editor"
            webview.create_window('ویرایشگر قالب فاکتور', editor_url, width=1300, height=900)
            webview.start()
        except ImportError:
            messagebox.showerror("خطا", "کتابخانه PyWebView نصب نیست. لطفاً با pip install pywebview نصب کنید.")
        except Exception as e:
            messagebox.showerror("خطا", f"خطا در باز کردن ویرایشگر: {str(e)}")

if __name__ == "__main__":
    base_path = get_base_path()
    config_path = os.path.join(base_path, 'db_config.json')
    if not os.path.exists(config_path):
        default_config = {"server": ".", "database": "KKTCOIG", "username": "sa", "password": "123"}
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(default_config, f, ensure_ascii=False, indent=2)
        print(f"فایل تنظیمات در {config_path} ایجاد شد. لطفاً ویرایش کنید.")
        sys.exit(1)
    app = App()
    app.mainloop()